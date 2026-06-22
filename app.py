# -*- coding: utf-8 -*-
import os, sys, threading
from datetime import datetime, timedelta
from flask import Flask, render_template_string, request, redirect, url_for, flash, jsonify, Response, make_response
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from functools import wraps
from sqlalchemy import cast, Integer, text
from sqlalchemy.exc import IntegrityError
from collections import defaultdict

_sn_lock = threading.Lock()

def bj_now():
    return datetime.utcnow().replace(microsecond=0) + timedelta(hours=8)

# ---------- 环境路径处理 (兼容单文件打包) ----------
if getattr(sys, 'frozen', False):
    BUNDLE_DIR = sys._MEIPASS
    EXE_DIR = os.path.dirname(sys.executable)
else:
    BUNDLE_DIR = os.path.dirname(os.path.abspath(__file__))
    EXE_DIR = BUNDLE_DIR

app = Flask(__name__, static_folder=os.path.join(BUNDLE_DIR, 'static'))
app.config['SECRET_KEY'] = 'ledger-2026-secret'
app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{os.path.join(EXE_DIR, "data.db")}'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
    'connect_args': {
        'check_same_thread': False,
    },
    'pool_pre_ping': True,
}

db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'

# ---------- 全局业务变量 ----------
DEPT_OPTIONS = ['研发','工艺','采购','工智','市场部','管控中心','质量部']
SHORTAGE_TYPES = ['原材料','外购件','机加件','复材件','外协件']

@app.context_processor
def inject_global_vars():
    return dict(dept_options=DEPT_OPTIONS, shortage_types=SHORTAGE_TYPES)

# ---------- 列定义逻辑 ----------
COLUMN_ORDER_RAW = [
    ('serial_no','序号'),('liquidated_damages','违约金'),('help_letter','求援函项'),
    ('first_article','首件鉴定'),('demand_no','需求单编号'),('responsible_person','负责人'),
    ('product_category','品种'), 
    ('stage','阶段'),
    ('product_model','产品型号'),('product_draw_no','产品图号'),('product_name','产品名称'),
    ('total_qty','总数量'),('plan_deliver_qty','计划交付数量'),('fixed_check_qty','定检数量'),
    ('batch_no','批次号'),('check_party','验收方'),('plan_attribute','计划属性'),
    ('contract_no','合同编号'),('plan_issue_time','主计划下达时间'),('plan_delivery_time','主计划要求交付时间'),
    ('specific_model','具体号型'),('plan_source','计划来源'),('contract_no2','合同编号'),
    ('demander','需求人'),('customer_name','客户名称'),('project_no','机型/项目流程编号'),
    ('unit_price','单价(万元)'),('output_value','产值'),
    ('production_time', '出产时间'),
    ('matching_time','配套日期'),
    ('cut_start','裁剪开始'),('cut_end','裁剪结束'),
    ('sew_start','缝纫开始'),('sew_end','缝纫结束'),
    ('glue_start','粘胶开始'),('glue_end','粘胶结束'),
    ('assembly1_start','总装1开始'),('assembly1_end','总装1结束'),
    ('assembly2_start','总装2开始'),('assembly2_end','总装2结束'),
    ('oxygen_start','氧调开始'),('oxygen_end','氧调结束'),
    ('heat_seal_start','热风热合开始'),('heat_seal_end','热风热合结束'),
    ('estimated_finish_time','预计提交总检时间'),('production_status','生产情况'),
    ('final_check_time','交总检日期'),('fixed_check_deliver_time','送定检日期'),
    ('fixed_check_finish_time','定检完成日期'),('military_check_plan_time','报军检计划时间'),
    ('military_check_time','军检时间'),('fixed_submit_item','固定提交项'),
    ('fixed_submit_finish_time','固定提交完成日期'),('storage_time','入库时间'),('remark','备注')
]

def reorder_columns():
    new_order = [('serial_no','序号'), ('responsible_person','负责人'), ('product_category','品种'), ('stage','阶段')]
    core_fields = ['product_model','product_draw_no','product_name','total_qty','plan_deliver_qty','fixed_check_qty']
    for f,n in COLUMN_ORDER_RAW:
        if f in core_fields: new_order.append((f,n))
    new_order.append(('batch_no','批次号'))
    attr_fields = ['liquidated_damages','help_letter','first_article','demand_no']
    for f,n in COLUMN_ORDER_RAW:
        if f in attr_fields: new_order.append((f,n))
    for f,n in COLUMN_ORDER_RAW:
        if (f,n) not in new_order: new_order.append((f,n))
    return new_order

COLUMN_ORDER = reorder_columns()
CENTER_COLUMNS = {'liquidated_damages','help_letter','first_article','total_qty','plan_deliver_qty','fixed_check_qty','fixed_submit_item'}

# ---------- 数据库模型 ----------
class Role(db.Model):
    __tablename__ = 'roles'
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(50), unique=True)

class User(UserMixin, db.Model):
    __tablename__ = 'users'
    id = db.Column(db.Integer, primary_key=True)
    account = db.Column(db.String(50), unique=True)
    username = db.Column(db.String(50))
    password_hash = db.Column(db.String(256))
    role_id = db.Column(db.Integer, db.ForeignKey('roles.id'))
    role = db.relationship('Role', backref='users')
    team = db.Column(db.String(50))
    production_line = db.Column(db.String(50))

class ColumnPermission(db.Model):
    __tablename__ = 'column_permissions'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id')) 
    column_name = db.Column(db.String(50))
    can_view = db.Column(db.Boolean, default=False)
    can_edit = db.Column(db.Boolean, default=False)

class TeamMember(db.Model):
    __tablename__ = 'team_members'
    id = db.Column(db.Integer, primary_key=True)
    team = db.Column(db.String(50))
    name = db.Column(db.String(50))

class ProductionTask(db.Model):
    __tablename__ = 'production_tasks'
    id = db.Column(db.Integer, primary_key=True)
    serial_no = db.Column(db.String(50)); liquidated_damages = db.Column(db.String(5)); help_letter = db.Column(db.String(5)); first_article = db.Column(db.String(5))
    demand_no = db.Column(db.String(50)); responsible_person = db.Column(db.String(50)); product_category = db.Column(db.String(50)); stage = db.Column(db.String(50)); product_model = db.Column(db.String(50)); product_draw_no = db.Column(db.String(50)); product_name = db.Column(db.String(100))
    total_qty = db.Column(db.Integer); plan_deliver_qty = db.Column(db.Integer); fixed_check_qty = db.Column(db.Integer)
    batch_no = db.Column(db.String(50)); check_party = db.Column(db.String(50)); plan_attribute = db.Column(db.String(50)); contract_no = db.Column(db.String(100))
    plan_issue_time = db.Column(db.DateTime); plan_delivery_time = db.Column(db.DateTime); specific_model = db.Column(db.String(100)); plan_source = db.Column(db.String(100)); contract_no2 = db.Column(db.String(100))
    demander = db.Column(db.String(50)); customer_name = db.Column(db.String(100)); project_no = db.Column(db.String(50)); unit_price = db.Column(db.Numeric(10,2)); output_value = db.Column(db.Numeric(12,2))
    production_time = db.Column(db.DateTime)
    matching_time = db.Column(db.DateTime)
    cut_start = db.Column(db.DateTime); cut_end = db.Column(db.DateTime); sew_start = db.Column(db.DateTime); sew_end = db.Column(db.DateTime)
    glue_start = db.Column(db.DateTime); glue_end = db.Column(db.DateTime); assembly1_start = db.Column(db.DateTime); assembly1_end = db.Column(db.DateTime)
    assembly2_start = db.Column(db.DateTime); assembly2_end = db.Column(db.DateTime); oxygen_start = db.Column(db.DateTime); oxygen_end = db.Column(db.DateTime)
    heat_seal_start = db.Column(db.DateTime); heat_seal_end = db.Column(db.DateTime)
    estimated_finish_time = db.Column(db.DateTime); production_status = db.Column(db.String(200))
    final_check_time = db.Column(db.DateTime); fixed_check_deliver_time = db.Column(db.DateTime); fixed_check_finish_time = db.Column(db.DateTime)
    military_check_plan_time = db.Column(db.DateTime); military_check_time = db.Column(db.DateTime); fixed_submit_item = db.Column(db.String(5)); fixed_submit_finish_time = db.Column(db.DateTime)
    storage_time = db.Column(db.DateTime); remark = db.Column(db.Text)
    operator = db.Column(db.String(200))
    created_by = db.Column(db.Integer); created_time = db.Column(db.DateTime, default=bj_now)
    updated_by = db.Column(db.Integer); updated_time = db.Column(db.DateTime, onupdate=bj_now)
    tech_mgmt_issues = db.relationship('IssueRecord', backref='task', lazy='dynamic', foreign_keys='IssueRecord.task_id')
    shortage_records = db.relationship('ShortageRecord', backref='task', lazy='dynamic', foreign_keys='ShortageRecord.task_id')

class IssueRecord(db.Model):
    __tablename__ = 'issue_records'
    id = db.Column(db.Integer, primary_key=True)
    task_id = db.Column(db.Integer, db.ForeignKey('production_tasks.id'))
    issue_type = db.Column(db.String(10))
    content = db.Column(db.Text)
    dept = db.Column(db.String(50))
    raise_time = db.Column(db.DateTime)
    finish_time = db.Column(db.DateTime)

class ShortageRecord(db.Model):
    __tablename__ = 'shortage_records'
    id = db.Column(db.Integer, primary_key=True)
    task_id = db.Column(db.Integer, db.ForeignKey('production_tasks.id'))
    shortage_type = db.Column(db.String(50))
    content = db.Column(db.Text)
    send_time = db.Column(db.DateTime)
    report_time = db.Column(db.DateTime)
    arrive_time = db.Column(db.DateTime)

class OperationLog(db.Model):
    __tablename__ = 'operation_logs'
    id = db.Column(db.Integer, primary_key=True); task_id = db.Column(db.Integer); operated_by = db.Column(db.Integer)
    operated_time = db.Column(db.DateTime, default=bj_now); field_name = db.Column(db.String(50))
    old_value = db.Column(db.Text); new_value = db.Column(db.Text); operation_type = db.Column(db.String(50))

class BranchNotice(db.Model):
    __tablename__ = 'branch_notices'
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(100))
    content = db.Column(db.Text)
    target = db.Column(db.String(20), default='全部')
    created_by = db.Column(db.Integer, db.ForeignKey('users.id'))
    created_time = db.Column(db.DateTime, default=bj_now)
    user = db.relationship('User', backref='branch_notices')

class SystemFeedback(db.Model):
    __tablename__ = 'system_feedback'
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200))
    content = db.Column(db.Text)
    category = db.Column(db.String(50))
    status = db.Column(db.String(20), default='待处理')
    reply = db.Column(db.Text)
    created_by = db.Column(db.Integer, db.ForeignKey('users.id'))
    created_time = db.Column(db.DateTime, default=bj_now)
    replied_by = db.Column(db.Integer)
    replied_time = db.Column(db.DateTime)
    user = db.relationship('User', backref='feedbacks')

@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))

def get_user_permissions():
    view, edit = set(), set()
    if current_user.is_authenticated:
        perms = ColumnPermission.query.filter_by(user_id=current_user.id).all()
        view = {p.column_name for p in perms if p.can_view}
        edit = {p.column_name for p in perms if p.can_edit}
        if current_user.role.name == '管理员':
            view = edit = {f for f,_ in COLUMN_ORDER}
    return view, edit

def compute_first_article(task):
    """Auto-compute first_article: stage change same product or production gap > 2 years."""
    draw_no = task.product_draw_no
    batch_no = task.batch_no
    if not draw_no or not batch_no:
        return '否'
    fl = batch_no.strip()[0].upper() if batch_no.strip() else ''
    if fl in 'CFSZ':
        current_stage = '试制'
    elif fl in 'DP':
        current_stage = '批产'
    else:
        return '否'
    existing = ProductionTask.query.filter(
        ProductionTask.product_draw_no == draw_no,
        ProductionTask.id != task.id if task.id else True
    ).all()
    for t in existing:
        bn = (t.batch_no or '').strip()
        if bn:
            fl2 = bn[0].upper()
            if fl2 in 'CFSZ' and current_stage == '批产':
                return '是'
            if fl2 in 'DP' and current_stage == '试制':
                return '是'
    # Production gap > 2 years: any previous storage_time vs new plan_delivery_time
    now = bj_now()
    threshold = task.plan_delivery_time or now
    for t in existing:
        if t.storage_time and (threshold - t.storage_time).days > 730:
            return '是'
    return '否'

def role_required(role_name):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if not current_user.is_authenticated:
                flash('请先登录', 'danger')
                return redirect(url_for('login'))
            if current_user.role.name == '管理员' or current_user.role.name == role_name:
                return f(*args, **kwargs)
            flash('无权限访问', 'danger')
            return redirect(url_for('login'))
        return decorated_function
    return decorator

def fmt_date(val):
    return val.strftime('%Y/%m/%d') if val else ''

def derive_stage_from_batch(batch_no):
    if not batch_no or not batch_no.strip():
        return ''
    fl = batch_no.strip()[0].upper()
    if fl in ('C', 'F', 'S', 'Z'):
        return '试制'
    if fl in ('D', 'P'):
        return '批产'
    return ''

def fmt_val(val):
    if val is None: return ''
    if isinstance(val, datetime): return fmt_date(val)
    if isinstance(val, bool): return '是' if val else '否'
    return str(val)

def parse_date_slash(date_str):
    if not date_str: return None
    date_str = date_str.strip().replace('/', '-')
    try:
        return datetime.strptime(date_str, '%Y-%m-%d')
    except:
        pass
    try:
        parts = date_str.split('-')
        if len(parts) == 3:
            y, m, d = int(parts[0]), int(parts[1]), int(parts[2])
            return datetime(y, m, d)
    except:
        pass
    return None

app.jinja_env.globals.update(fmt_date=fmt_date, fmt_val=fmt_val)

# ---------- 路由 ----------
@app.route('/login', methods=['GET','POST'])
def login():
    if request.method == 'POST':
        user = User.query.filter_by(account=request.form['account']).first()
        if user and check_password_hash(user.password_hash, request.form['password']):
            login_user(user)
            return redirect(url_for('dashboard'))
        flash('账号或密码错误','danger')
    return render_template_string(LOGIN_HTML)

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

@app.route('/')
@login_required
def index():
    return redirect(url_for('dashboard'))

@app.route('/tasks')
@login_required
def task_list():
    view_cols, edit_cols = get_user_permissions()
    view_cols.update({'id','serial_no','responsible_person','product_category','product_name','plan_delivery_time'})
    tasks = ProductionTask.query.order_by(ProductionTask.id).all()
    task_ids = [t.id for t in tasks]
    
    all_issues = IssueRecord.query.filter(IssueRecord.task_id.in_(task_ids)).all()
    issues_map = defaultdict(list)
    for issue in all_issues:
        issues_map[issue.task_id].append(issue)
        
    all_shortages = ShortageRecord.query.filter(ShortageRecord.task_id.in_(task_ids)).all()
    shortages_map = defaultdict(list)
    for sr in all_shortages:
        shortages_map[sr.task_id].append(sr)
        
    rows = []
    for t in tasks:
        # 生产状态判断逻辑
        if t.final_check_time:
            prod_status = '已完成'
        elif any([t.cut_start, t.sew_start, t.glue_start, t.assembly1_start, t.assembly2_start, t.oxygen_start, t.heat_seal_start]):
            prod_status = '生产中'
        else:
            prod_status = '未开工'
            
        processes = [
            ('裁剪', t.cut_start, t.cut_end), ('缝纫', t.sew_start, t.sew_end),
            ('粘胶', t.glue_start, t.glue_end), ('总装1', t.assembly1_start, t.assembly1_end),
            ('总装2', t.assembly2_start, t.assembly2_end), ('氧调', t.oxygen_start, t.oxygen_end),
            ('热风热合', t.heat_seal_start, t.heat_seal_end)
        ]
        
        hover_lines = []
        for p_name, p_start, p_end in processes:
            start_str = fmt_date(p_start) if p_start else ''
            end_str = fmt_date(p_end) if p_end else ''
            if p_end: stat = f"已完成 ({start_str} 至 {end_str})"
            elif p_start: stat = f"生产中 (开始: {start_str})"
            else: stat = "未开工"
            hover_lines.append(f"{p_name}: {stat}")
        
        row = {
            'id': t.id, 
            'production_status_calc': prod_status,
            'hover_text': '\n'.join(hover_lines)
        }
        
        t_issues = issues_map.get(t.id, [])
        t_shortages = shortages_map.get(t.id, [])
        row['has_tech'] = any(issue.issue_type == 'tech' and issue.finish_time is None for issue in t_issues)
        row['has_mgmt'] = any(issue.issue_type == 'mgmt' and issue.finish_time is None for issue in t_issues)
        row['has_shortage'] = any(sr.arrive_time is None for sr in t_shortages)
        row['has_issue'] = row['has_tech'] or row['has_mgmt'] or row['has_shortage']
        
        for field, _ in COLUMN_ORDER:
            if field in view_cols:
                if field == 'stage':
                    if t.stage and t.stage.strip():
                        row['stage'] = t.stage
                    else:
                        batch = t.batch_no or ''
                        if batch:
                            fl = batch[0].upper()
                            if fl in ('C','F','S','Z'):
                                row['stage'] = '试制'
                            elif fl in ('D','P'):
                                row['stage'] = '批产'
                            else:
                                row['stage'] = ''
                        else:
                            row['stage'] = ''
                else:
                    row[field] = fmt_val(getattr(t, field))
        rows.append(row)
        
    final_visible = []
    for f,n in COLUMN_ORDER:
        if f in view_cols:
            if f == 'serial_no':
                final_visible.append((f,n))
                final_visible.append(('has_tech', '技术问题'))
                final_visible.append(('has_mgmt', '管理问题'))
                final_visible.append(('has_shortage', '缺件问题'))
                final_visible.append(('production_status_calc', '生产状态'))
            else:
                final_visible.append((f,n))
                
    return render_template_string(TASK_LIST_HTML, rows=rows, columns=final_visible, edit_cols=edit_cols, center_cols=CENTER_COLUMNS)

@app.route('/api/task/new_empty', methods=['POST'])
@login_required
def new_empty_task():
    MAX_RETRIES = 5
    for attempt in range(MAX_RETRIES):
        try:
            with _sn_lock:
                task = ProductionTask()
                max_sn = db.session.query(db.func.max(cast(ProductionTask.serial_no, Integer))).scalar()
                task.serial_no = '1' if max_sn is None else str(max_sn + 1)
                task.created_by = current_user.id
                task.created_time = bj_now()
                task.first_article = '否'
                db.session.add(task)
                db.session.flush()
                db.session.add(OperationLog(
                    task_id=task.id, operated_by=current_user.id, operated_time=bj_now(),
                    operation_type='创建任务', field_name='serial_no',
                    new_value=f'序号{task.serial_no}'
                ))
                db.session.commit()
                return jsonify({'success': True, 'task_id': task.id})
        except IntegrityError:
            db.session.rollback()
            if attempt == MAX_RETRIES - 1:
                return jsonify({'error': '创建失败：序号冲突，请重试'}), 500
        except Exception:
            db.session.rollback()
            raise
    return jsonify({'error': '创建失败'}), 500

@app.route('/api/task/batch_add', methods=['POST'])
@login_required
def batch_add_tasks():
    data = request.get_json()
    text = data.get('text', '')
    if not text.strip(): return jsonify({'error': '内容为空'}), 400
    lines = [l for l in text.strip().splitlines() if l.strip()]
    if len(lines) < 2: return jsonify({'error': '至少需要表头和一行数据'}), 400
    headers = lines[0].split('\t')
    name_to_field = {name: field for field, name in COLUMN_ORDER}
    field_map = [name_to_field.get(h.strip()) for h in headers]
    if not any(field_map): return jsonify({'error': '表头中没有可识别的列'}), 400

    # Find serial_no column index
    serial_no_idx = None
    for idx, h in enumerate(headers):
        if h.strip() == '序号':
            serial_no_idx = idx
            break

    with _sn_lock:
        max_sn = db.session.query(db.func.max(cast(ProductionTask.serial_no, Integer))).scalar()
        start_sn = 1 if max_sn is None else max_sn + 1
        new_count = 0
        update_count = 0
        first_task_id = None
        for line in lines[1:]:
            cols = line.split('\t')
            task = None
            is_existing = False
            if serial_no_idx is not None and serial_no_idx < len(cols):
                sn = cols[serial_no_idx].strip()
                if sn:
                    task = ProductionTask.query.filter_by(serial_no=sn).first()
                    if not task:
                        try:
                            sn_int = int(sn)
                            task = ProductionTask.query.filter_by(serial_no=str(sn_int)).first()
                        except:
                            pass

            if task:
                is_existing = True
                update_count += 1
            else:
                task = ProductionTask()
                task.serial_no = str(start_sn)
                start_sn += 1
                task.created_by = current_user.id
                task.created_time = bj_now()
                new_count += 1

            if first_task_id is None:
                db.session.add(task)
                db.session.flush()
                first_task_id = task.id

            for idx, field in enumerate(field_map):
                if field is None: continue
                value = cols[idx].strip() if idx < len(cols) else ''
                if not value: continue
                col_type = getattr(ProductionTask, field).type
                try:
                    if isinstance(col_type, db.Boolean): new_val = value in ['1', '是', 'true']
                    elif isinstance(col_type, db.DateTime): new_val = parse_date_slash(value) if value else None
                    elif isinstance(col_type, db.Integer): new_val = int(value) if value else None
                    elif isinstance(col_type, db.Numeric): new_val = float(value) if value else None
                    else: new_val = value
                except:
                    new_val = None
                setattr(task, field, new_val)

            task.first_article = compute_first_article(task)
            task.updated_by = current_user.id
            task.updated_time = bj_now()
            db.session.add(task)
            db.session.flush()
            db.session.add(OperationLog(
                task_id=task.id, operated_by=current_user.id, operated_time=bj_now(),
                operation_type='批量新增' if not is_existing else '批量更新',
                field_name='serial_no', new_value=f'序号{task.serial_no}'
            ))
        db.session.commit()
    msg = f'新增{new_count}项'
    if update_count > 0:
        msg += f'，更新{update_count}项'
    return jsonify({'success': True, 'count': new_count + update_count, 'message': msg, 'first_id': first_task_id})

@app.route('/api/task/<int:task_id>/delete', methods=['POST'])
@login_required
@role_required('管理员')
def delete_task(task_id):
    task = ProductionTask.query.get_or_404(task_id)
    IssueRecord.query.filter_by(task_id=task_id).delete()
    ShortageRecord.query.filter_by(task_id=task_id).delete()
    db.session.add(OperationLog(
        task_id=task.id, operated_by=current_user.id, operated_time=bj_now(),
        operation_type='删除任务', field_name='serial_no',
        new_value=f'删除序号{task.serial_no}'
    ))
    db.session.delete(task)
    db.session.commit()
    return jsonify({'success': True})

@app.route('/api/task/<int:task_id>/field', methods=['PUT'])
@login_required
def update_field(task_id):
    task = ProductionTask.query.get_or_404(task_id)
    _, edit_cols = get_user_permissions()
    data = request.get_json()
    field = data.get('field')
    value = data.get('value')
    if field not in edit_cols: return jsonify({'error': '无权限'}), 403
    old_val = getattr(task, field)
    col_type = getattr(ProductionTask, field).type
    try:
        if isinstance(col_type, db.Boolean): new_val = value in ['1', '是', 'true']
        elif isinstance(col_type, db.DateTime): new_val = parse_date_slash(value) if value else None
        elif isinstance(col_type, db.Integer): new_val = int(value) if value else None
        elif isinstance(col_type, db.Numeric): new_val = float(value) if value else None
        else: new_val = value
    except Exception as e:
        return jsonify({'error': f'格式错误: {str(e)}'}), 400
    if str(old_val) != str(new_val):
        setattr(task, field, new_val)
        task.updated_by = current_user.id
        task.updated_time = bj_now()
        db.session.add(OperationLog(task_id=task.id, operated_by=current_user.id, field_name=field, old_value=str(old_val), new_value=str(new_val), operation_type='编辑字段'))
        if field in ('product_draw_no', 'batch_no', 'plan_delivery_time', 'storage_time'):
            task.first_article = compute_first_article(task)
        if field == 'batch_no':
            task.stage = derive_stage_from_batch(new_val)
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()
            return jsonify({'error': '保存失败，请重试'}), 500
    resp = {'success': True, 'new_value': fmt_val(new_val)}
    if field in ('product_draw_no', 'batch_no', 'plan_delivery_time', 'storage_time'):
        resp['first_article'] = task.first_article
    if field == 'batch_no':
        resp['stage'] = task.stage
    return jsonify(resp)

@app.route('/api/task/batch_update', methods=['POST'])
@login_required
def batch_update():
    data = request.get_json()
    updates = data.get('updates', [])
    _, edit_cols = get_user_permissions()
    success_count = 0
    results = []
    for item in updates:
        task_id = item['task_id']
        field = item['field']
        value = item['value']
        if field not in edit_cols: continue
        task = ProductionTask.query.get(task_id)
        if not task: continue
        col_type = getattr(ProductionTask, field).type
        try:
            if isinstance(col_type, db.Boolean): new_val = value in ['1', '是', 'true']
            elif isinstance(col_type, db.DateTime): new_val = parse_date_slash(value) if value else None
            elif isinstance(col_type, db.Integer): new_val = int(value) if value else None
            elif isinstance(col_type, db.Numeric): new_val = float(value) if value else None
            else: new_val = value
        except: continue
        old_val = getattr(task, field)
        if str(old_val) != str(new_val):
            setattr(task, field, new_val)
            task.updated_by = current_user.id
            task.updated_time = bj_now()
            db.session.add(OperationLog(task_id=task.id, operated_by=current_user.id, field_name=field, old_value=str(old_val), new_value=str(new_val), operation_type='编辑字段'))
            if field in ('product_draw_no', 'batch_no', 'plan_delivery_time', 'storage_time'):
                task.first_article = compute_first_article(task)
            if field == 'batch_no':
                task.stage = derive_stage_from_batch(new_val)
            success_count += 1
            item_result = {'task_id': task_id, 'field': field, 'new_value': fmt_val(new_val)}
            if field in ('product_draw_no', 'batch_no', 'plan_delivery_time', 'storage_time'):
                item_result['first_article'] = task.first_article
            if field == 'batch_no':
                item_result['stage'] = task.stage
            results.append(item_result)
    try:
        db.session.commit()
    except Exception:
        db.session.rollback()
        return jsonify({'error': '保存失败，请重试'}), 500
    return jsonify({'success': True, 'updated': success_count, 'results': results})

@app.route('/api/task/batch_restore', methods=['POST'])
@login_required
def batch_restore():
    data = request.get_json()
    updates = data.get('updates', [])
    _, edit_cols = get_user_permissions()
    restored = 0
    results = []
    for item in updates:
        task_id = item['task_id']
        field = item['field']
        old_value = item['old_value']
        if field not in edit_cols: continue
        task = ProductionTask.query.get(task_id)
        if not task: continue
        col_type = getattr(ProductionTask, field).type
        try:
            if isinstance(col_type, db.Boolean): new_val = old_value in ['1', '是', 'true']
            elif isinstance(col_type, db.DateTime): new_val = parse_date_slash(old_value) if old_value else None
            elif isinstance(col_type, db.Integer): new_val = int(old_value) if old_value else None
            elif isinstance(col_type, db.Numeric): new_val = float(old_value) if old_value else None
            else: new_val = old_value
        except: continue
        prev_val = getattr(task, field)
        setattr(task, field, new_val)
        task.updated_by = current_user.id
        task.updated_time = bj_now()
        db.session.add(OperationLog(task_id=task.id, operated_by=current_user.id, field_name=field, old_value=str(prev_val or ''), new_value=str(new_val or ''), operation_type='撤销恢复'))
        restored += 1
        results.append({'task_id': task_id, 'field': field, 'new_value': fmt_val(new_val)})
    try:
        db.session.commit()
    except Exception:
        db.session.rollback()
        return jsonify({'error': '撤销恢复失败，请重试'}), 500
    return jsonify({'success': True, 'restored': restored, 'results': results})

@app.route('/export/doc')
@login_required
def export_doc():
    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from io import BytesIO

    view_cols, _ = get_user_permissions()
    tasks = ProductionTask.query.order_by(ProductionTask.id).all()
    visible_fields = [(f, n) for f, n in COLUMN_ORDER if f in view_cols]

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = '任务列表'
    thin = Side(style='thin')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    header_fill = PatternFill(start_color='F2F2F2', end_color='F2F2F2', fill_type='solid')
    header_font = Font(bold=True, size=11)
    cell_font = Font(size=11)
    center_align = Alignment(horizontal='center', vertical='center')

    for ci, (f, n) in enumerate(visible_fields, 1):
        cell = ws.cell(row=1, column=ci, value=n)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = border
        cell.alignment = center_align

    for ri, t in enumerate(tasks, 2):
        for ci, (f, n) in enumerate(visible_fields, 1):
            if f == 'stage':
                batch = t.batch_no or ''
                if batch:
                    fl = batch[0].upper()
                    if fl in ('C','F','S','Z'): val = '试制'
                    elif fl in ('D','P'): val = '批产'
                    else: val = ''
                else: val = ''
            else:
                val = fmt_val(getattr(t, f))
            cell = ws.cell(row=ri, column=ci, value=val)
            cell.font = cell_font
            cell.border = border
            cell.alignment = center_align

    for ci in range(1, len(visible_fields) + 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(ci)].width = 14

    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return Response(output.getvalue(), mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', headers={'Content-Disposition': 'attachment;filename=tasks_export.xlsx'})

@app.route('/task/<int:task_id>/details', methods=['GET','POST'])
@login_required
def task_details(task_id):
    task = ProductionTask.query.get_or_404(task_id)
    if request.method == 'POST':
        for key in request.form:
            if key.startswith('issue_id_'):
                is_id = request.form[key]
                if is_id:
                    obj = IssueRecord.query.get(int(is_id))
                    if obj and obj.task_id == task.id:
                        obj.issue_type = request.form.get(f'issue_type_{obj.id}')
                        obj.content = request.form.get(f'issue_content_{obj.id}','')
                        obj.dept = request.form.get(f'issue_dept_{obj.id}','')
                        obj.raise_time = parse_date_slash(request.form.get(f'issue_raise_{obj.id}'))
                        obj.finish_time = parse_date_slash(request.form.get(f'issue_finish_{obj.id}'))
        new_types = request.form.getlist('new_issue_type[]')
        for i,t in enumerate(new_types):
            if t:
                db.session.add(IssueRecord(
                    task_id=task.id, issue_type=t,
                    content=request.form.getlist('new_issue_content[]')[i],
                    dept=request.form.getlist('new_issue_dept[]')[i],
                    raise_time=parse_date_slash(request.form.getlist('new_issue_raise[]')[i]),
                    finish_time=parse_date_slash(request.form.getlist('new_issue_finish[]')[i])))
        for did in request.form.get('delete_issues','').split(','):
            if did: IssueRecord.query.filter_by(id=int(did), task_id=task.id).delete()
        for key in request.form:
            if key.startswith('sh_id_'):
                sid = request.form[key]
                if sid:
                    obj = ShortageRecord.query.get(int(sid))
                    if obj and obj.task_id == task.id:
                        obj.shortage_type = request.form.get(f'sh_type_{obj.id}')
                        obj.content = request.form.get(f'sh_content_{obj.id}','')
                        obj.send_time = parse_date_slash(request.form.get(f'sh_send_{obj.id}'))
                        obj.report_time = parse_date_slash(request.form.get(f'sh_report_{obj.id}'))
                        obj.arrive_time = parse_date_slash(request.form.get(f'sh_arrive_{obj.id}'))
        new_sh = request.form.getlist('new_sh_type[]')
        for i,st in enumerate(new_sh):
            if st:
                db.session.add(ShortageRecord(
                    task_id=task.id, shortage_type=st,
                    content=request.form.getlist('new_sh_content[]')[i],
                    send_time=parse_date_slash(request.form.getlist('new_sh_send[]')[i]),
                    report_time=parse_date_slash(request.form.getlist('new_sh_report[]')[i]),
                    arrive_time=parse_date_slash(request.form.getlist('new_sh_arrive[]')[i])))
        for did in request.form.get('delete_shortages','').split(','):
            if did: ShortageRecord.query.filter_by(id=int(did), task_id=task.id).delete()
        now_val = bj_now()
        for did in request.form.get('delete_issues','').split(','):
            if did: db.session.add(OperationLog(task_id=task.id, operated_by=current_user.id, operated_time=now_val, field_name='issue_records', old_value=did, new_value='(已删除)', operation_type='删除问题'))
        for did in request.form.get('delete_shortages','').split(','):
            if did: db.session.add(OperationLog(task_id=task.id, operated_by=current_user.id, operated_time=now_val, field_name='shortage_records', old_value=did, new_value='(已删除)', operation_type='删除缺件'))
        new_types = request.form.getlist('new_issue_type[]')
        for i,t in enumerate(new_types):
            if t:
                db.session.add(OperationLog(task_id=task.id, operated_by=current_user.id, operated_time=now_val, field_name='issue_records', old_value='', new_value=request.form.getlist('new_issue_content[]')[i] or '', operation_type='新增问题'))
        new_sh = request.form.getlist('new_sh_type[]')
        for i,st in enumerate(new_sh):
            if st:
                db.session.add(OperationLog(task_id=task.id, operated_by=current_user.id, operated_time=now_val, field_name='shortage_records', old_value='', new_value=request.form.getlist('new_sh_content[]')[i] or '', operation_type='新增缺件'))
        db.session.commit()
        flash('详情已保存','success')
        return redirect(url_for('task_details', task_id=task.id))
    today = bj_now().strftime('%Y/%m/%d')
    return render_template_string(DETAILS_HTML, task=task, today=today)

# ---------- 看板路由 ----------
@app.route('/dashboard')
@login_required
def dashboard():
    now = bj_now()
    try: selected_year = int(request.args.get('year', now.year))
    except: selected_year = now.year
    try: selected_month = int(request.args.get('month', now.month))
    except: selected_month = now.month

    month_start = datetime(selected_year, selected_month, 1)
    if selected_month == 12: month_end = datetime(selected_year + 1, 1, 1)
    else: month_end = datetime(selected_year, selected_month + 1, 1)

    tasks = ProductionTask.query.all()
    task_ids = [t.id for t in tasks]

    all_issues = IssueRecord.query.filter(IssueRecord.task_id.in_(task_ids)).all()
    issues_map = defaultdict(list)
    for issue in all_issues: issues_map[issue.task_id].append(issue)

    all_shortages = ShortageRecord.query.filter(ShortageRecord.task_id.in_(task_ids)).all()
    shortages_map = defaultdict(list)
    for sr in all_shortages: shortages_map[sr.task_id].append(sr)

    def get_issues(tid): return issues_map.get(tid, [])
    def get_shortages(tid): return shortages_map.get(tid, [])
    
    current_month_tasks = [t for t in tasks if t.plan_delivery_time and month_start <= t.plan_delivery_time < month_end]
    total_monthly = len(current_month_tasks)
    
    past_tasks = [t for t in tasks if t.plan_delivery_time and t.plan_delivery_time < month_start]
    past_unfinished_tasks = [t for t in past_tasks if not t.storage_time]
    past_unfinished_count = len(past_unfinished_tasks)
    
    base_total = total_monthly + past_unfinished_count
    
    on_time_tasks = [t for t in current_month_tasks if t.storage_time and month_start <= t.storage_time < month_end]
    on_time = len(on_time_tasks)
    rate1 = round(on_time / base_total * 100, 1) if base_total else 0
    on_time_output = sum(float(t.output_value or 0) for t in on_time_tasks)

    done_tasks = [t for t in current_month_tasks if t.storage_time]
    done_count = len(done_tasks)
    rate2 = round(done_count / base_total * 100, 1) if base_total else 0
    done_output = sum(float(t.output_value or 0) for t in done_tasks)

    past_done_tasks = [t for t in past_tasks if t.storage_time]
    past_done = len(past_done_tasks)
    rate3 = round(past_done / base_total * 100, 1) if base_total else 0
    delayed_done_output = sum(float(t.output_value or 0) for t in past_done_tasks)

    penalty_tasks = [t for t in tasks if t.plan_delivery_time and month_start <= t.plan_delivery_time < month_end and not t.storage_time and t.liquidated_damages == '是']
    penalty_count = len(penalty_tasks)
    penalty_amount = sum(float(t.output_value or 0) for t in penalty_tasks)

    monthly_output_value = sum(float(t.output_value or 0) for t in current_month_tasks)
    not_on_time_output = monthly_output_value - on_time_output

    done_not_stored_tasks = [t for t in tasks if t.final_check_time and not t.storage_time]
    deadline = now + timedelta(days=7)
    warnings = [t for t in tasks if t.plan_delivery_time and t.plan_delivery_time <= deadline and not t.storage_time]

    quality_list = []

    for t in current_month_tasks:
        t_issues = get_issues(t.id)
        t_shortages = get_shortages(t.id)
        if not t_issues and not t_shortages:
            continue
        t_val = float(t.output_value or 0)
        
        # Production status
        if t.final_check_time: prod_status = '已完成'
        elif any([t.cut_start, t.sew_start, t.glue_start, t.assembly1_start, t.assembly2_start, t.oxygen_start, t.heat_seal_start]): prod_status = '生产中'
        else: prod_status = '未开工'
        
        # Compute max delay and nature
        max_days = -1
        max_nature = ''
        max_detail = ''
        for iss in t_issues:
            if iss.raise_time:
                d = ((iss.finish_time or now) - iss.raise_time).days
                d = max(0, d)
                nature = iss.dept or ('技术' if iss.issue_type=='tech' else '管理')
                detail = f"[技术/管理] {iss.dept or ''} | 内容:{iss.content or ''} | 提出:{fmt_date(iss.raise_time)} | 完成:{fmt_date(iss.finish_time or now)} | 延期:{d}天"
                if d >= max_days: max_days = d; max_nature = nature; max_detail = detail
        for sr in t_shortages:
            if sr.report_time:
                d = ((sr.arrive_time or now) - sr.report_time).days
                d = max(0, d)
                nature = sr.shortage_type or ''
                detail = f"[物料缺件] {nature} | 内容:{sr.content or ''} | 报缺:{fmt_date(sr.report_time)} | 到位:{fmt_date(sr.arrive_time or now)} | 延期:{d}天"
                if d >= max_days: max_days = d; max_nature = nature; max_detail = detail
        if max_days < 0: max_days = 0; max_nature = '无'
        
        quality_list.append({
            'id': t.id, 'serial_no': t.serial_no,
            'product_category': t.product_category or '',
            'product_draw_no': t.product_draw_no or '',
            'product_name': t.product_name,
            'batch_no': t.batch_no or '',
            'total_qty': t.total_qty,
            'prod_status': prod_status,
            'days': max_days,
            'nature': max_nature,
            'output_value': t_val,
            'detail': max_detail
        })

    monthly_issues = [iss for iss in all_issues if iss.raise_time and month_start <= iss.raise_time < month_end]
    monthly_shortages = [sh for sh in all_shortages if sh.report_time and month_start <= sh.report_time < month_end]

    task_dict = {t.id: t for t in tasks}
    m_issues_data = []
    for iss in monthly_issues:
        tk = task_dict.get(iss.task_id)
        if tk:
            m_issues_data.append({'serial_no': tk.serial_no, 'product_draw_no': tk.product_draw_no, 'product_name': tk.product_name, 'batch_no': tk.batch_no, 'total_qty': tk.total_qty, 'issue_type': iss.issue_type, 'raise_time': iss.raise_time})
    m_issues_data.sort(key=lambda x: x['raise_time'] or datetime.min, reverse=True)

    m_shortages_data = []
    for sh in monthly_shortages:
        tk = task_dict.get(sh.task_id)
        if tk:
            m_shortages_data.append({'serial_no': tk.serial_no, 'product_draw_no': tk.product_draw_no, 'product_name': tk.product_name, 'batch_no': tk.batch_no, 'total_qty': tk.total_qty, 'shortage_type': sh.shortage_type, 'report_time': sh.report_time})
    m_shortages_data.sort(key=lambda x: x['report_time'] or datetime.min, reverse=True)

    in_progress_detail = {
        '裁剪': [t for t in tasks if t.cut_start and not t.cut_end], '缝纫': [t for t in tasks if t.sew_start and not t.sew_end],
        '粘胶': [t for t in tasks if t.glue_start and not t.glue_end], '总装1': [t for t in tasks if t.assembly1_start and not t.assembly1_end],
        '总装2': [t for t in tasks if t.assembly2_start and not t.assembly2_end], '氧调': [t for t in tasks if t.oxygen_start and not t.oxygen_end],
        '热风热合': [t for t in tasks if t.heat_seal_start and not t.heat_seal_end],
    }
    in_progress_for_template = {}
    for team, tlist in in_progress_detail.items():
        items = []
        for t in tlist:
            st = None
            if team == '裁剪': st = t.cut_start
            elif team == '缝纫': st = t.sew_start
            elif team == '粘胶': st = t.glue_start
            elif team == '总装1': st = t.assembly1_start
            elif team == '总装2': st = t.assembly2_start
            elif team == '氧调': st = t.oxygen_start
            elif team == '热风热合': st = t.heat_seal_start
            items.append({'serial_no': t.serial_no, 'product_draw_no': t.product_draw_no or '', 'product_name': t.product_name, 'specific_model': t.specific_model or '', 'batch_no': t.batch_no or '', 'total_qty': t.total_qty, 'start_time': st})
        in_progress_for_template[team] = items

    monthly_delivery_list = []
    for t in current_month_tasks:
        monthly_delivery_list.append({'serial_no': t.serial_no, 'product_draw_no': t.product_draw_no or '', 'product_name': t.product_name, 'batch_no': t.batch_no or '', 'specific_model': t.specific_model or '', 'total_qty': t.total_qty, 'output_value': float(t.output_value or 0), 'status': '已完成' if t.final_check_time else '未完成'})

    year_range = list(range(2024, now.year + 6))
    
    return render_template_string(DASHBOARD_HTML, rate1=rate1, on_time_output=on_time_output, rate2=rate2, done_output=done_output, rate3=rate3, delayed_done_output=delayed_done_output, penalty_count=penalty_count, penalty_amount=penalty_amount, penalty_tasks=penalty_tasks, not_on_time_output=not_on_time_output, quality_list=quality_list, in_progress_for_template=in_progress_for_template, monthly_delivery_list=monthly_delivery_list, done_not_stored_tasks=done_not_stored_tasks, warnings=warnings, m_issues_data=m_issues_data, m_shortages_data=m_shortages_data, on_time=on_time, total_monthly=total_monthly, selected_year=selected_year, selected_month=selected_month, year_range=year_range)

# ---------- 质量问题一本账导出 ----------
@app.route('/export/quality')
@login_required
def export_quality():
    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side
    from io import BytesIO
    import traceback
    
    try:
        now = bj_now()
        try: selected_year = int(request.args.get('year', now.year))
        except: selected_year = now.year
        try: selected_month = int(request.args.get('month', now.month))
        except: selected_month = now.month
        month_start = datetime(selected_year, selected_month, 1)
        if selected_month == 12: month_end = datetime(selected_year + 1, 1, 1)
        else: month_end = datetime(selected_year, selected_month + 1, 1)
        
        tasks = ProductionTask.query.all()
        plan_monthly = [t for t in tasks if t.plan_delivery_time and month_start <= t.plan_delivery_time < month_end]
        all_issues = IssueRecord.query.all()
        all_shortages = ShortageRecord.query.all()
        im = defaultdict(list); sm = defaultdict(list)
        for iss in all_issues: im[iss.task_id].append(iss)
        for sr in all_shortages: sm[sr.task_id].append(sr)
        
        rows = []
        for t in plan_monthly:
            t_issues = im.get(t.id, [])
            t_shortages = sm.get(t.id, [])
            if not t_issues and not t_shortages: continue
            t_val = float(t.output_value or 0)
            if t.final_check_time: prod_status = '已完成'
            elif any([t.cut_start, t.sew_start, t.glue_start, t.assembly1_start, t.assembly2_start, t.oxygen_start, t.heat_seal_start]): prod_status = '生产中'
            else: prod_status = '未开工'
            max_days = -1; max_nature = ''; max_detail = ''
            for iss in t_issues:
                if iss.raise_time:
                    d = max(0, ((iss.finish_time or now) - iss.raise_time).days)
                    nature = iss.dept or ('技术' if iss.issue_type=='tech' else '管理')
                    detail = f'[技术/管理] {iss.dept or ""} | 内容:{iss.content or ""} | 提出:{fmt_date(iss.raise_time)} | 完成:{fmt_date(iss.finish_time or now)} | 延期:{d}天'
                    if d >= max_days: max_days = d; max_nature = nature; max_detail = detail
            for sr in t_shortages:
                if sr.report_time:
                    d = max(0, ((sr.arrive_time or now) - sr.report_time).days)
                    nature = sr.shortage_type or ''
                    detail = f'[物料缺件] {nature} | 内容:{sr.content or ""} | 报缺:{fmt_date(sr.report_time)} | 到位:{fmt_date(sr.arrive_time or now)} | 延期:{d}天'
                    if d >= max_days: max_days = d; max_nature = nature; max_detail = detail
            if max_days < 0: max_days = 0; max_nature = '无'
            rows.append([t.serial_no, prod_status, t.product_category or '', t.product_draw_no or '',
                t.batch_no or '', str(t.total_qty or ''), f'{max_days}天', max_nature, f'{t_val:.2f}', max_detail])
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = '问题台账'
        headers = ['序号','生产状态','品种','图号','批次','数量','用时(天)','定性','产值(W)','详情']
        thin = Side(style='thin')
        border = Border(left=thin,right=thin,top=thin,bottom=thin)
        hf = Font(bold=True, size=11); cf = Font(size=11)
        ca = Alignment(horizontal='center', vertical='center', wrap_text=True)
        for ci,h in enumerate(headers,1):
            c = ws.cell(row=1,column=ci,value=h)
            c.font = hf; c.border = border; c.alignment = ca
        for ri,row in enumerate(rows,2):
            for ci,val in enumerate(row,1):
                c = ws.cell(row=ri,column=ci,value=val)
                c.font = cf; c.border = border; c.alignment = ca
        # Auto-fit column widths
        for ci in range(1,11):
            max_len = 0
            for ri in range(1, len(rows)+2):
                val = ws.cell(row=ri,column=ci).value
                if val:
                    clen = sum(2 if ord(ch)>127 else 1 for ch in str(val))
                    max_len = max(max_len, clen)
            cap = 80 if ci == 10 else 55
            ws.column_dimensions[openpyxl.utils.get_column_letter(ci)].width = min(max_len+4, cap)
        ws.row_dimensions[1].height = 26
        for ri in range(2, len(rows)+2):
            detail_text = str(ws.cell(row=ri, column=10).value or '')
            if detail_text:
                total_w = sum(2 if ord(ch)>127 else 1 for ch in detail_text)
                col_w = ws.column_dimensions['J'].width or 55
                wrapped_lines = max(1, int(total_w / (col_w * 0.8)) + detail_text.count('\n'))
                ws.row_dimensions[ri].height = max(22, wrapped_lines * 16)
            else:
                ws.row_dimensions[ri].height = 22
        
        output = BytesIO()
        wb.save(output)
        output.seek(0)
        from urllib.parse import quote
        filename = quote(f'问题台账_{selected_year}_{selected_month}.xlsx')
        cd = "attachment;filename*=UTF-8''" + filename
        return Response(output.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={'Content-Disposition': cd})
    except Exception as e:
        return f'Export error: {str(e)}\n{traceback.format_exc()}', 500

# ---------- 违约金任务分析报告导出 ----------
@app.route('/export/penalty')
@login_required
def export_penalty():
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from io import BytesIO
    import traceback

    def set_cell_shading(cell, color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        shading.set(qn('w:val'), 'clear')
        tcPr.append(shading)

    try:
        now = bj_now()
        try: selected_year = int(request.args.get('year', now.year))
        except: selected_year = now.year
        try: selected_month = int(request.args.get('month', now.month))
        except: selected_month = now.month
        month_start = datetime(selected_year, selected_month, 1)
        if selected_month == 12: month_end = datetime(selected_year + 1, 1, 1)
        else: month_end = datetime(selected_year, selected_month + 1, 1)

        tasks = ProductionTask.query.all()
        all_penalty = [t for t in tasks if t.plan_delivery_time and month_start <= t.plan_delivery_time < month_end and t.liquidated_damages == '是']

        penalty_ids = [t.id for t in all_penalty]
        all_issues = IssueRecord.query.filter(IssueRecord.task_id.in_(penalty_ids)).all() if penalty_ids else []
        all_shortages = ShortageRecord.query.filter(ShortageRecord.task_id.in_(penalty_ids)).all() if penalty_ids else []
        im = defaultdict(list); sm = defaultdict(list)
        for iss in all_issues: im[iss.task_id].append(iss)
        for sr in all_shortages: sm[sr.task_id].append(sr)

        def calc_delay(t):
            max_d = 0
            for iss in im.get(t.id, []):
                if iss.raise_time:
                    d = max(0, ((iss.finish_time or now) - iss.raise_time).days)
                    if d > max_d: max_d = d
            for sr in sm.get(t.id, []):
                if sr.report_time:
                    d = max(0, ((sr.arrive_time or now) - sr.report_time).days)
                    if d > max_d: max_d = d
            if max_d == 0:
                max_d = max(0, ((t.storage_time or now) - t.plan_delivery_time).days)
            return max_d

        on_time_tasks = [t for t in all_penalty if t.storage_time and t.storage_time <= t.plan_delivery_time]
        delayed_done_tasks = [t for t in all_penalty if t.storage_time and t.storage_time > t.plan_delivery_time]
        unfinished_tasks = [t for t in all_penalty if not t.storage_time]

        on_time_val = sum(float(t.output_value or 0) for t in on_time_tasks)
        delayed_done_val = sum(float(t.output_value or 0) for t in delayed_done_tasks)
        unfinished_val = sum(float(t.output_value or 0) for t in unfinished_tasks)
        total_val = on_time_val + delayed_done_val + unfinished_val

        cats = {
            'tech_quality': {'count': 0, 'val': 0, 'delays': []},
            'raw_material': {'count': 0, 'val': 0, 'delays': []},
            'purchased': {'count': 0, 'val': 0, 'delays': []},
            'composite': {'count': 0, 'val': 0, 'delays': []},
            'machined': {'count': 0, 'val': 0, 'delays': []},
            'outsourced': {'count': 0, 'val': 0, 'delays': []},
            'fixed_check': {'count': 0, 'val': 0, 'delays': []},
            'contract': {'count': 0, 'val': 0, 'delays': []},
            'military_check': {'count': 0, 'val': 0, 'delays': []},
        }

        for t in unfinished_tasks:
            t_val = float(t.output_value or 0)
            has_tech = False; has_mgmt = False; tech_delays = []; mgmt_delays = []
            for iss in im.get(t.id, []):
                if iss.raise_time:
                    d = max(0, ((iss.finish_time or now) - iss.raise_time).days)
                    if iss.issue_type == 'tech': tech_delays.append(d)
                    else: mgmt_delays.append(d)

            if tech_delays or mgmt_delays:
                all_d = tech_delays + mgmt_delays
                cats['tech_quality']['count'] += 1
                cats['tech_quality']['val'] += t_val
                cats['tech_quality']['delays'].append(max(all_d))

            for stype, ckey in [('原材料','raw_material'),('外购件','purchased'),('复材件','composite'),('机加件','machined'),('外协件','outsourced')]:
                st_delays = []
                for sr in sm.get(t.id, []):
                    if sr.shortage_type == stype and sr.report_time:
                        d = max(0, ((sr.arrive_time or now) - sr.report_time).days)
                        st_delays.append(d)
                if st_delays:
                    cats[ckey]['count'] += 1
                    cats[ckey]['val'] += t_val
                    cats[ckey]['delays'].append(max(st_delays))

            if t.fixed_check_finish_time and t.plan_delivery_time and t.fixed_check_finish_time > t.plan_delivery_time:
                d = max(0, (now - t.plan_delivery_time).days)
                cats['fixed_check']['count'] += 1
                cats['fixed_check']['val'] += t_val
                cats['fixed_check']['delays'].append(d)

            if t.final_check_time:
                d = max(0, (now - t.plan_delivery_time).days)
                cats['contract']['count'] += 1
                cats['contract']['val'] += t_val
                if d > 0:
                    cats['contract']['delays'].append(d)

            if t.military_check_plan_time and not t.military_check_time:
                cats['military_check']['count'] += 1
                cats['military_check']['val'] += t_val
                d = max(0, (now - t.military_check_plan_time).days)
                cats['military_check']['delays'].append(d)

        doc = Document()
        style = doc.styles['Normal']
        style.font.size = Pt(10.5)
        style.font.name = '宋体'
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        rPr = style.element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '宋体')
        rPr.append(rFonts)

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f'{selected_month}月违约金任务分析报告')
        run.bold = True; run.font.size = Pt(16)

        doc.add_paragraph()
        doc.add_paragraph(f'一、{selected_month}月共有{len(all_penalty)}项违约金产品需要出产，涉及产值{total_val:.2f}万元，实际完成情况如下:')
        for label, tasks_set, val in [
            ('准时完成', on_time_tasks, on_time_val),
            ('延期但当月实际完成', delayed_done_tasks, delayed_done_val),
            ('未完成', unfinished_tasks, unfinished_val)
        ]:
            p = doc.add_paragraph(f'{label}{len(tasks_set)}项，涉及产值{val:.2f}万元。')
            p.paragraph_format.left_indent = Cm(0.75)

        doc.add_paragraph('二、对未完成任务进行分析，情况如下:')
        cat_labels = [
            ('tech_quality', '受技术质量问题影响'),
            ('raw_material', '受原材料缺项影响'),
            ('purchased', '受外购件缺项影响'),
            ('composite', '受复材件缺项影响'),
            ('machined', '受机加件缺项影响'),
            ('outsourced', '受外协件缺项影响'),
            ('fixed_check', '受定检影响'),
            ('contract', '受合同/监管协议影响'),
            ('military_check', '现有待军检任务'),
        ]
        for idx, (ckey, clabel) in enumerate(cat_labels):
            c = cats[ckey]
            cir_num = chr(ord('①') + idx)
            text = f'{cir_num} {clabel}{c["count"]}项，涉及产值{c["val"]:.2f}万元'
            if c['delays']:
                avg_d = sum(c['delays']) / len(c['delays'])
                max_d = max(c['delays'])
                text += f'，平均延期{avg_d:.1f}天，最长延期天数{max_d}天'
            text += '。'
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Cm(0.75)

        doc.add_paragraph('三、具体内容：')
        detail_cols = ['序号', '图号', '产品', '批次', '数量', '产值', '最长延期天数']
        table = doc.add_table(rows=1, cols=len(detail_cols))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for ci, h in enumerate(detail_cols):
            cell = table.cell(0, ci)
            cell.text = h
            set_cell_shading(cell, 'D9D9D9')
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.bold = True; run.font.size = Pt(9)

        for t in unfinished_tasks:
            t_val = float(t.output_value or 0)
            max_d = calc_delay(t)
            delay_str = f'{max_d}天' if max_d > 0 else '0天'
            row = table.add_row()
            vals = [t.serial_no or '', t.product_draw_no or '', t.product_name or '', t.batch_no or '', str(t.total_qty or ''), f'{t_val:.2f}', delay_str]
            for ci, v in enumerate(vals):
                cell = row.cells[ci]
                cell.text = v
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.line_spacing = 1.0
                    for run in p.runs:
                        run.font.size = Pt(9)

            detail_row = table.add_row()
            detail_cell = detail_row.cells[0]
            for ci in range(1, len(detail_cols)):
                detail_cell.merge(detail_row.cells[ci])
            set_cell_shading(detail_cell, 'F5F5F5')

            t_issues = im.get(t.id, [])
            t_shortages = sm.get(t.id, [])

            def add_para(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT):
                p = cell.add_paragraph()
                p.alignment = align
                p.paragraph_format.line_spacing = 1.0
                run = p.add_run(text)
                run.font.size = Pt(9)
                return p

            # Set first paragraph to '详情：'
            p0 = detail_cell.paragraphs[0]
            p0.clear()
            p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p0.paragraph_format.line_spacing = 1.0
            run0 = p0.add_run('详情：')
            run0.font.size = Pt(9)

            if not t_issues and not t_shortages:
                add_para(detail_cell, '无问题/缺件记录')
            else:
                for iss in t_issues:
                    i_type = '技术问题' if iss.issue_type == 'tech' else '管理问题'
                    add_para(detail_cell, f'【{i_type}】责任部门:{iss.dept or ""}')
                    add_para(detail_cell, f'提出时间:{fmt_date(iss.raise_time)}', WD_ALIGN_PARAGRAPH.RIGHT)
                    add_para(detail_cell, f'问题原因:{iss.content or ""}')
                    add_para(detail_cell, f'完成时间:{fmt_date(iss.finish_time or now)}', WD_ALIGN_PARAGRAPH.RIGHT)
                for sr in t_shortages:
                    add_para(detail_cell, f'【缺件-{sr.shortage_type or ""}】')
                    add_para(detail_cell, f'报缺时间:{fmt_date(sr.report_time)}', WD_ALIGN_PARAGRAPH.RIGHT)
                    add_para(detail_cell, f'问题原因:{sr.content or ""}')
                    add_para(detail_cell, f'到位时间:{fmt_date(sr.arrive_time or now)}', WD_ALIGN_PARAGRAPH.RIGHT)

        col_widths = [Cm(1.5), Cm(2.7), Cm(2.7), Cm(1.95), Cm(1.95), Cm(1.95), Cm(2.7)]
        for row in table.rows:
            for ci, width in enumerate(col_widths):
                if ci < len(row.cells):
                    row.cells[ci].width = width

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        from urllib.parse import quote
        filename = quote(f'违约金任务分析报告_{selected_year}年{selected_month}月.docx')
        cd = "attachment;filename*=UTF-8''" + filename
        return Response(output.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': cd})
    except Exception as e:
        return f'Export error: {str(e)}\n{traceback.format_exc()}', 500

# ---------- 数据中心导出 ----------
@app.route('/datacenter/export/<mid>')
@login_required
def datacenter_export(mid):
    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from io import BytesIO

    now = bj_now()
    try: selected_year = int(request.args.get('year', now.year))
    except: selected_year = now.year
    try: selected_month = int(request.args.get('month', now.month))
    except: selected_month = now.month

    month_start = datetime(selected_year, selected_month, 1)
    if selected_month == 12: month_end = datetime(selected_year + 1, 1, 1)
    else: month_end = datetime(selected_year, selected_month + 1, 1)
    year_start = datetime(selected_year, 1, 1)
    year_end = datetime(selected_year + 1, 1, 1)

    tasks = ProductionTask.query.all()
    denom6 = [t for t in tasks if t.plan_delivery_time and month_start <= t.plan_delivery_time < month_end]

    m_map = {
        'm1': [t for t in tasks if t.storage_time and month_start <= t.storage_time < month_end],
        'm2': [t for t in tasks if t.plan_delivery_time and year_start <= t.plan_delivery_time < year_end and t.storage_time],
        'm3': [t for t in tasks if t.plan_delivery_time and month_start <= t.plan_delivery_time < month_end and not t.storage_time],
        'm4': [t for t in tasks if t.plan_delivery_time and year_start <= t.plan_delivery_time < year_end and not t.storage_time],
        'm5': [t for t in tasks if t.plan_issue_time and month_start <= t.plan_issue_time < month_end],
        'm6': [t for t in denom6 if t.storage_time],
        'm7': [t for t in denom6 if t.final_check_time and t.storage_time and t.plan_delivery_time and t.storage_time <= t.plan_delivery_time],
        'm8': [t for t in tasks if t.storage_time and month_start <= t.storage_time < month_end and t.plan_delivery_time and t.storage_time <= t.plan_delivery_time],
        'm9': [t for t in tasks if t.final_check_time and year_start <= t.final_check_time < year_end],
        'm10': [t for t in tasks if t.storage_time and year_start <= t.storage_time < year_end],
        'm11': [t for t in tasks if t.fixed_check_finish_time and year_start <= t.fixed_check_finish_time < year_end],
        'm12': [t for t in tasks if t.fixed_check_deliver_time and not t.fixed_check_finish_time],
        'm13': [t for t in tasks if t.fixed_check_qty and not t.fixed_check_deliver_time and t.estimated_finish_time and (t.estimated_finish_time - now).days < 60],
        'm14': [t for t in tasks if t.plan_delivery_time and month_start <= t.plan_delivery_time < month_end and t.liquidated_damages == '是' and not t.storage_time],
        'm15': [t for t in tasks if t.plan_delivery_time and month_start <= t.plan_delivery_time < month_end and t.liquidated_damages == '是'],
        'm16': [t for t in tasks if t.plan_delivery_time and year_start <= t.plan_delivery_time < year_end and t.liquidated_damages == '是'],
        'm17': [t for t in tasks if t.plan_delivery_time and month_start <= t.plan_delivery_time < month_end],
        'm18': [t for t in tasks if t.plan_delivery_time and month_start <= t.plan_delivery_time < month_end and (t.storage_time or now) > t.plan_delivery_time],
    }

    cost_time_map = {}
    for t in m_map.get('m11', []):
        if t.fixed_check_deliver_time and t.fixed_check_finish_time:
            cost_time_map[('m11', t.id)] = f'{(t.fixed_check_finish_time - t.fixed_check_deliver_time).days}天'
    for t in m_map.get('m12', []):
        if t.fixed_check_deliver_time:
            cost_time_map[('m12', t.id)] = f'{(now - t.fixed_check_deliver_time).days}天'

    def fc(key, tid):
        return cost_time_map.get((key, tid), '')

    def compute_prod_status(t):
        if t.final_check_time: return '已完成'
        if any([t.cut_start, t.sew_start, t.glue_start, t.assembly1_start, t.assembly2_start, t.oxygen_start, t.heat_seal_start]): return '生产中'
        return '未开工'

    mc = {
        'm1': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划要求交付时间','plan_delivery_time'),('产值','output_value')],
        'm2': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划要求交付时间','plan_delivery_time'),('产值','output_value')],
        'm3': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划要求交付时间','plan_delivery_time'),('产值','output_value')],
        'm4': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划要求交付时间','plan_delivery_time'),('产值','output_value')],
        'm5': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划下达时间','plan_issue_time'),('产值','output_value')],
        'm6': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划要求交付时间','plan_delivery_time'),('交总检日期','final_check_time'),('产值','output_value')],
        'm7': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划要求交付时间','plan_delivery_time'),('交总检日期','final_check_time'),('产值','output_value')],
        'm8': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划要求交付时间','plan_delivery_time'),('入库时间','storage_time'),('产值','output_value')],
        'm9': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划要求交付时间','plan_delivery_time'),('交总检日期','final_check_time'),('产值','output_value')],
        'm10': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划要求交付时间','plan_delivery_time'),('入库时间','storage_time'),('产值','output_value')],
        'm11': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('送定检日期','fixed_check_deliver_time'),('入库时间','storage_time'),('耗用时间','__cost'),('产值','output_value')],
        'm12': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('送定检日期','fixed_check_deliver_time'),('入库时间','storage_time'),('耗用时间','__cost'),('产值','output_value')],
        'm13': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('预计提交总检时间','estimated_finish_time'),('产值','output_value')],
        'm14': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划要求交付时间','plan_delivery_time'),('入库日期','storage_time'),('产值','output_value')],
        'm15': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划要求交付时间','plan_delivery_time'),('产值','output_value')],
        'm16': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划要求交付时间','plan_delivery_time'),('完成情况','__completion_status'),('产值','output_value')],
        'm17': [('序号','serial_no'),('生产状态','production_status'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('违约金','liquidated_damages'),('主计划要求交付时间','plan_delivery_time'),('交总检日期','final_check_time'),('入库日期','storage_time'),('产值','output_value')],
        'm18': [('序号','serial_no'),('生产状态','production_status'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('违约金','liquidated_damages'),('主计划要求交付时间','plan_delivery_time'),('入库时间','storage_time'),('核心原因','__core_reason'),('延期天数','__delay_days'),('产值','output_value')],
    }

    # m18 extra computation
    em18 = m_map.get('m18', [])
    em18_ids = [t.id for t in em18]
    em18_issues = IssueRecord.query.filter(IssueRecord.task_id.in_(em18_ids)).all() if em18_ids else []
    em18_shortages = ShortageRecord.query.filter(ShortageRecord.task_id.in_(em18_ids)).all() if em18_ids else []
    from collections import defaultdict
    eim = defaultdict(list); esm = defaultdict(list)
    for iss in em18_issues: eim[iss.task_id].append(iss)
    for sr in em18_shortages: esm[sr.task_id].append(sr)
    em18_extra = {}
    for t in em18:
        max_d = 0; core = '无'
        for iss in eim.get(t.id, []):
            if iss.raise_time:
                d = ((iss.finish_time or now) - iss.raise_time).days
                if d > max_d: max_d = d; core = ('技术' if iss.issue_type=='tech' else '管理') + (f'-{iss.dept}' if iss.dept else '')
        for sr in esm.get(t.id, []):
            if sr.report_time:
                d = ((sr.arrive_time or now) - sr.report_time).days
                if d > max_d: max_d = d; core = sr.shortage_type or ''
        em18_extra[t.id] = (core, max_d)

    cols = mc.get(mid, [])
    t_list = m_map.get(mid, [])
    rows_data = []
    for t in t_list:
        row = {}
        for _, f in cols:
            if f == '__cost': row[f] = fc(mid, t.id)
            elif f == '__completion_status': row[f] = '已完成' if t.storage_time else '未完成'
            elif f == '__core_reason': row[f] = em18_extra.get(t.id, ('无',0))[0]
            elif f == '__delay_days': days = em18_extra.get(t.id, ('无',0))[1]; row[f] = f'{days}天' if days > 0 else '无延期'
            elif f == 'production_status': row[f] = compute_prod_status(t)
            elif f == 'product_name': row[f] = str(t.product_name or '')
            else: row[f] = fmt_val(getattr(t, f))
        rows_data.append(row)

    # m18 → Word export; others → Excel
    if mid == 'm18':
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        def set_cell_shading(cell, color):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), color)
            shading.set(qn('w:val'), 'clear')
            tcPr.append(shading)

        doc = Document()
        section = doc.sections[0]
        from docx.enum.section import WD_ORIENT
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        style = doc.styles['Normal']
        style.font.size = Pt(10)
        style.font.name = '宋体'
        rPr = style.element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '宋体')
        rPr.append(rFonts)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.space_before = Pt(2)

        title = doc.add_heading(f'影响因素分析明细 — {selected_year}年{selected_month}月', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        task_cols = [c for c in cols if c[1] not in ('__link',)]
        task_col_names = [cn for cn, _ in task_cols]
        task_col_fields = [f for _, f in task_cols]

        for idx, t in enumerate(t_list):
            row = rows_data[idx]
            if idx > 0:
                doc.add_paragraph('—' * 40)

            # Task header
            doc.add_paragraph(f'【{idx+1}】 {t.product_name or ""}  (序号: {t.serial_no})', style='List Number')
            
            # Task info table
            ncols = len(task_cols)
            table = doc.add_table(rows=2, cols=ncols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            for ci, cn in enumerate(task_col_names):
                cell = table.cell(0, ci)
                cell.text = cn
                set_cell_shading(cell, 'D9D9D9')
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(9)
            for ci, f in enumerate(task_col_fields):
                cell = table.cell(1, ci)
                cell.text = str(row.get(f, ''))
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(9)

            # Detail analysis
            doc.add_paragraph('▼ 详情分析')
            iss_list = eim.get(t.id, [])
            sh_list = esm.get(t.id, [])

            if not iss_list and not sh_list:
                doc.add_paragraph('  无技术/管理报错和物料缺件记录')

            if iss_list:
                doc.add_paragraph('  【技术/管理报错】')
                itbl = doc.add_table(rows=len(iss_list)+1, cols=6)
                itbl.style = 'Table Grid'
                itbl.autofit = True
                iheaders = ['类型','部门','内容','提出时间','完成时间','延期天数']
                for ci, h in enumerate(iheaders):
                    cell = itbl.cell(0, ci)
                    cell.text = h
                    set_cell_shading(cell, 'F2F2F2')
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
                            run.font.size = Pt(8)
                for ri, iss in enumerate(iss_list, 1):
                    d = ((iss.finish_time or now) - iss.raise_time).days if iss.raise_time else 0
                    vals = [
                        '技术' if iss.issue_type=='tech' else '管理',
                        iss.dept or '',
                        iss.content or '',
                        fmt_date(iss.raise_time),
                        fmt_date(iss.finish_time or now),
                        f'{d}天'
                    ]
                    for ci, v in enumerate(vals):
                        cell = itbl.cell(ri, ci)
                        cell.text = v
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(8)

            if sh_list:
                doc.add_paragraph('  【物料缺件报备】')
                stbl = doc.add_table(rows=len(sh_list)+1, cols=7)
                stbl.style = 'Table Grid'
                stbl.autofit = True
                sheaders = ['类型','内容','发送时间','报缺时间','到位时间','延期天数','备注']
                for ci, h in enumerate(sheaders):
                    cell = stbl.cell(0, ci)
                    cell.text = h
                    set_cell_shading(cell, 'F2F2F2')
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
                            run.font.size = Pt(8)
                for ri, sr in enumerate(sh_list, 1):
                    d = ((sr.arrive_time or now) - sr.report_time).days if sr.report_time else 0
                    vals = [
                        sr.shortage_type or '',
                        sr.content or '',
                        fmt_date(sr.send_time),
                        fmt_date(sr.report_time),
                        fmt_date(sr.arrive_time or now),
                        f'{d}天',
                        ''
                    ]
                    for ci, v in enumerate(vals):
                        cell = stbl.cell(ri, ci)
                        cell.text = v
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(8)

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return Response(output.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment;filename=m18_{selected_year}_{selected_month}.docx'})

    # Other metrics → Excel (unchanged)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = '导出数据'
    thin = Side(style='thin')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    hf = Font(bold=True, size=11)
    cf = Font(size=11)
    ca = Alignment(horizontal='center', vertical='center', wrap_text=True)

    for ci, (cn, _) in enumerate(cols, 1):
        c = ws.cell(row=1, column=ci, value=cn)
        c.font = hf; c.border = border; c.alignment = ca

    for ri, row in enumerate(rows_data, 2):
        for ci, (_, f) in enumerate(cols, 1):
            c = ws.cell(row=ri, column=ci, value=row.get(f, ''))
            c.font = cf; c.border = border; c.alignment = ca

    # Auto-fit column widths and row heights
    for ci in range(1, len(cols) + 1):
        max_len = 0
        for ri in range(1, len(rows_data) + 2):
            val = ws.cell(row=ri, column=ci).value
            if val:
                clen = sum(2 if ord(ch) > 127 else 1 for ch in str(val))
                max_len = max(max_len, clen)
        col_letter = openpyxl.utils.get_column_letter(ci)
        ws.column_dimensions[col_letter].width = min(max_len + 4, 55)
    ws.row_dimensions[1].height = 26
    for ri in range(2, len(rows_data) + 2):
        ws.row_dimensions[ri].height = 22

    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return Response(output.getvalue(), mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', headers={'Content-Disposition': f'attachment;filename={mid}_{selected_year}_{selected_month}.xlsx'})

# ---------- 数据中心路由 ----------
@app.route('/datacenter')
@login_required
def datacenter():
    now = bj_now()
    try: selected_year = int(request.args.get('year', now.year))
    except: selected_year = now.year
    try: selected_month = int(request.args.get('month', now.month))
    except: selected_month = now.month
    month_start = datetime(selected_year, selected_month, 1)
    if selected_month == 12: month_end = datetime(selected_year + 1, 1, 1)
    else: month_end = datetime(selected_year, selected_month + 1, 1)
    year_start = datetime(selected_year, 1, 1)
    year_end = datetime(selected_year + 1, 1, 1)
    tasks = ProductionTask.query.all()
    def sum_val(lst): return sum(float(t.output_value or 0) for t in lst)
    m1 = [t for t in tasks if t.storage_time and month_start <= t.storage_time < month_end]
    m2 = [t for t in tasks if t.plan_delivery_time and year_start <= t.plan_delivery_time < year_end and t.storage_time]
    m3 = [t for t in tasks if t.plan_delivery_time and month_start <= t.plan_delivery_time < month_end and not t.storage_time]
    m4 = [t for t in tasks if t.plan_delivery_time and year_start <= t.plan_delivery_time < year_end and not t.storage_time]
    m5 = [t for t in tasks if t.plan_issue_time and month_start <= t.plan_issue_time < month_end]
    denom6 = [t for t in tasks if t.plan_delivery_time and month_start <= t.plan_delivery_time < month_end]
    m6 = [t for t in denom6 if t.storage_time]
    rate6 = round(len(m6)/len(denom6)*100,1) if denom6 else 0
    m7 = [t for t in denom6 if t.final_check_time and t.storage_time and t.plan_delivery_time and t.storage_time <= t.plan_delivery_time]
    rate7 = round(len(m7)/len(denom6)*100,1) if denom6 else 0
    m8 = [t for t in tasks if t.storage_time and month_start <= t.storage_time < month_end and t.plan_delivery_time and t.storage_time <= t.plan_delivery_time]
    rate8 = round(len(m8)/len(denom6)*100,1) if denom6 else 0
    denom9 = [t for t in tasks if t.plan_delivery_time and year_start <= t.plan_delivery_time <= now]
    m9 = [t for t in tasks if t.final_check_time and year_start <= t.final_check_time < year_end]
    rate9 = round(len(m9)/len(denom9)*100,1) if denom9 else 0
    m10 = [t for t in tasks if t.storage_time and year_start <= t.storage_time < year_end]
    rate10 = round(len(m10)/len(denom9)*100,1) if denom9 else 0
    m11 = [t for t in tasks if t.fixed_check_finish_time and year_start <= t.fixed_check_finish_time < year_end]
    m12 = [t for t in tasks if t.fixed_check_deliver_time and not t.fixed_check_finish_time]
    m13 = [t for t in tasks if t.fixed_check_qty and not t.fixed_check_deliver_time and t.estimated_finish_time and (t.estimated_finish_time - now).days < 60]
    m14 = [t for t in tasks if t.plan_delivery_time and month_start <= t.plan_delivery_time < month_end and t.liquidated_damages == '是' and not t.storage_time]
    m15 = [t for t in tasks if t.plan_delivery_time and month_start <= t.plan_delivery_time < month_end and t.liquidated_damages == '是']
    m16 = [t for t in tasks if t.plan_delivery_time and year_start <= t.plan_delivery_time < year_end and t.liquidated_damages == '是']
    m17 = [t for t in tasks if t.plan_delivery_time and month_start <= t.plan_delivery_time < month_end]
    m18 = [t for t in tasks if t.plan_delivery_time and month_start <= t.plan_delivery_time < month_end and (t.storage_time or now) > t.plan_delivery_time]
    m18_ids = [t.id for t in m18]
    m18_issues = IssueRecord.query.filter(IssueRecord.task_id.in_(m18_ids)).all() if m18_ids else []
    m18_shortages = ShortageRecord.query.filter(ShortageRecord.task_id.in_(m18_ids)).all() if m18_ids else []
    m18_issues_map = defaultdict(list)
    for iss in m18_issues: m18_issues_map[iss.task_id].append(iss)
    m18_shortages_map = defaultdict(list)
    for sr in m18_shortages: m18_shortages_map[sr.task_id].append(sr)
    m18_extra = {}
    for t in m18:
        max_d = 0; core = '无'
        for iss in m18_issues_map.get(t.id, []):
            if iss.raise_time:
                d = ((iss.finish_time or now) - iss.raise_time).days
                if d > max_d: max_d = d; core = ('技术' if iss.issue_type=='tech' else '管理') + (f'-{iss.dept}' if iss.dept else '')
        for sr in m18_shortages_map.get(t.id, []):
            if sr.report_time:
                d = ((sr.arrive_time or now) - sr.report_time).days
                if d > max_d: max_d = d; core = sr.shortage_type or ''
        m18_extra[t.id] = (core, max_d)
    cost_time_map = {}
    for t in m11:
        if t.fixed_check_deliver_time and t.fixed_check_finish_time:
            cost_time_map[('m11', t.id)] = f'{(t.fixed_check_finish_time - t.fixed_check_deliver_time).days}天'
    for t in m12:
        if t.fixed_check_deliver_time:
            cost_time_map[('m12', t.id)] = f'{(now - t.fixed_check_deliver_time).days}天'
    def fmt_cost(key, tid):
        return cost_time_map.get((key, tid), '')
    metric_cols = {
        'm1': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划要求交付时间','plan_delivery_time'),('产值','output_value')],
        'm2': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划要求交付时间','plan_delivery_time'),('产值','output_value')],
        'm3': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划要求交付时间','plan_delivery_time'),('产值','output_value')],
        'm4': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划要求交付时间','plan_delivery_time'),('产值','output_value')],
        'm5': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划下达时间','plan_issue_time'),('产值','output_value')],
        'm6': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划要求交付时间','plan_delivery_time'),('交总检日期','final_check_time'),('产值','output_value')],
        'm7': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划要求交付时间','plan_delivery_time'),('交总检日期','final_check_time'),('产值','output_value')],
        'm8': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划要求交付时间','plan_delivery_time'),('入库时间','storage_time'),('产值','output_value')],
        'm9': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划要求交付时间','plan_delivery_time'),('交总检日期','final_check_time'),('产值','output_value')],
        'm10': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划要求交付时间','plan_delivery_time'),('入库时间','storage_time'),('产值','output_value')],
        'm11': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('送定检日期','fixed_check_deliver_time'),('入库时间','storage_time'),('耗用时间','__cost'),('产值','output_value')],
        'm12': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('送定检日期','fixed_check_deliver_time'),('入库时间','storage_time'),('耗用时间','__cost'),('产值','output_value')],
        'm13': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('预计提交总检时间','estimated_finish_time'),('产值','output_value')],
        'm14': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划要求交付时间','plan_delivery_time'),('入库日期','storage_time'),('产值','output_value')],
        'm15': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划要求交付时间','plan_delivery_time'),('产值','output_value')],
        'm16': [('序号','serial_no'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('主计划要求交付时间','plan_delivery_time'),('完成情况','__completion_status'),('产值','output_value')],
        'm17': [('序号','serial_no'),('生产状态','production_status'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('违约金','liquidated_damages'),('主计划要求交付时间','plan_delivery_time'),('交总检日期','final_check_time'),('入库日期','storage_time'),('产值','output_value'),('','__hover_text')],
        'm18': [('序号','serial_no'),('生产状态','production_status'),('品种','product_category'),('图号','product_draw_no'),('产品','product_name'),('批次','batch_no'),('数量','total_qty'),('违约金','liquidated_damages'),('主计划要求交付时间','plan_delivery_time'),('入库时间','storage_time'),('核心原因','__core_reason'),('延期天数','__delay_days'),('产值','output_value'),('操作','__link')],
    }
    def prep_rows(t_list, mid):
        rows = []; cols = metric_cols[mid]
        for t in t_list:
            row = {}
            for _, f in cols:
                if f == '__cost': row[f] = fmt_cost(mid, t.id)
                elif f == 'product_name': row[f] = str(t.product_name or '')
                elif f == '__completion_status': row[f] = '已完成' if t.storage_time else '未完成'
                elif f == 'production_status':
                    if t.final_check_time: row[f] = '已完成'
                    elif any([t.cut_start,t.sew_start,t.glue_start,t.assembly1_start,t.assembly2_start,t.oxygen_start,t.heat_seal_start]): row[f] = '生产中'
                    else: row[f] = '未开工'
                elif f == '__hover_text':
                    procs = [('裁剪',t.cut_start,t.cut_end),('缝纫',t.sew_start,t.sew_end),('粘胶',t.glue_start,t.glue_end),('总装1',t.assembly1_start,t.assembly1_end),('总装2',t.assembly2_start,t.assembly2_end),('氧调',t.oxygen_start,t.oxygen_end),('热风热合',t.heat_seal_start,t.heat_seal_end)]
                    lines = []
                    for pn,ps,pe in procs:
                        if pe: lines.append(f"{pn}: 已完成 ({fmt_date(ps)} 至 {fmt_date(pe)})")
                        elif ps: lines.append(f"{pn}: 生产中 (开始: {fmt_date(ps)})")
                        else: lines.append(f"{pn}: 未开工")
                    row[f] = '\n'.join(lines)
                elif f == '__core_reason':
                    if mid == 'm18': row[f] = m18_extra.get(t.id,('无',0))[0]
                    else: row[f] = ''
                elif f == '__delay_days':
                    if mid == 'm18': d = m18_extra.get(t.id,('无',0))[1]; row[f] = f'{d}天' if d>0 else '无延期'
                    else: row[f] = ''
                elif f == '__link': row[f] = str(t.id)
                else: row[f] = fmt_val(getattr(t, f))
            rows.append(row)
        return rows
    metrics = [
        {'id':'m17','title':'当月生产任务','rows':prep_rows(m17,'m17'),'count':len(m17),'val':sum_val(m17),'rate':None},
        {'id':'m1','title':'当月入库项数','rows':prep_rows(m1,'m1'),'count':len(m1),'val':sum_val(m1),'rate':None},
        {'id':'m2','title':'累积入库项数','rows':prep_rows(m2,'m2'),'count':len(m2),'val':sum_val(m2),'rate':None},
        {'id':'m3','title':'当月未入库项数','rows':prep_rows(m3,'m3'),'count':len(m3),'val':sum_val(m3),'rate':None},
        {'id':'m4','title':'累计未入库项数','rows':prep_rows(m4,'m4'),'count':len(m4),'val':sum_val(m4),'rate':None},
        {'id':'m5','title':'月度下发计划数','rows':prep_rows(m5,'m5'),'count':len(m5),'val':sum_val(m5),'rate':None},
        {'id':'m6','title':'月度完成率','rows':prep_rows(m6,'m6'),'count':len(m6),'val':sum_val(m6),'rate':rate6},
        {'id':'m7','title':'月度准时完成率','rows':prep_rows(m7,'m7'),'count':len(m7),'val':sum_val(m7),'rate':rate7},
        {'id':'m8','title':'月度准时入库率','rows':prep_rows(m8,'m8'),'count':len(m8),'val':sum_val(m8),'rate':rate8},
        {'id':'m9','title':'累积完成率','rows':prep_rows(m9,'m9'),'count':len(m9),'val':sum_val(m9),'rate':rate9},
        {'id':'m10','title':'累积入库率','rows':prep_rows(m10,'m10'),'count':len(m10),'val':sum_val(m10),'rate':rate10},
        {'id':'m11','title':'定检完成任务','rows':prep_rows(m11,'m11'),'count':len(m11),'val':sum_val(m11),'rate':None},
        {'id':'m12','title':'在定检任务','rows':prep_rows(m12,'m12'),'count':len(m12),'val':sum_val(m12),'rate':None},
        {'id':'m13','title':'待定检任务','rows':prep_rows(m13,'m13'),'count':len(m13),'val':sum_val(m13),'rate':None},
        {'id':'m14','title':'当月未完成违约金','rows':prep_rows(m14,'m14'),'count':len(m14),'val':sum_val(m14),'rate':None},
        {'id':'m15','title':'当月违约金任务','rows':prep_rows(m15,'m15'),'count':len(m15),'val':sum_val(m15),'rate':None},
        {'id':'m16','title':'累积违约金任务','rows':prep_rows(m16,'m16'),'count':len(m16),'val':sum_val(m16),'rate':None},
        {'id':'m18','title':'影响因素分析','rows':prep_rows(m18,'m18'),'count':len(m18),'val':sum_val(m18),'rate':None},
    ]
    year_range = list(range(2024, now.year + 6))
    return render_template_string(DATACENTER_HTML, metrics=metrics, selected_year=selected_year, selected_month=selected_month, year_range=year_range, metric_cols=metric_cols)

# ---------- 班组模块 ----------
@app.route('/teamkanban')
@login_required
def team_kanban():
    month_filter = request.args.get('month', '')
    selected_team = request.args.get('team', '')
    db.session.expire_all()
    tasks = ProductionTask.query.all()
    now = bj_now()

    available_months = set()
    for t in tasks:
        if t.plan_delivery_time:
            available_months.add(t.plan_delivery_time.strftime('%Y-%m'))
    available_months = sorted(available_months, reverse=True)

    if month_filter:
        year, month = month_filter.split('-')
        tasks = [t for t in tasks if t.plan_delivery_time and
                 t.plan_delivery_time.year == int(year) and
                 t.plan_delivery_time.month == int(month)]

    teams = ['裁剪','缝纫','粘胶','总装1','总装2','氧调','热风热合']
    team_field_map = {
        '裁剪': ('cut_start', 'cut_end'),
        '缝纫': ('sew_start', 'sew_end'),
        '粘胶': ('glue_start', 'glue_end'),
        '总装1': ('assembly1_start', 'assembly1_end'),
        '总装2': ('assembly2_start', 'assembly2_end'),
        '氧调': ('oxygen_start', 'oxygen_end'),
        '热风热合': ('heat_seal_start', 'heat_seal_end'),
    }

    kanban_data = {}
    for team in teams:
        start_field, end_field = team_field_map[team]
        todo, doing, done, skipped = [], [], [], []
        for t in tasks:
            start_val = getattr(t, start_field)
            end_val = getattr(t, end_field)
            is_overdue = False
            overdue_days = 0
            if not end_val and t.plan_delivery_time:
                is_overdue = t.plan_delivery_time < now
                if is_overdue:
                    overdue_days = (now - t.plan_delivery_time).days

            duration_days = 0
            if start_val and not end_val:
                duration_days = (now - start_val).days

            item = {
                'id': t.id,
                'serial_no': t.serial_no,
                'product_name': t.product_name,
                'product_draw_no': t.product_draw_no,
                'batch_no': t.batch_no,
                'total_qty': t.total_qty,
                'specific_model': t.specific_model,
                'plan_delivery_time': t.plan_delivery_time,
                'start_time': start_val,
                'end_time': end_val,
                'is_overdue': is_overdue,
                'overdue_days': overdue_days,
                'duration_days': duration_days,
                'product_category': t.product_category or '',
                'stage': t.stage or '',
                'output_value': float(t.output_value or 0),
                'operator': t.operator or '',
            }
            if end_val:
                if start_val and start_val == end_val:
                    skipped.append(item)
                else:
                    done.append(item)
            elif start_val:
                doing.append(item)
            else:
                todo.append(item)

        team_overdue_count = sum(1 for item in todo if item['is_overdue']) + sum(1 for item in doing if item['is_overdue'])

        kanban_data[team] = {
            'todo': todo,
            'doing': doing,
            'done': done,
            'skipped': skipped,
            'overdue_count': team_overdue_count,
            'todo_overdue': sum(1 for item in todo if item['is_overdue']),
            'doing_overdue': sum(1 for item in doing if item['is_overdue']),
        }

        member_stats = {}
        for item in todo:
            op = item.get('operator', '')
            if not op: continue
            for single_op in op.split(','):
                single_op = single_op.strip()
                if not single_op: continue
                member_stats.setdefault(single_op, {'todo': 0, 'doing': 0, 'done': 0})
                member_stats[single_op]['todo'] += 1
        for item in doing:
            op = item.get('operator', '')
            if not op: continue
            for single_op in op.split(','):
                single_op = single_op.strip()
                if not single_op: continue
                member_stats.setdefault(single_op, {'todo': 0, 'doing': 0, 'done': 0})
                member_stats[single_op]['doing'] += 1
        for item in done:
            op = item.get('operator', '')
            if not op: continue
            for single_op in op.split(','):
                single_op = single_op.strip()
                if not single_op: continue
                member_stats.setdefault(single_op, {'todo': 0, 'doing': 0, 'done': 0})
                member_stats[single_op]['done'] += 1
        kanban_data[team]['member_stats'] = member_stats

        member_tasks = {}
        for item in doing:
            op = item.get('operator', '')
            if not op: continue
            for single_op in op.split(','):
                single_op = single_op.strip()
                if not single_op: continue
                member_tasks.setdefault(single_op, []).append({
                'id': item['id'],
                'serial_no': item['serial_no'],
                'product_draw_no': item.get('product_draw_no', ''),
                'product_name': item.get('product_name', ''),
                'batch_no': item.get('batch_no', ''),
                'total_qty': item.get('total_qty', ''),
            })
        kanban_data[team]['member_tasks'] = member_tasks

    for v in kanban_data.values():
        v['todo'].reverse()
        v['doing'].reverse()
        v['done'].reverse()
        v['skipped'].reverse()

    total_todo = sum(len(v['todo']) for v in kanban_data.values())
    total_doing = sum(len(v['doing']) for v in kanban_data.values())
    total_done = sum(len(v['done']) for v in kanban_data.values())
    total_overdue = sum(v['overdue_count'] for v in kanban_data.values())
    global_stats = {'todo': total_todo, 'doing': total_doing, 'done': total_done, 'overdue': total_overdue}

    members_data = {}
    for m in TeamMember.query.all():
        members_data.setdefault(m.team, []).append({'id': m.id, 'name': m.name})

    branch_notices = BranchNotice.query.filter(BranchNotice.target.in_(['班组','全部'])).order_by(BranchNotice.created_time.desc()).limit(5).all()

    resp = make_response(render_template_string(TEAM_KANBAN_HTML,
        kanban_data=kanban_data, teams=teams, fmt_date=fmt_date,
        now=now, global_stats=global_stats, members_data=members_data,
        month_filter=month_filter, available_months=available_months,
        selected_team=selected_team, branch_notices=branch_notices))
    resp.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate, max-age=0'
    resp.headers['Pragma'] = 'no-cache'
    resp.headers['Expires'] = '0'
    return resp

@app.route('/teamkanban/notice/add', methods=['POST'])
@login_required
@role_required('管理员')
def add_branch_notice():
    title = request.form.get('title','').strip()
    content = request.form.get('content','').strip()
    if not title or not content:
        flash('请填写通知标题和内容','danger')
        return redirect(url_for('team_kanban', month=request.form.get('month',''), team=request.form.get('team','')))
    db.session.add(BranchNotice(title=title[:100], content=content, created_by=current_user.id))
    db.session.commit()
    flash('通知已发布','success')
    return redirect(url_for('team_kanban', month=request.form.get('month',''), team=request.form.get('team','')))

@app.route('/teamkanban/notice/delete/<int:notice_id>', methods=['POST'])
@login_required
@role_required('管理员')
def delete_branch_notice(notice_id):
    notice = db.session.get(BranchNotice, notice_id)
    if not notice:
        flash('通知不存在','danger')
        return redirect(url_for('team_kanban', month=request.form.get('month',''), team=request.form.get('team','')))
    db.session.delete(notice)
    db.session.commit()
    flash('通知已删除','success')
    return redirect(url_for('team_kanban', month=request.form.get('month',''), team=request.form.get('team','')))

@app.route('/teamkanban/done/<team>')
@login_required
def team_kanban_done(team):
    month_filter = request.args.get('month', '')
    tasks = ProductionTask.query.all()
    if month_filter:
        year, month = month_filter.split('-')
        tasks = [t for t in tasks if t.plan_delivery_time and
                 t.plan_delivery_time.year == int(year) and
                 t.plan_delivery_time.month == int(month)]
    team_field_map = {
        '裁剪': ('cut_start', 'cut_end'), '缝纫': ('sew_start', 'sew_end'),
        '粘胶': ('glue_start', 'glue_end'), '总装1': ('assembly1_start', 'assembly1_end'),
        '总装2': ('assembly2_start', 'assembly2_end'), '氧调': ('oxygen_start', 'oxygen_end'),
        '热风热合': ('heat_seal_start', 'heat_seal_end'),
    }
    pair = team_field_map.get(team)
    if not pair:
        return '无效班组', 404
    start_f, end_f = pair
    done_list = []
    for t in tasks:
        start_val = getattr(t, start_f)
        end_val = getattr(t, end_f)
        if end_val:
            done_list.append({
                'id': t.id, 'serial_no': t.serial_no,
                'product_name': t.product_name or '',
                'product_draw_no': t.product_draw_no or '',
                'batch_no': t.batch_no or '',
                'total_qty': t.total_qty,
                'specific_model': t.specific_model or '',
                'start_time': fmt_date(start_val),
                'end_time': fmt_date(end_val),
                'operator': t.operator or '',
            })
    return render_template_string(DONE_LIST_HTML, team=team, items=done_list,
        fmt_date=fmt_date, month_filter=month_filter)

@app.route('/teamkanban/done/<team>/export')
@login_required
def team_kanban_done_export(team):
    month_filter = request.args.get('month', '')
    tasks = ProductionTask.query.all()
    if month_filter:
        year, month = month_filter.split('-')
        tasks = [t for t in tasks if t.plan_delivery_time and
                 t.plan_delivery_time.year == int(year) and
                 t.plan_delivery_time.month == int(month)]
    team_field_map = {
        '裁剪': ('cut_start', 'cut_end'), '缝纫': ('sew_start', 'sew_end'),
        '粘胶': ('glue_start', 'glue_end'), '总装1': ('assembly1_start', 'assembly1_end'),
        '总装2': ('assembly2_start', 'assembly2_end'), '氧调': ('oxygen_start', 'oxygen_end'),
        '热风热合': ('heat_seal_start', 'heat_seal_end'),
    }
    pair = team_field_map.get(team)
    if not pair:
        return '无效班组', 404
    start_f, end_f = pair

    import openpyxl, io
    from urllib.parse import quote
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f'{team}已完成'
    ws.append(['序号','图号','产品名称','型号','批次','数量','操作人员','开始','完成'])
    for t in tasks:
        end_val = getattr(t, end_f)
        if not end_val:
            continue
        start_val = getattr(t, start_f)
        ws.append([
            t.serial_no,
            t.product_draw_no or '',
            t.product_name or '',
            t.specific_model or '',
            t.batch_no or '',
            str(t.total_qty) if t.total_qty else '',
            t.operator or '',
            fmt_date(start_val),
            fmt_date(end_val),
        ])
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    filename = quote(f'{team}_已完成任务.xlsx')
    resp = Response(output.getvalue(), mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    resp.headers['Content-Disposition'] = f"attachment; filename*=UTF-8''{filename}"
    return resp

# ---------- 看板操作API ----------
@app.route('/api/kanban/start/<int:task_id>', methods=['POST'])
@login_required
def api_kanban_start(task_id):
    team = request.form.get('team')
    if not team:
        return jsonify({'ok': False, 'msg': '未指定班组'}), 400
    team_field_map = {
        '裁剪': 'cut_start', '缝纫': 'sew_start', '粘胶': 'glue_start',
        '总装1': 'assembly1_start', '总装2': 'assembly2_start',
        '氧调': 'oxygen_start', '热风热合': 'heat_seal_start',
    }
    field = team_field_map.get(team)
    if not field:
        return jsonify({'ok': False, 'msg': '无效班组'}), 400
    task = db.session.get(ProductionTask, task_id)
    if not task:
        return jsonify({'ok': False, 'msg': '任务不存在'}), 404
    if getattr(task, field):
        return jsonify({'ok': False, 'msg': '已开始，请勿重复操作'}), 400
    old_val = getattr(task, field)
    now_val = bj_now()
    setattr(task, field, now_val)
    db.session.add(OperationLog(task_id=task.id, operated_by=current_user.id, operated_time=now_val, field_name=field, old_value=str(old_val or ''), new_value=str(now_val), operation_type='班组操作-开始'))
    db.session.commit()
    db.session.close()
    return jsonify({'ok': True, 'msg': '已开始进行'})

@app.route('/api/kanban/finish/<int:task_id>', methods=['POST'])
@login_required
def api_kanban_finish(task_id):
    team = request.form.get('team')
    if not team:
        return jsonify({'ok': False, 'msg': '未指定班组'}), 400
    team_field_map = {
        '裁剪': 'cut_end', '缝纫': 'sew_end', '粘胶': 'glue_end',
        '总装1': 'assembly1_end', '总装2': 'assembly2_end',
        '氧调': 'oxygen_end', '热风热合': 'heat_seal_end',
    }
    end_f = team_field_map.get(team)
    if not end_f:
        return jsonify({'ok': False, 'msg': '无效班组'}), 400
    task = db.session.get(ProductionTask, task_id)
    if not task:
        return jsonify({'ok': False, 'msg': '任务不存在'}), 404
    start_f = end_f.replace('_end', '_start')
    if not getattr(task, start_f):
        return jsonify({'ok': False, 'msg': '请先记录开始时间'}), 400
    if getattr(task, end_f):
        return jsonify({'ok': False, 'msg': '已完成，请勿重复操作'}), 400
    old_val = getattr(task, end_f)
    now_val = bj_now()
    setattr(task, end_f, now_val)
    db.session.add(OperationLog(task_id=task.id, operated_by=current_user.id, operated_time=now_val, field_name=end_f, old_value=str(old_val or ''), new_value=str(now_val), operation_type='班组操作-完成'))
    db.session.commit()
    db.session.close()
    return jsonify({'ok': True, 'msg': '已完成'})

@app.route('/api/kanban/cancel/<int:task_id>', methods=['POST'])
@login_required
def api_kanban_cancel(task_id):
    team = request.form.get('team')
    if not team:
        return jsonify({'ok': False, 'msg': '未指定班组'}), 400
    team_field_map = {
        '裁剪': 'cut_start', '缝纫': 'sew_start', '粘胶': 'glue_start',
        '总装1': 'assembly1_start', '总装2': 'assembly2_start',
        '氧调': 'oxygen_start', '热风热合': 'heat_seal_start',
    }
    start_f = team_field_map.get(team)
    if not start_f:
        return jsonify({'ok': False, 'msg': '无效班组'}), 400
    end_f = start_f.replace('_start', '_end')
    task = db.session.get(ProductionTask, task_id)
    if not task:
        return jsonify({'ok': False, 'msg': '任务不存在'}), 404
    old_start = getattr(task, start_f)
    old_end = getattr(task, end_f)
    now_val = bj_now()
    setattr(task, start_f, None)
    setattr(task, end_f, None)
    db.session.add(OperationLog(task_id=task.id, operated_by=current_user.id, operated_time=now_val, field_name=start_f, old_value=str(old_start or ''), new_value='(已清空)', operation_type='班组操作-取消'))
    db.session.add(OperationLog(task_id=task.id, operated_by=current_user.id, operated_time=now_val, field_name=end_f, old_value=str(old_end or ''), new_value='(已清空)', operation_type='班组操作-取消'))
    db.session.commit()
    db.session.close()
    return jsonify({'ok': True, 'msg': '已取消开始'})

@app.route('/api/kanban/skip/<int:task_id>', methods=['POST'])
@login_required
def api_kanban_skip(task_id):
    team = request.form.get('team')
    if not team:
        return jsonify({'ok': False, 'msg': '未指定班组'}), 400
    team_field_map = {
        '裁剪': ('cut_start', 'cut_end'), '缝纫': ('sew_start', 'sew_end'),
        '粘胶': ('glue_start', 'glue_end'), '总装1': ('assembly1_start', 'assembly1_end'),
        '总装2': ('assembly2_start', 'assembly2_end'), '氧调': ('oxygen_start', 'oxygen_end'),
        '热风热合': ('heat_seal_start', 'heat_seal_end'),
    }
    pair = team_field_map.get(team)
    if not pair:
        return jsonify({'ok': False, 'msg': '无效班组'}), 400
    start_f, end_f = pair
    task = db.session.get(ProductionTask, task_id)
    if not task:
        return jsonify({'ok': False, 'msg': '任务不存在'}), 404
    if getattr(task, start_f) and getattr(task, end_f):
        return jsonify({'ok': False, 'msg': '已处理，请勿重复操作'}), 400
    now = bj_now()
    if not getattr(task, start_f):
        setattr(task, start_f, now)
        db.session.add(OperationLog(task_id=task.id, operated_by=current_user.id, operated_time=now, field_name=start_f, old_value='', new_value=str(now), operation_type='班组操作-跳过'))
    if not getattr(task, end_f):
        setattr(task, end_f, now)
        db.session.add(OperationLog(task_id=task.id, operated_by=current_user.id, operated_time=now, field_name=end_f, old_value='', new_value=str(now), operation_type='班组操作-跳过'))
    db.session.commit()
    db.session.close()
    return jsonify({'ok': True, 'msg': '已标记不涉及'})

@app.route('/api/kanban/operator/<int:task_id>', methods=['POST'])
@login_required
def api_kanban_operator(task_id):
    operator = request.form.get('operator', '').strip()
    task = db.session.get(ProductionTask, task_id)
    if not task:
        return jsonify({'ok': False, 'msg': '任务不存在'}), 404
    old_op = task.operator or ''
    task.operator = operator if operator else None
    db.session.add(OperationLog(task_id=task.id, operated_by=current_user.id, operated_time=bj_now(), field_name='operator', old_value=old_op, new_value=operator or '', operation_type='班组操作-指派人员'))
    db.session.commit()
    db.session.close()
    return jsonify({'ok': True, 'msg': '操作人员已更新'})

@app.route('/api/kanban/members/<team>')
@login_required
def api_kanban_members(team):
    members = TeamMember.query.filter_by(team=team).all()
    return jsonify({'ok': True, 'members': [{'id': m.id, 'name': m.name} for m in members]})

@app.route('/api/kanban/members/add', methods=['POST'])
@login_required
def api_kanban_members_add():
    team = request.form.get('team', '')
    name = request.form.get('name', '').strip()
    if not team or not name:
        return jsonify({'ok': False, 'msg': '班组和姓名不能为空'}), 400
    existing = TeamMember.query.filter_by(team=team, name=name).first()
    if existing:
        return jsonify({'ok': False, 'msg': '该成员已存在'}), 400
    member = TeamMember(team=team, name=name)
    db.session.add(member)
    db.session.commit()
    return jsonify({'ok': True, 'msg': '成员已添加', 'id': member.id})

@app.route('/api/kanban/members/remove/<int:member_id>', methods=['POST'])
@login_required
def api_kanban_members_remove(member_id):
    member = db.session.get(TeamMember, member_id)
    if not member:
        return jsonify({'ok': False, 'msg': '成员不存在'}), 404
    db.session.delete(member)
    db.session.commit()
    return jsonify({'ok': True, 'msg': '成员已移除'})

@app.route('/api/kanban/members/edit/<int:member_id>', methods=['POST'])
@login_required
def api_kanban_members_edit(member_id):
    member = db.session.get(TeamMember, member_id)
    if not member:
        return jsonify({'ok': False, 'msg': '成员不存在'}), 404
    name = request.form.get('name', '').strip()
    if not name:
        return jsonify({'ok': False, 'msg': '姓名不能为空'}), 400
    member.name = name
    db.session.commit()
    return jsonify({'ok': True, 'msg': '成员已更新'})

# ---------- 问题/缺件添加API ----------
@app.route('/api/issue/add', methods=['POST'])
@login_required
@role_required('计调员')
def api_issue_add():
    data = request.get_json()
    task_id = data.get('task_id')
    if not task_id:
        return jsonify({'error': '请选择任务'}), 400
    task = db.session.get(ProductionTask, int(task_id))
    if not task:
        return jsonify({'error': '任务不存在'}), 404
    now = bj_now()
    iss = IssueRecord(
        task_id=task.id,
        issue_type=data.get('issue_type','')[:10],
        content=data.get('content','')[:500],
        dept=data.get('dept','')[:50],
        raise_time=parse_date_slash(data.get('raise_time')) or now,
    )
    db.session.add(iss)
    db.session.add(OperationLog(task_id=task.id, operated_by=current_user.id, operated_time=now, field_name='issue', operation_type='添加问题', new_value=iss.content[:50] if iss.content else ''))
    db.session.commit()
    return jsonify({'success': True, 'id': iss.id})

@app.route('/api/shortage/add', methods=['POST'])
@login_required
@role_required('计调员')
def api_shortage_add():
    data = request.get_json()
    task_id = data.get('task_id')
    if not task_id:
        return jsonify({'error': '请选择任务'}), 400
    task = db.session.get(ProductionTask, int(task_id))
    if not task:
        return jsonify({'error': '任务不存在'}), 404
    now = bj_now()
    sr = ShortageRecord(
        task_id=task.id,
        shortage_type=data.get('shortage_type','')[:50],
        content=data.get('content','')[:500],
        report_time=parse_date_slash(data.get('report_time')) or now,
    )
    db.session.add(sr)
    db.session.add(OperationLog(task_id=task.id, operated_by=current_user.id, operated_time=now, field_name='shortage', operation_type='添加缺件', new_value=sr.content[:50] if sr.content else ''))
    db.session.commit()
    return jsonify({'success': True, 'id': sr.id})

@app.route('/planning-task-report')
@login_required
@role_required('计调员')
def planning_task_report():
    db.session.expire_all()
    tasks = ProductionTask.query.order_by(ProductionTask.serial_no).all()
    task_list = []
    for t in tasks:
        task_list.append({
            'id': t.id, 'serial_no': t.serial_no, 'product_name': t.product_name,
            'product_draw_no': t.product_draw_no or '', 'batch_no': t.batch_no or '',
            'product_category': t.product_category or '',
            'matching_time': t.matching_time,
            'fixed_check_deliver_time': t.fixed_check_deliver_time,
            'fixed_check_finish_time': t.fixed_check_finish_time,
            'military_check_plan_time': t.military_check_plan_time,
            'military_check_time': t.military_check_time,
            'fixed_submit_finish_time': t.fixed_submit_finish_time,
            'storage_time': t.storage_time,
        })
    resp = make_response(render_template_string(TASK_REPORT_HTML, tasks=task_list, fmt_date=fmt_date))
    resp.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate, max-age=0'
    resp.headers['Pragma'] = 'no-cache'
    resp.headers['Expires'] = '0'
    return resp

# ---------- 管理员页面 ----------
@app.route('/admin/users')
@login_required
@role_required('管理员')
def admin_users():
    users = User.query.all()
    roles = Role.query.all()
    return render_template_string(ADMIN_USERS_HTML, users=users, roles=roles)

@app.route('/admin/user/edit/<int:user_id>', methods=['POST'])
@login_required
@role_required('管理员')
def edit_user(user_id):
    user = db.session.get(User, user_id)
    if not user:
        flash('用户不存在','danger')
        return redirect(url_for('admin_users'))
    new_role = request.form.get('role_id')
    new_account = request.form.get('account','').strip()
    new_username = request.form.get('username','').strip()
    new_password = request.form.get('password','').strip()
    new_team = request.form.get('team','').strip()
    new_production_line = request.form.get('production_line','').strip()
    if new_role:
        user.role_id = int(new_role)
    if new_account:
        existing = User.query.filter_by(account=new_account).first()
        if existing and existing.id != user_id:
            flash('账号已存在','danger')
            return redirect(url_for('admin_users'))
        user.account = new_account
    if new_username:
        user.username = new_username
    if new_password:
        user.password_hash = generate_password_hash(new_password)
    user.team = new_team if new_team else None
    user.production_line = new_production_line if new_production_line else None
    db.session.commit()
    flash(f'用户 {user.account} 已更新','success')
    return redirect(url_for('admin_users'))

@app.route('/admin/user/delete/<int:user_id>', methods=['POST'])
@login_required
@role_required('管理员')
def delete_user(user_id):
    user = db.session.get(User, user_id)
    if not user:
        flash('用户不存在','danger')
        return redirect(url_for('admin_users'))
    if user.account == 'admin':
        flash('不能删除admin用户','danger')
        return redirect(url_for('admin_users'))
    acct = user.account
    ColumnPermission.query.filter_by(user_id=user_id).delete()
    db.session.delete(user)
    db.session.commit()
    flash(f'用户 {acct} 已删除','success')
    return redirect(url_for('admin_users'))

@app.route('/admin/user/add', methods=['POST'])
@login_required
@role_required('管理员')
def add_user():
    account = request.form['account']
    username = request.form.get('username','')
    password = request.form['password']
    role_id = request.form['role_id']
    team = request.form.get('team','').strip()
    production_line = request.form.get('production_line','').strip()
    if User.query.filter_by(account=account).first():
        flash('账号已存在','danger')
    else:
        db.session.add(User(account=account, username=username, password_hash=generate_password_hash(password), role_id=role_id,
            team=team if team else None, production_line=production_line if production_line else None))
        db.session.commit()
        flash('添加成功','success')
    return redirect(url_for('admin_users'))

@app.route('/admin/permissions')
@login_required
@role_required('管理员')
def admin_permissions():
    users = User.query.all()
    groups = {
        '基本信息': [(f,n) for f,n in COLUMN_ORDER if f in ['serial_no','responsible_person','product_category','product_model','product_draw_no','product_name','total_qty','plan_deliver_qty','fixed_check_qty','batch_no','liquidated_damages','help_letter','first_article','demand_no','check_party','plan_attribute','contract_no','plan_issue_time','plan_delivery_time','specific_model','plan_source','contract_no2','demander','customer_name','project_no','unit_price','output_value','production_time']],
        '工序日期': [(f,n) for f,n in COLUMN_ORDER if f in ['matching_time','cut_start','cut_end','sew_start','sew_end','glue_start','glue_end','assembly1_start','assembly1_end','assembly2_start','assembly2_end','oxygen_start','oxygen_end','heat_seal_start','heat_seal_end']],
        '后续状态': [(f,n) for f,n in COLUMN_ORDER if f in ['estimated_finish_time','production_status','final_check_time','fixed_check_deliver_time','fixed_check_finish_time','military_check_plan_time','military_check_time','fixed_submit_item','fixed_submit_finish_time','storage_time','remark']],
    }
    return render_template_string(ADMIN_PERM_HTML, users=users, groups=groups)

@app.route('/admin/permissions/set', methods=['POST'])
@login_required
@role_required('管理员')
def set_permissions():
    user_id = request.form['user_id']
    ColumnPermission.query.filter_by(user_id=user_id).delete()
    for col, _ in COLUMN_ORDER:
        if f'view_{col}' in request.form:
            db.session.add(ColumnPermission(user_id=user_id, column_name=col, can_view=True, can_edit=(f'edit_{col}' in request.form)))
    db.session.commit()
    flash('权限已更新','success')
    return redirect(url_for('admin_permissions'))

@app.route('/admin/logs')
@login_required
@role_required('管理员')
def admin_logs():
    return render_template_string(ADMIN_LOGS_HTML)

@app.route('/api/logs')
@login_required
@role_required('管理员')
def api_logs():
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 30, type=int)
    op_type = request.args.get('operation_type', '').strip()
    start_date = request.args.get('start_date', '').strip()
    end_date = request.args.get('end_date', '').strip()

    query = db.session.query(
        OperationLog, User.username, User.account
    ).outerjoin(User, OperationLog.operated_by == User.id)

    if op_type:
        query = query.filter(OperationLog.operation_type == op_type)
    if start_date:
        try:
            sd = parse_date_slash(start_date)
            if sd:
                query = query.filter(OperationLog.operated_time >= sd)
        except:
            pass
    if end_date:
        try:
            ed = parse_date_slash(end_date) + timedelta(days=1)
            if ed:
                query = query.filter(OperationLog.operated_time < ed)
        except:
            pass

    total = query.count()
    results = query.order_by(OperationLog.operated_time.desc()).offset(
        (page - 1) * per_page
    ).limit(per_page).all()

    logs = []
    for log, username, account in results:
        op_type = log.operation_type or ('编辑字段' if log.field_name and log.field_name != 'serial_no' else '')
        logs.append({
            'id': log.id,
            'operated_time': log.operated_time.strftime('%Y-%m-%d %H:%M:%S') if log.operated_time else '',
            'operation_type': op_type,
            'task_id': log.task_id,
            'field_name': log.field_name or '',
            'old_value': log.old_value or '',
            'new_value': log.new_value or '',
            'operator': username or account or f'用户{log.operated_by}'
        })

    return jsonify({
        'total': total,
        'page': page,
        'per_page': per_page,
        'total_pages': (total + per_page - 1) // per_page,
        'logs': logs
    })

# ---------- 计划导入 ----------
@app.route('/planning-report')
@login_required
@role_required('计调员')
def planning_report():
    recent = ProductionTask.query.filter_by(created_by=current_user.id).order_by(ProductionTask.created_time.desc()).limit(10).all()
    resp = make_response(render_template_string(PLANNING_REPORT_HTML, recent=recent))
    resp.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate, max-age=0'
    resp.headers['Pragma'] = 'no-cache'
    resp.headers['Expires'] = '0'
    return resp

@app.route('/api/planning-report/submit', methods=['POST'])
@login_required
@role_required('计调员')
def planning_report_submit():
    data = request.get_json()
    with _sn_lock:
        max_sn = db.session.query(db.func.max(cast(ProductionTask.serial_no, Integer))).scalar()
        serial = '1' if max_sn is None else str(max_sn + 1)
        task = ProductionTask(serial_no=serial, created_by=current_user.id, created_time=bj_now(), updated_by=current_user.id, updated_time=bj_now())
        for field, value in data.items():
            if not hasattr(ProductionTask, field):
                continue
            col_type = getattr(ProductionTask, field).type
            try:
                if isinstance(col_type, db.DateTime):
                    setattr(task, field, parse_date_slash(value) if value else None)
                elif isinstance(col_type, db.Integer):
                    setattr(task, field, int(value) if value else None)
                elif isinstance(col_type, db.Numeric):
                    setattr(task, field, float(value) if value else None)
                elif isinstance(col_type, db.Boolean):
                    setattr(task, field, value in ['1','是','true'])
                else:
                    setattr(task, field, str(value) if value else None)
            except:
                pass
        task.first_article = compute_first_article(task)
        task.stage = task.stage or derive_stage_from_batch(task.batch_no)
        db.session.add(task)
        db.session.flush()
        db.session.add(OperationLog(
            task_id=task.id, operated_by=current_user.id, operated_time=bj_now(),
            operation_type='创建任务', field_name='serial_no', new_value=f'序号{task.serial_no}'
        ))
        db.session.commit()
    return jsonify({'success': True, 'task_id': task.id, 'serial_no': task.serial_no})

# ---------- 计调模块 ----------
@app.route('/planning-dashboard')
@login_required
@role_required('计调员')
def planning_dashboard():
    now = bj_now()
    db.session.expire_all()
    tasks = ProductionTask.query.all()

    # Basic stats
    total = len(tasks)
    not_started = 0
    in_progress = 0
    completed = 0
    overdue = 0

    teams = ['裁剪','缝纫','粘胶','总装1','总装2','氧调','热风热合']
    team_field_map = {
        '裁剪': ('cut_start','cut_end'), '缝纫': ('sew_start','sew_end'),
        '粘胶': ('glue_start','glue_end'), '总装1': ('assembly1_start','assembly1_end'),
        '总装2': ('assembly2_start','assembly2_end'), '氧调': ('oxygen_start','oxygen_end'),
        '热风热合': ('heat_seal_start','heat_seal_end'),
    }
    team_stats = {t: {'todo':0,'doing':0,'done':0,'overdue':0} for t in teams}

    overdue_detail = []

    for t in tasks:
        has_start = any([t.cut_start, t.sew_start, t.glue_start, t.assembly1_start, t.assembly2_start, t.oxygen_start, t.heat_seal_start])
        has_end = all([
            t.cut_end is not None if t.cut_start else True,
            t.sew_end is not None if t.sew_start else True,
            t.glue_end is not None if t.glue_start else True,
            t.assembly1_end is not None if t.assembly1_start else True,
            t.assembly2_end is not None if t.assembly2_start else True,
            t.oxygen_end is not None if t.oxygen_start else True,
            t.heat_seal_end is not None if t.heat_seal_start else True,
        ])
        if t.storage_time: completed += 1
        elif has_end: completed += 1
        elif has_start: in_progress += 1
        else: not_started += 1

        deadline_passed = t.plan_delivery_time and t.plan_delivery_time < now and not t.storage_time
        if deadline_passed:
            overdue += 1
            overdue_detail.append({
                'id': t.id, 'serial_no': t.serial_no, 'product_name': t.product_name,
                'product_draw_no': t.product_draw_no or '', 'batch_no': t.batch_no or '',
                'product_category': t.product_category or '', 'responsible_person': t.responsible_person or '',
                'plan_delivery_time': t.plan_delivery_time,
                'overdue_days': (now - t.plan_delivery_time).days,
                'production_status': '未开工' if not has_start else '生产中',
            })

        # Team stats
        for team in teams:
            sf, ef = team_field_map[team]
            sv = getattr(t, sf); ev = getattr(t, ef)
            if ev: team_stats[team]['done'] += 1
            elif sv: team_stats[team]['doing'] += 1
            else: team_stats[team]['todo'] += 1
            if deadline_passed and not ev:
                team_stats[team]['overdue'] += 1

    # Sort overdue by days descending
    overdue_detail.sort(key=lambda x: x['overdue_days'], reverse=True)

    branch_notices = BranchNotice.query.filter(BranchNotice.target.in_(['计调','全部'])).order_by(BranchNotice.created_time.desc()).limit(5).all()

    all_issues = IssueRecord.query.order_by(IssueRecord.raise_time.desc()).limit(30).all()
    issue_ledger = []
    for iss in all_issues:
        task = db.session.get(ProductionTask, iss.task_id)
        if not task: continue
        issue_ledger.append({
            'task_id': iss.task_id, 'serial_no': task.serial_no,
            'product_draw_no': task.product_draw_no or '', 'batch_no': task.batch_no or '',
            'issue_type': '技术' if iss.issue_type == 'tech' else '管理',
            'dept': iss.dept or '', 'content': (iss.content or '')[:50],
            'raise_time': iss.raise_time, 'finish_time': iss.finish_time,
            'days': ((iss.finish_time or now) - iss.raise_time).days if iss.raise_time else 0,
        })
    all_shortages = ShortageRecord.query.order_by(ShortageRecord.report_time.desc()).limit(30).all()
    shortage_ledger = []
    for sr in all_shortages:
        task = db.session.get(ProductionTask, sr.task_id)
        if not task: continue
        shortage_ledger.append({
            'task_id': sr.task_id, 'serial_no': task.serial_no,
            'product_draw_no': task.product_draw_no or '', 'batch_no': task.batch_no or '',
            'shortage_type': sr.shortage_type or '', 'content': (sr.content or '')[:50],
            'report_time': sr.report_time, 'arrive_time': sr.arrive_time,
            'days': ((sr.arrive_time or now) - sr.report_time).days if sr.report_time else 0,
        })

    team_ops = ['班组操作-开始','班组操作-完成','班组操作-跳过','班组操作-指派人员']
    team_logs = OperationLog.query.filter(OperationLog.operation_type.in_(team_ops)).order_by(OperationLog.operated_time.desc()).limit(20).all()
    field_stage_map = {
        'cut_start':'裁剪','cut_end':'裁剪','sew_start':'缝纫','sew_end':'缝纫',
        'glue_start':'粘胶','glue_end':'粘胶','assembly1_start':'总装1','assembly1_end':'总装1',
        'assembly2_start':'总装2','assembly2_end':'总装2','oxygen_start':'氧调','oxygen_end':'氧调',
        'heat_seal_start':'热风热合','heat_seal_end':'热风热合',
    }
    team_dynamics = []
    for log in team_logs:
        task = db.session.get(ProductionTask, log.task_id)
        if not task: continue
        user = db.session.get(User, log.operated_by)
        op_user = (user.username or user.account) if user else '系统'
        stage = field_stage_map.get(log.field_name, '')
        desc = ''
        draw_info = (task.product_draw_no or '') + ('/' + task.batch_no if task.batch_no else '')
        if log.operation_type == '班组操作-开始':
            desc = '{} 开始 {} {}'.format(op_user, stage, draw_info)
        elif log.operation_type == '班组操作-完成':
            desc = '{} 完成 {} {}'.format(op_user, stage, draw_info)
        elif log.operation_type == '班组操作-跳过':
            desc = '{} 跳过 {} {}'.format(op_user, stage, draw_info)
        elif log.operation_type == '班组操作-指派人员':
            new_ops = (log.new_value or '').replace(',','、')
            if new_ops:
                desc = '{} 指派 {} 到 {}'.format(op_user, new_ops, draw_info)
            else:
                desc = '{} 清空 #{} 指派'.format(op_user, task.serial_no)
        team_dynamics.append({
            'task_id': log.task_id, 'serial_no': task.serial_no,
            'product_name': task.product_name,
            'description': desc, 'operation_type': log.operation_type,
            'operated_time': log.operated_time, 'operator': op_user,
        })

    resp = make_response(render_template_string(PLANNING_DASHBOARD_HTML,
        total=total, not_started=not_started, in_progress=in_progress,
        completed=completed, overdue=overdue, team_stats=team_stats,
        teams=teams, overdue_detail=overdue_detail,
        branch_notices=branch_notices, team_dynamics=team_dynamics,
        issue_ledger=issue_ledger, shortage_ledger=shortage_ledger,
        all_tasks=tasks, dept_options=DEPT_OPTIONS, shortage_types=SHORTAGE_TYPES,
        now=now, fmt_date=fmt_date))
    resp.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate, max-age=0'
    resp.headers['Pragma'] = 'no-cache'
    resp.headers['Expires'] = '0'
    return resp

# ---------- 通知公告（领导模块） ----------
@app.route('/admin/notices')
@login_required
@role_required('部门领导')
def admin_notices():
    notices = BranchNotice.query.order_by(BranchNotice.created_time.desc()).all()
    resp = make_response(render_template_string(NOTICES_HTML, notices=notices, fmt_date=fmt_date))
    resp.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate, max-age=0'
    resp.headers['Pragma'] = 'no-cache'
    resp.headers['Expires'] = '0'
    return resp

@app.route('/admin/notices/add', methods=['POST'])
@login_required
@role_required('部门领导')
def admin_notices_add():
    title = (request.form.get('title') or '').strip()
    content = (request.form.get('content') or '').strip()
    if not title or not content:
        flash('请填写标题和内容', 'danger')
        return redirect(url_for('admin_notices'))
    db.session.add(BranchNotice(title=title[:100], content=content, target=(request.form.get('target') or '全部').strip(), created_by=current_user.id))
    db.session.commit()
    flash('通知已发布', 'success')
    return redirect(url_for('admin_notices'))

@app.route('/admin/notices/delete/<int:notice_id>', methods=['POST'])
@login_required
@role_required('部门领导')
def admin_notices_delete(notice_id):
    notice = db.session.get(BranchNotice, notice_id)
    if notice:
        db.session.delete(notice)
        db.session.commit()
        flash('通知已删除', 'success')
    return redirect(url_for('admin_notices'))

# ---------- 修改意见 ----------
@app.route('/feedback')
@login_required
def feedback_page():
    feedbacks = SystemFeedback.query.order_by(SystemFeedback.created_time.desc()).all()
    resp = make_response(render_template_string(FEEDBACK_HTML, feedbacks=feedbacks, fmt_date=fmt_date))
    resp.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate, max-age=0'
    resp.headers['Pragma'] = 'no-cache'
    resp.headers['Expires'] = '0'
    return resp

@app.route('/api/feedback/submit', methods=['POST'])
@login_required
def feedback_submit():
    data = request.get_json()
    title = (data.get('title') or '').strip()
    content = (data.get('content') or '').strip()
    category = (data.get('category') or '功能建议').strip()
    if not title or not content:
        return jsonify({'error': '标题和内容不能为空'}), 400
    fb = SystemFeedback(title=title[:200], content=content, category=category, created_by=current_user.id)
    db.session.add(fb)
    db.session.commit()
    return jsonify({'success': True, 'id': fb.id})

@app.route('/api/feedback/reply/<int:fb_id>', methods=['POST'])
@login_required
@role_required('管理员')
def feedback_reply(fb_id):
    fb = db.session.get(SystemFeedback, fb_id)
    if not fb:
        return jsonify({'error': '记录不存在'}), 404
    data = request.get_json()
    fb.reply = (data.get('reply') or '').strip()
    fb.status = (data.get('status') or '已回复').strip()
    fb.replied_by = current_user.id
    fb.replied_time = bj_now()
    db.session.commit()
    return jsonify({'success': True})

# ================== HTML 模板部分 ==================

BASE_HTML = '''
<!DOCTYPE html><html lang="zh"><head><meta charset="utf-8"><title>三分厂生产计划可视化管理系统</title>
<link href="/static/bootstrap.min.css" rel="stylesheet">
<script src="/static/chart.umd.min.js"></script>
<script src="/static/chartjs-plugin-datalabels.min.js"></script>
<style>
body{font-size:14px; background: #f3f6fa; min-height: 100vh; font-family: "Microsoft YaHei", "Segoe UI", Roboto, sans-serif; color:#1f2937;}
.navbar{background:linear-gradient(90deg,#182334,#27384f)!important; padding:.45rem .75rem;}
.navbar .container-fluid{gap:10px; align-items:center;}
.navbar-brand{font-weight:700; font-size:1.28rem; letter-spacing:0; white-space:nowrap}
.navbar .d-flex.align-items-center{gap:6px; flex-wrap:wrap; justify-content:flex-end;}
.navbar .btn{border-radius:7px!important; font-weight:600; box-shadow:none!important;}
.navbar .text-light{white-space:nowrap;}
.dropdown-menu{border:1px solid #e5e7eb; border-radius:8px; box-shadow:0 12px 28px rgba(15,23,42,.14); padding:6px;}
.dropdown-item{border-radius:6px; font-size:13px; padding:7px 10px;}
.dropdown-item:hover{background:#eef4ff; color:#1d4ed8;}
.container-fluid.px-2{padding-left:14px!important; padding-right:14px!important; margin-top:12px!important;}
.alert{border-radius:8px;}
.card,.dashboard-table-wrap,.table-outer{border-color:#e2e8f0!important; box-shadow:0 8px 22px rgba(15,23,42,.05)!important;}
.card{border-radius:8px; overflow:hidden;}
.card-header{border-bottom-color:#e5e7eb!important;}
.btn{border-radius:6px; font-weight:600;}
.btn-sm{line-height:1.35;}
.form-control,.form-select{border-color:#d6dee8; border-radius:6px;}
.form-control:focus,.form-select:focus{border-color:#3b82f6; box-shadow:0 0 0 .18rem rgba(59,130,246,.16);}
.table{--bs-table-hover-bg:#f8fafc;}
.table thead th{background:#f8fafc!important; color:#334155; font-weight:700; border-bottom-color:#dbe4ef!important;}
.table tbody td{color:#334155;}
.btn-group-actions .btn {
    background-color: #ffffff; border-radius: 4px; padding: 5px 16px; margin-right: 8px;
    font-size: 13px; font-weight: 600; letter-spacing: 0.3px; transition: all 0.2s ease;
    box-shadow: 0 1px 2px rgba(0,0,0,0.05); border: 1px solid #ccc; color: #333;
}
.btn-group-actions .btn:hover { background-color: #333; color: #fff; transform: translateY(-1px); box-shadow: 0 2px 6px rgba(0,0,0,0.08); }

.table-outer { width: 100%; height: calc(100vh - 150px); overflow: auto; border: 1px solid #dee2e6; border-radius: 8px; background-color: #fff; }
.table-inner { width: max-content; min-width: 100%; overflow: visible !important; }
#taskTable, #taskTable th, #taskTable td { border: 1px solid #ccc !important; }
#taskTable { border-collapse: collapse; border-spacing: 0; }
#taskTable th{background-color:#f8fafc !important;text-align:center;vertical-align:middle;white-space:nowrap;padding:5px 7px;font-size:14px;font-weight:700;}
#taskTable td{padding:4px 7px !important;white-space:nowrap;vertical-align:middle;font-size:14px;line-height:1.35;}
#taskTable .btn{font-size:15px !important;padding:0 6px !important;}
#taskTable .btn-del{padding:0 6px !important;font-size:15px !important;}
#taskTable .badge{font-size:15px !important;padding:0 6px !important;vertical-align:middle !important;line-height:1.1;display:inline-block !important;}
#taskTable thead th { position: sticky; top: 0; z-index: 10; border-bottom: 2px solid #94a3b8 !important; border-top: none !important; }
#taskTable td:first-child, #taskTable th:first-child { position: sticky; left: 0; z-index: 11; background-color: #fff; }
#taskTable th:first-child { background-color: #f8fafc !important; z-index: 12; }
#taskTable td:nth-child(2), #taskTable th:nth-child(2) { border-left: 1px solid #ccc !important; }

#reportTable, #reportTable th, #reportTable td { border: 1px solid #ccc !important; }
#reportTable { border-collapse: collapse; border-spacing: 0; }
#reportTable th{background-color:#f8fafc !important;text-align:center;vertical-align:middle;white-space:nowrap;padding:4px 8px;font-size:14px;font-weight:bold;}
#reportTable td{padding:4px 8px !important;white-space:nowrap;vertical-align:middle;font-size:14px;line-height:1.3;}
#reportTable thead th { position: sticky; top: 0; z-index: 10; }
#reportTable td:first-child, #reportTable th:first-child { position: sticky; left: 0; z-index: 11; background-color: #fff; }
#reportTable th:first-child { background-color: #f8fafc !important; z-index: 12; }
#reportTable .can-edit{cursor:pointer}
#reportTable .can-edit:hover{background-color:#e8f0fe}

.table td.seq{text-align:center;font-weight:bold}
.can-edit{cursor:pointer}
.row-issue{background-color:#fff0f0 !important}
.center-text{text-align:center !important}
.selected-cell{background-color:#d4edda !important; box-shadow: inset 0 0 0 2px #198754;}
tr.highlight-row td{background-color:#f0f8ff !important}
tr.highlight-row td.highlight-col{background-color:#fff3cd !important}
tr.highlight-row td.selected-cell{background-color:#d4edda !important; box-shadow: inset 0 0 0 2px #198754;}
td.highlight-col{background-color:#fff3cd !important}
.btn-del { background: none; border: 1px solid #c0392b; color: #c0392b; border-radius: 4px; padding: 2px 12px; font-size: 12px; transition: all 0.2s; }
.btn-del:hover { background: #c0392b; color: white; }

.card-metric { border: 1px solid #e2e8f0; border-radius: 8px; background: #fff; margin-bottom: 0; box-shadow: 0 8px 20px rgba(15,23,42,0.05); position: relative; overflow: hidden; }
.card-metric .card-body { display: flex; justify-content: space-between; align-items: center; padding: 0.9rem; position: relative; z-index: 2; }
.card-metric .metric-value { font-size: 1.6rem; font-weight: bold; color: #1e293b; margin-top: 5px; }
.card-metric .metric-label { font-size: 0.9rem; color: #64748b; font-weight:600; }
.metric-icon-bg { position: absolute; right: 10px; top: 50%; transform: translateY(-50%); font-size: 3.5rem; opacity: 0.08; z-index: 1; user-select: none; }

.section-title { font-size: 15px; font-weight: bold; color: #1e293b; margin-bottom: 12px; border-left: 4px solid #3b82f6; padding-left: 10px;}
.dashboard-table-wrap { background: #fff; border: 1px solid #e2e8f0; border-radius: 8px; padding: 12px; box-shadow: 0 2px 6px rgba(0,0,0,0.03); margin-bottom: 24px;}

.dashboard-table { border-collapse: separate !important; border-spacing: 0; width: 100%; background-color:#fff; border: 1px solid #e2e8f0; border-radius: 6px; overflow: hidden; }
.dashboard-table thead th { background-color: #f8fafc !important; color: #475569; border-bottom: 1px solid #e2e8f0 !important; border-top: none !important; }
.dashboard-table th, .dashboard-table td { border: none !important; border-bottom: 1px solid #f1f5f9 !important; text-align: center; padding: 5px 8px; vertical-align: middle; font-size: 13px; height: 34px; }
.dashboard-table tbody tr:last-child td { border-bottom: none !important; }
.dashboard-table tbody tr:hover td { background-color: #f8fafc !important; }

.nav-tabs .nav-link { color: #64748b; font-weight: 500; border-radius: 6px 6px 0 0; }
.nav-tabs .nav-link.active { font-weight: bold; color: #0f172a; border-bottom-color: #fff; background-color: #fff; }

.btn-circle-page { width: 26px; height: 26px; padding: 0; border-radius: 50%; display: inline-flex; align-items: center; justify-content: center; font-weight: bold; border: 1px solid #cbd5e1; background: #fff; color: #475569; transition: all 0.2s; font-size: 12px; margin: 0 4px;}
.btn-circle-page:hover:not(:disabled) { background: #e2e8f0; color: #0f172a; }
.btn-circle-page:disabled { opacity: 0.5; cursor: not-allowed; }

.table td, .table th { white-space: nowrap; }
.dashboard-table td, .dashboard-table th { white-space: nowrap; }
.pagination-nav { flex-shrink: 0; }
.card-metric { flex-shrink: 0; }
#colSortList .list-group-item.drag-over { background-color: #e8f0fe; border-color: #3b82f6; }
#colSortList .list-group-item { cursor: default; user-select: none; }
.sub-dropdown { display: none; position: absolute; left: 100%; top: 0; margin-top: -1px; }
.dropend:hover > .sub-dropdown { display: block; }
.dropend { position: relative; }
@media (max-width: 1200px){
  .navbar .container-fluid{align-items:flex-start;}
  .navbar-brand{font-size:1.08rem; margin-top:4px;}
  .navbar .d-flex.align-items-center{justify-content:flex-start;}
}
@media (max-width: 768px){
  body{font-size:13px;}
  .container-fluid.px-2{padding-left:8px!important; padding-right:8px!important;}
  .navbar-brand{white-space:normal; line-height:1.25;}
  h4{font-size:1.05rem;}
  .table-outer{height:calc(100vh - 185px);}
}
</style>
</head><body>
<nav class="navbar navbar-expand-lg navbar-dark bg-dark shadow-sm"><div class="container-fluid">
<a class="navbar-brand" href="/">三分厂生产计划可视化管理系统</a>
<div class="d-flex align-items-center">
{% if current_user.is_authenticated %}
<a class="btn btn-outline-light btn-sm me-2 px-3 rounded-pill" href="/tasks">任务列表</a>
{% if current_user.role.name in ('计调员','管理员') %}
<a class="btn btn-outline-primary btn-sm me-2 px-3 rounded-pill" href="/planning-report">计划导入</a>
{% endif %}
<div class="btn-group me-2">
  <a class="btn btn-outline-info btn-sm dropdown-toggle rounded-pill px-3" data-bs-toggle="dropdown" href="#">功能模块</a>
  <ul class="dropdown-menu">
    {% if current_user.role.name in ('部门领导','管理员') %}
    <li class="dropend">
      <a class="dropdown-item dropdown-toggle" href="#" data-bs-toggle="dropdown">领导模块</a>
      <ul class="dropdown-menu sub-dropdown">
        <li><a class="dropdown-item" href="/dashboard">计划看板</a></li>
        <li><a class="dropdown-item" href="/datacenter">数据中心</a></li>
        <li><a class="dropdown-item" href="/admin/notices">发布通知</a></li>
      </ul>
    </li>
    <li><hr class="dropdown-divider"></li>
    {% endif %}
    {% if current_user.role.name in ('计调员','管理员') %}
    <li class="dropend">
      <a class="dropdown-item dropdown-toggle" href="#" data-bs-toggle="dropdown">计调模块</a>
      <ul class="dropdown-menu sub-dropdown">
        <li><a class="dropdown-item" href="/planning-dashboard">计调看板</a></li>
        <li><a class="dropdown-item" href="/planning-task-report">计调任务填报</a></li>
      </ul>
    </li>
    <li><hr class="dropdown-divider"></li>
    {% endif %}
    <li><a class="dropdown-item" href="/teamkanban">班组模块</a></li>
  </ul>
</div>
{% if current_user.role.name == '管理员' %}
<div class="btn-group me-2">
  <a class="btn btn-outline-warning btn-sm dropdown-toggle rounded-pill px-3" data-bs-toggle="dropdown" href="#">用户</a>
  <ul class="dropdown-menu">
    <li><a class="dropdown-item" href="/admin/users">用户管理</a></li>
    <li><a class="dropdown-item" href="/admin/permissions">权限设置</a></li>
  </ul>
</div>
<a class="btn btn-outline-warning btn-sm me-2 px-3 rounded-pill" href="/admin/logs">操作日志</a>
{% endif %}
<a class="btn btn-outline-secondary btn-sm me-2 px-3 rounded-pill" href="/feedback">💬 修改意见</a>
<span class="text-light me-3 opacity-75">👤 {{ current_user.username or current_user.account }} ({{ current_user.role.name }})</span>
<a class="btn btn-danger btn-sm rounded-pill px-3" href="/logout">退出</a>
{% endif %}
</div></div></nav>
<div class="container-fluid px-2" style="margin-top:8px;">
{% with msgs = get_flashed_messages(with_categories=true) %}{% if msgs %}{% for c,m in msgs %}<div class="alert alert-{{c}} alert-dismissible fade show border-0 shadow-sm">{{m}}<button class="btn-close" data-bs-dismiss="alert"></button></div>{% endfor %}{% endif %}{% endwith %}
{% block content %}{% endblock %}
</div>
<script src="/static/bootstrap.bundle.min.js"></script>
<script>
function parseDate(val) {
    if (!val) return null;
    var s = val.replace(/\//g, '-');
    var d = new Date(s);
    if (isNaN(d.getTime())) return null;
    return d;
}
function makeEditable(td) {
  if (!td.classList.contains('can-edit')) return;
  var oldText = td.innerText;
  var field = td.getAttribute('data-field');
  var id = td.getAttribute('data-id');
  var input;
  
  if (['liquidated_damages','help_letter','first_article','fixed_submit_item'].includes(field)) {
    input = document.createElement('select');
    input.className = "form-select form-select-sm";
    input.innerHTML = '<option value="">-</option><option value="是">是</option><option value="否">否</option>';
    input.value = oldText;
  } else if (field === 'product_category') {
    input = document.createElement('select');
    input.className = "form-select form-select-sm";
    input.innerHTML = '<option value="">-</option><option value="服装">服装</option><option value="头盔">头盔</option><option value="面罩">面罩</option><option value="船囊">船囊</option><option value="救生衣">救生衣</option>';
    input.value = oldText;
  } else if (field.includes('dept')) {
    input = document.createElement('select');
    input.className = "form-select form-select-sm";
    var options = {{ dept_options | tojson }};
    input.innerHTML = '<option value="">-</option>' + options.map(d => `<option value="${d}">${d}</option>`).join('');
    input.value = oldText;
  } else {
    input = document.createElement('input');
    input.type = 'text';
    input.className = "form-control form-control-sm";
    input.value = oldText;
  }

  td.innerHTML = '';
  td.appendChild(input);
  input.focus();
  function save() {
    var newVal = input.value;
    fetch('/api/task/' + id + '/field', {
      method: 'PUT',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({field: field, value: newVal})
    })
    .then(res => res.json())
    .then(data => {
      if (data.success) { 
        td.innerText = data.new_value || '';
        if (field === 'batch_no') {
          var tr = td.parentNode;
          var stageCells = tr.querySelectorAll('td[data-field="stage"]');
          if (stageCells.length > 0) {
            var bn = newVal || '';
            var stage = '';
            if (bn) {
              var fl = bn.trim().charAt(0).toUpperCase();
              if ('CFSZ'.indexOf(fl) >= 0) stage = '试制';
              else if ('DP'.indexOf(fl) >= 0) stage = '批产';
            }
            stageCells[0].innerText = stage;
          }
        }
        if (data.first_article) {
          var tr = td.parentNode;
          var faCells = tr.querySelectorAll('td[data-field="first_article"]');
          if (faCells.length > 0) faCells[0].innerText = data.first_article;
        }
      }
      else { alert(data.error || '保存失败'); td.innerText = oldText; }
      td.classList.add('can-edit');
    });
  }
  input.addEventListener('blur', save);
  input.addEventListener('keypress', function(e) { if (e.key === 'Enter') save(); });
}
</script>
</body></html>
'''

LOGIN_HTML = BASE_HTML.replace('{% block content %}{% endblock %}','''
{% block content %}<div class="row justify-content-center mt-5"><div class="col-md-4 p-4 shadow-sm border bg-white rounded">
<h4 class="text-center mb-4 text-dark fw-bold">登录三分厂生产计划可视化管理系统</h4>
<form method="post"><div class="mb-3"><input class="form-control" name="account" placeholder="账号" required></div>
<div class="mb-3"><input type="password" class="form-control" name="password" placeholder="密码" required></div>
<button class="btn btn-primary w-100">登录</button></form></div></div>{% endblock %}
''')

TASK_LIST_HTML = BASE_HTML.replace('{% block content %}{% endblock %}','''
{% block content %}
<div class="d-flex justify-content-between align-items-center mb-3">
  <div class="btn-group-actions">
    <button id="clearFilterBtn" class="btn">✕ 取消筛选</button>
    <button id="colSettingsBtn" class="btn">⚙ 列设置</button>
    <a href="/export/doc" class="btn">📥 导出Excel</a>
  </div>
  <div>
    <input type="text" id="searchInput" class="form-control form-control-sm" style="width: 250px;" placeholder="全局模糊搜索...">
  </div>
</div>

<div class="table-outer shadow-sm">
  <div class="table-inner">
    <table class="table table-bordered table-sm" id="taskTable">
    <thead>
        <tr>
            {% for c,n in columns %}
                <th data-field="{{c}}">{{ n }}</th>
            {% endfor %}
            <th>操作</th>
        </tr>
    </thead>
    <tbody>
    {% for row in rows %}
    <tr class="{% if row.has_issue %}row-issue{% endif %}" data-task-id="{{ row['id'] }}">
      {% for c,n in columns %}
        {% if c in ['has_tech','has_mgmt','has_shortage'] %}
          <td class="text-center">{% if row[c] %}<span class="text-danger fw-bold" style="font-size:15px;">✕</span>{% else %}<span class="text-success fw-bold" style="font-size:15px;">✓</span>{% endif %}</td>
        {% elif c == 'production_status_calc' %}
          <td class="text-center" title="{{ row['hover_text'] }}" style="cursor: help;">
            {% if row.production_status_calc == '未开工' %}
                <span class="badge bg-danger">未开工</span>
            {% elif row.production_status_calc == '生产中' %}
                <span class="badge bg-warning text-dark">生产中</span>
            {% else %}
                <span class="badge bg-success">已完成</span>
            {% endif %}
          </td>
        {% else %}
          <td data-field="{{c}}" data-id="{{row['id']}}" class="{% if c=='serial_no' %}seq{% endif %} {% if c in edit_cols %}can-edit{% endif %} {% if c in center_cols %}center-text{% endif %}" {% if c != '__neveredit__' %}ondblclick="makeEditable(this)"{% endif %}>{{ row[c] if c in row else '' }}</td>
        {% endif %}
      {% endfor %}
      <td class="text-center" style="white-space:nowrap;">
        <a href="/task/{{row['id']}}/details" class="btn btn-outline-secondary btn-sm py-0">详情</a>
        {% if current_user.role.name == '管理员' %}
        <button class="btn-del btn-sm" data-task-id="{{row['id']}}">删除</button>
        {% endif %}
      </td>
    </tr>
    {% endfor %}
    </tbody>
    </table>
  </div>
</div>

<!-- 列设置弹窗 -->
<div class="modal fade" id="colSettingsModal" tabindex="-1">
  <div class="modal-dialog modal-sm">
    <div class="modal-content">
      <div class="modal-header"><h5 class="modal-title">列设置</h5><button type="button" class="btn-close" data-bs-dismiss="modal"></button></div>
      <div class="modal-body">
        <div class="d-flex gap-1 mb-2">
          <button type="button" class="btn btn-outline-secondary btn-sm" id="colSelectAll">全选</button>
          <button type="button" class="btn btn-outline-secondary btn-sm" id="colSelectNone">全不选</button>
          <button type="button" class="btn btn-outline-secondary btn-sm" id="colResetDefault">恢复默认</button>
        </div>
        <div id="colSortList" class="list-group" style="max-height:450px;overflow-y:auto;"></div>
      </div>
      <div class="modal-footer"><button type="button" class="btn btn-secondary" data-bs-dismiss="modal">取消</button><button type="button" class="btn btn-primary" id="colApplyBtn">应用</button></div>
    </div>
  </div>
</div>

<script>
let selectedCell = null;
let selectedColIndex = null;

function cellClickHandler(e) {
  if (selectedCell) {
    selectedCell.classList.remove('selected-cell');
    selectedCell.parentNode.classList.remove('highlight-row');
    document.querySelectorAll('#taskTable tbody tr').forEach(r => {
      if (r.cells[selectedColIndex]) r.cells[selectedColIndex].classList.remove('highlight-col');
    });
  }
  selectedCell = this;
  selectedColIndex = this.cellIndex;
  selectedCell.parentNode.classList.add('highlight-row');
  selectedCell.classList.add('selected-cell');
  document.querySelectorAll('#taskTable tbody tr').forEach(r => {
    if (r.cells[selectedColIndex]) r.cells[selectedColIndex].classList.add('highlight-col');
  });
}
document.querySelectorAll('#taskTable tbody td.can-edit').forEach(td => td.addEventListener('click', cellClickHandler));

document.addEventListener("DOMContentLoaded", function() {
    const table = document.getElementById("taskTable");
    const headers = table.querySelectorAll("thead th");
    const rows = Array.from(table.querySelectorAll("tbody tr"));
    window._taskFilters = {}; 

    headers.forEach((th, colIndex) => {
        if (colIndex === headers.length - 1) return;

        let icon = document.createElement("span");
        icon.innerHTML = " 🔽";
        icon.style.cursor = "pointer";
        icon.style.fontSize = "10px";
        icon.onclick = function(e) {
            e.stopPropagation();
            showFilterMenu(th, colIndex, e.pageX, e.pageY);
        };
        th.appendChild(icon);
    });

    function showFilterMenu(th, colIndex, pageX, pageY) {
        let existing = document.getElementById("filterMenu");
        if(existing) existing.remove();

        let uniqueValues = new Set();
        rows.forEach(r => {
            let td = r.cells[colIndex];
            if(td) {
                let text = td.innerText.trim();
                uniqueValues.add(text);
            }
        });
        let sortedValues = Array.from(uniqueValues).sort();

        let menu = document.createElement("div");
        menu.id = "filterMenu";
        menu.style.position = "absolute";
        menu.style.left = pageX + "px";
        menu.style.top = pageY + "px";
        menu.style.background = "#fff";
        menu.style.border = "1px solid #ccc";
        menu.style.padding = "10px";
        menu.style.zIndex = "9999";
        menu.style.maxHeight = "400px";
        menu.style.overflowY = "auto";
        menu.style.boxShadow = "0 4px 8px rgba(0,0,0,0.1)";
        menu.style.minWidth = "180px";

        let currentFilter = window._taskFilters[colIndex] || { text: '', values: uniqueValues };
        let currentText = currentFilter.text || '';
        let currentValues = currentFilter.values || uniqueValues;

        let html = '<div><input type="text" id="filterTextInput" class="form-control form-control-sm" placeholder="输入关键词筛选..." value="' + currentText.replace(/"/g,'&quot;') + '" style="margin-bottom:6px;"></div>';
        html += '<div><label><input type="checkbox" id="selectAllFilter" checked> <strong>(全选)</strong></label></div><hr style="margin:5px 0;">';
        
        sortedValues.forEach(val => {
            let checked = currentValues.has(val) ? "checked" : "";
            let displayVal = val || '(空白)';
            html += '<div><label><input type="checkbox" class="filter-val" value="' + val.replace(/"/g,'&quot;') + '" ' + checked + '> ' + displayVal + '</label></div>';
        });
        html += '<hr style="margin:5px 0;"><button class="btn btn-sm btn-primary w-100" id="applyFilterBtn">确定</button>';
        html += '<button class="btn btn-sm btn-outline-secondary w-100 mt-1" id="clearFilterBtn">清除筛选</button>';
        menu.innerHTML = html;
        document.body.appendChild(menu);

        let selectAll = menu.querySelector('#selectAllFilter');
        let checkboxes = menu.querySelectorAll('.filter-val');
        let textInput = menu.querySelector('#filterTextInput');
        
        let allChecked = Array.from(checkboxes).every(c => c.checked);
        selectAll.checked = allChecked;

        selectAll.onchange = function() {
            checkboxes.forEach(cb => cb.checked = selectAll.checked);
        };
        checkboxes.forEach(cb => {
            cb.onchange = function() {
                if(!cb.checked) selectAll.checked = false;
                else {
                    let allC = Array.from(checkboxes).every(c => c.checked);
                    if(allC) selectAll.checked = true;
                }
            };
        });

        // Free-text input: filter checkbox list in real-time
        textInput.addEventListener('input', function() {
            let q = textInput.value.toLowerCase();
            checkboxes.forEach(cb => {
                let label = cb.parentNode;
                let row = label.parentNode;
                let val = (cb.value || '(空白)').toLowerCase();
                row.style.display = val.includes(q) ? '' : 'none';
            });
        });

        menu.querySelector('#applyFilterBtn').onclick = function() {
            let selected = new Set();
            checkboxes.forEach(cb => { if(cb.checked) selected.add(cb.value); });
            let txtVal = textInput.value.trim();
            
            if(selected.size === uniqueValues.size && !txtVal) { delete window._taskFilters[colIndex]; } 
            else { window._taskFilters[colIndex] = { text: txtVal, values: selected }; }
            applyFilters();
            menu.remove();
        };

        menu.querySelector('#clearFilterBtn').onclick = function() {
            delete window._taskFilters[colIndex];
            applyFilters();
            menu.remove();
        };

        setTimeout(() => {
            textInput.focus();
            document.addEventListener('click', function closeMenu(e) {
                if(!menu.contains(e.target)) {
                    menu.remove();
                    document.removeEventListener('click', closeMenu);
                }
            });
        }, 10);
    }

    function applyFilters() {
        let search = document.getElementById('searchInput').value.toLowerCase();
        rows.forEach(r => {
            let show = true;
            for(let col in window._taskFilters) {
                let td = r.cells[col];
                if(td) {
                    let text = td.innerText.trim();
                    let f = window._taskFilters[col];
                    let matchText = true, matchValues = true;
                    if (f.text) {
                        matchText = text.toLowerCase().includes(f.text.toLowerCase());
                    }
                    if (f.values && f.values.size > 0) {
                        matchValues = f.values.has(text);
                    }
                    if (!matchText || !matchValues) { show = false; break; }
                }
            }
            if (show && search) {
                if (!r.innerText.toLowerCase().includes(search)) show = false;
            }
            r.style.display = show ? "" : "none";
        });
    }
    
    document.getElementById('searchInput').addEventListener('keyup', applyFilters);
    document.getElementById('clearFilterBtn').addEventListener('click', function() {
        window._taskFilters = {};
        document.getElementById('searchInput').value = '';
        applyFilters();
    });
});

// ===== 列设置功能 =====
(function() {
  var STORAGE_KEY = 'taskTableColSettings';
  var table = document.getElementById('taskTable');

  function getDefaultOrder() {
    var order = [];
    table.querySelectorAll('thead th').forEach(function(th) {
      var f = th.getAttribute('data-field');
      if (f) order.push({ field: f, name: th.innerText.trim(), visible: true });
    });
    return order;
  }

  function loadSettings() {
    try { return JSON.parse(localStorage.getItem(STORAGE_KEY)); } catch(e) { return null; }
  }

  function saveSettings(s) {
    localStorage.setItem(STORAGE_KEY, JSON.stringify(s));
  }

  function buildList(settings) {
    var list = document.getElementById('colSortList');
    list.innerHTML = '';
    settings.forEach(function(item, i) {
      if (!item.field) return;
      var div = document.createElement('div');
      div.className = 'list-group-item d-flex align-items-center py-1 px-2';
      div.setAttribute('draggable', 'true');
      div.dataset.field = item.field;
      div.innerHTML = '<span class="me-1 text-muted" style="cursor:grab;font-size:12px;">☰</span>' +
        '<input type="checkbox" class="form-check-input me-2 col-chk" ' + (item.visible !== false ? 'checked' : '') + '>' +
        '<span style="font-size:13px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;">' + item.name + '</span>';

      div.addEventListener('dragstart', function(e) {
        e.dataTransfer.effectAllowed = 'move';
        div.style.opacity = '0.4';
        e.dataTransfer.setData('text/plain', item.field);
      });
      div.addEventListener('dragend', function() { div.style.opacity = '1'; });
      div.addEventListener('dragover', function(e) { e.preventDefault(); div.classList.add('drag-over'); });
      div.addEventListener('dragleave', function() { div.classList.remove('drag-over'); });
      div.addEventListener('drop', function(e) {
        e.preventDefault();
        div.classList.remove('drag-over');
        var fromField = e.dataTransfer.getData('text/plain');
        var fromEl = list.querySelector('[data-field="' + fromField + '"]');
        if (!fromEl || fromEl === div) return;
        var nodes = Array.from(list.children);
        var fromIdx = nodes.indexOf(fromEl);
        var toIdx = nodes.indexOf(div);
        if (fromIdx < toIdx) {
          list.insertBefore(fromEl, div.nextSibling);
        } else {
          list.insertBefore(fromEl, div);
        }
      });
      list.appendChild(div);
    });
  }

  function readSettings() {
    var items = document.getElementById('colSortList').children;
    var settings = [];
    for (var i = 0; i < items.length; i++) {
      var el = items[i];
      var chk = el.querySelector('.col-chk');
      var span = el.querySelector('span:last-child');
      settings.push({ field: el.dataset.field, name: span ? span.innerText : '', visible: chk.checked });
    }
    return settings;
  }

  function applyColumnOrder(settings) {
    var theadTr = table.querySelector('thead tr');
    var rows = table.querySelectorAll('tbody tr');
    var headerArr = Array.from(table.querySelectorAll('thead th'));

    // Build field->element maps
    var fieldToTh = {};
    headerArr.forEach(function(th) { var f = th.getAttribute('data-field'); if (f) fieldToTh[f] = th; });

    var orderedFields = [];
    var visibleSet = {};
    settings.forEach(function(s) { orderedFields.push(s.field); if (s.visible) visibleSet[s.field] = true; });

    // Reorder headers
    headerArr.forEach(function(th) { theadTr.removeChild(th); });
    orderedFields.forEach(function(f) { if (fieldToTh[f]) theadTr.appendChild(fieldToTh[f]); });
    theadTr.appendChild(headerArr[headerArr.length - 1]); // 操作列最后

    // Reorder & show/hide body cells
    rows.forEach(function(r) {
      var cells = Array.from(r.cells);
      var fieldToCell = {};
      cells.forEach(function(td) { var f = td.getAttribute('data-field'); if (f) fieldToCell[f] = td; });
      cells.forEach(function(td) { r.removeChild(td); });
      orderedFields.forEach(function(f) {
        if (fieldToCell[f]) {
          r.appendChild(fieldToCell[f]);
          fieldToCell[f].style.display = visibleSet[f] ? '' : 'none';
        }
      });
      r.appendChild(cells[cells.length - 1]); // 操作列
    });

    // Show/hide headers
    var newHeaders = table.querySelectorAll('thead th');
    newHeaders.forEach(function(th) {
      var f = th.getAttribute('data-field');
      if (f) th.style.display = visibleSet[f] ? '' : 'none';
    });

    saveSettings(settings);
  }

  // Button events
  document.getElementById('colSettingsBtn').addEventListener('click', function() {
    var settings = loadSettings();
    if (!settings || settings.length === 0) settings = getDefaultOrder();
    // Merge any new columns
    var currentFields = {};
    table.querySelectorAll('thead th').forEach(function(th) { var f = th.getAttribute('data-field'); if (f) currentFields[f] = th.innerText.trim(); });
    var existing = new Set(settings.map(function(s) { return s.field; }));
    for (var f in currentFields) {
      if (!existing.has(f)) settings.push({ field: f, name: currentFields[f], visible: true });
    }
    // Remove settings for non-existent columns
    settings = settings.filter(function(s) { return currentFields[s.field]; });
    buildList(settings);
    var modal = new bootstrap.Modal(document.getElementById('colSettingsModal'));
    modal.show();
  });

  document.getElementById('colSelectAll').addEventListener('click', function() {
    document.querySelectorAll('#colSortList .col-chk').forEach(function(c) { c.checked = true; });
  });
  document.getElementById('colSelectNone').addEventListener('click', function() {
    document.querySelectorAll('#colSortList .col-chk').forEach(function(c) { c.checked = false; });
  });
  document.getElementById('colResetDefault').addEventListener('click', function() {
    buildList(getDefaultOrder());
  });
  document.getElementById('colApplyBtn').addEventListener('click', function() {
    var settings = readSettings();
    var modal = bootstrap.Modal.getInstance(document.getElementById('colSettingsModal'));
    applyColumnOrder(settings);
    if (modal) modal.hide();
  });
})();

document.addEventListener('paste', async function(e) {
  // If any modal backdrop is visible, skip table paste
  if (document.querySelector('.modal-backdrop')) return;
  // Let textareas and inputs handle paste natively
  if (e.target.tagName === 'TEXTAREA' || e.target.tagName === 'INPUT') return;

  if (!selectedCell) return;
  e.preventDefault();
  const text = (e.clipboardData || window.clipboardData).getData('text');
  const rowsData = text.split(/\\r?\\n/).filter(line => line.trim() !== '');
  const matrix = rowsData.map(line => line.split('\\t'));
  
  const startTr = selectedCell.parentNode;
  const tbody = document.querySelector('#taskTable tbody');
  const dataRows = Array.from(tbody.rows);
  const startRowIndex = dataRows.indexOf(startTr);
  const startColIndex = selectedCell.cellIndex;
  
  const updates = [];

  for (let r = 0; r < matrix.length; r++) {
    const targetRowIndex = startRowIndex + r;
    if (targetRowIndex >= dataRows.length) break;
    const targetRow = dataRows[targetRowIndex];
    const taskId = targetRow.getAttribute('data-task-id');
    for (let c = 0; c < matrix[r].length; c++) {
      const targetColIndex = startColIndex + c;
      const targetCell = targetRow.cells[targetColIndex];
      if (targetCell && targetCell.classList.contains('can-edit')) {
        const field = targetCell.getAttribute('data-field');
        updates.push({ task_id: parseInt(taskId), field: field, value: matrix[r][c] });
      }
    }
  }
  
  if (updates.length > 0) {
      const res = await fetch('/api/task/batch_update', { method: 'POST', headers: {'Content-Type': 'application/json'}, body: JSON.stringify({updates: updates}) });
      const result = await res.json();
      if (result.success) {
        result.results.forEach(r => {
          let td = document.querySelector(`td[data-field="${r.field}"][data-id="${r.task_id}"]`);
          if (td) {
            td.innerText = r.new_value;
            if (r.first_article) {
              let faTd = document.querySelector(`td[data-field="first_article"][data-id="${r.task_id}"]`);
              if (faTd) faTd.innerText = r.first_article;
            }
            if (r.stage) {
              let stTd = document.querySelector(`td[data-field="stage"][data-id="${r.task_id}"]`);
              if (stTd) stTd.innerText = r.stage;
            }
          }
        });
      }
      if (selectedCell) {
        selectedCell.classList.remove('selected-cell');
        selectedCell.parentNode.classList.remove('highlight-row');
        document.querySelectorAll('#taskTable tbody tr').forEach(r => {
          if (r.cells[selectedColIndex]) r.cells[selectedColIndex].classList.remove('highlight-col');
        });
        selectedCell = null;
        selectedColIndex = null;
      }
  }
});


document.querySelectorAll('.btn-del').forEach(btn => btn.addEventListener('click', async function(){
    if(!confirm('确认删除？')) return;
    const res = await fetch('/api/task/'+this.getAttribute('data-task-id')+'/delete', {method:'POST'});
    if((await res.json()).success) this.closest('tr').remove();
}));

// Scroll to and select newly added task after batch import
(function() {
    var params = new URLSearchParams(window.location.search);
    var newId = params.get('new');
    if (newId) {
        var row = document.querySelector('#taskTable tbody tr[data-task-id="' + newId + '"]');
        if (row) {
            row.scrollIntoView({ behavior: 'smooth', block: 'center' });
            var editableCell = row.querySelector('td.can-edit');
            if (editableCell) {
                editableCell.click();
            }
        }
        history.replaceState(null, '', window.location.pathname);
    }
})();
</script>
{% endblock %}
''')

DETAILS_HTML = BASE_HTML.replace('{% block content %}{% endblock %}','''
{% block content %}
<div class="mb-4">
  <div class="section-title">📋 任务详情：{{ task.product_name }} ({{task.product_draw_no}})</div>
  <div class="dashboard-table-wrap">
    <form method="post" id="detailsForm">
      <h6 class="fw-bold mb-3">技术/管理问题</h6>
      <table class="dashboard-table text-center" id="issueTable">
        <thead><tr><th>类型</th><th>内容</th><th>部门</th><th>提出</th><th>完成</th><th></th></tr></thead>
        <tbody>
          {% for r in task.tech_mgmt_issues %}
          <tr>
            <td><select class="form-select form-select-sm" name="issue_type_{{r.id}}"><option value="tech" {% if r.issue_type=='tech' %}selected{% endif %}>技术</option><option value="mgmt" {% if r.issue_type=='mgmt' %}selected{% endif %}>管理</option></select></td>
            <td><input class="form-control form-control-sm" name="issue_content_{{r.id}}" value="{{r.content or ''}}"></td>
            <td><select class="form-select form-select-sm" name="issue_dept_{{r.id}}"><option value="">-</option>{% for d in dept_options %}<option value="{{d}}" {% if r.dept==d %}selected{% endif %}>{{d}}</option>{% endfor %}</select></td>
            <td><input class="form-control form-control-sm" name="issue_raise_{{r.id}}" value="{{fmt_date(r.raise_time)}}"></td>
            <td><input class="form-control form-control-sm" name="issue_finish_{{r.id}}" value="{{fmt_date(r.finish_time)}}"></td>
            <td><input type="hidden" name="issue_id_{{loop.index0}}" value="{{r.id}}"><button type="button" class="btn btn-sm btn-outline-danger py-0" onclick="removeIssueRow(this,{{r.id}})">✕</button></td>
          </tr>
          {% endfor %}
        </tbody>
      </table>
      <button type="button" class="btn btn-sm btn-secondary mb-4" onclick="addIssueRow()">+ 添加问题</button>

      <h6 class="fw-bold mb-3">缺件记录</h6>
      <table class="dashboard-table text-center" id="shTable">
        <thead><tr><th>类别</th><th>内容</th><th>下发</th><th>报缺</th><th>到位</th><th></th></tr></thead>
        <tbody>
          {% for r in task.shortage_records %}
          <tr>
            <td><select class="form-select form-select-sm" name="sh_type_{{r.id}}">{% for t in shortage_types %}<option value="{{t}}" {% if r.shortage_type==t %}selected{% endif %}>{{t}}</option>{% endfor %}</select></td>
            <td><input class="form-control form-control-sm" name="sh_content_{{r.id}}" value="{{r.content or ''}}"></td>
            <td><input class="form-control form-control-sm" name="sh_send_{{r.id}}" value="{{fmt_date(r.send_time)}}"></td>
            <td><input class="form-control form-control-sm" name="sh_report_{{r.id}}" value="{{fmt_date(r.report_time)}}"></td>
            <td><input class="form-control form-control-sm" name="sh_arrive_{{r.id}}" value="{{fmt_date(r.arrive_time)}}"></td>
            <td><input type="hidden" name="sh_id_{{loop.index0}}" value="{{r.id}}"><button type="button" class="btn btn-sm btn-outline-danger py-0" onclick="removeShRow(this,{{r.id}})">✕</button></td>
          </tr>
          {% endfor %}
        </tbody>
      </table>
      <button type="button" class="btn btn-sm btn-secondary" onclick="addShRow()">+ 添加缺件</button>

      <div class="mt-4 pt-3" style="border-top:1px solid #dee2e6;">
        <input type="hidden" name="delete_issues" id="delete_issues" value="">
        <input type="hidden" name="delete_shortages" id="delete_shortages" value="">
        <button class="btn btn-primary">💾 保存详情信息</button> 
        <a href="/tasks" class="btn btn-outline-secondary">返回列表</a>
      </div>
    </form>
  </div>
</div>
<script>
function addIssueRow() {
    let row = document.querySelector('#issueTable tbody').insertRow();
    row.innerHTML = `<td><select class="form-select form-select-sm" name="new_issue_type[]"><option value="tech">技术</option><option value="mgmt">管理</option></select></td>
    <td><input class="form-control form-control-sm" name="new_issue_content[]"></td>
    <td><select class="form-select form-select-sm" name="new_issue_dept[]"><option value="">-</option>{% for d in dept_options %}<option value="{{d}}">{{d}}</option>{% endfor %}</select></td>
    <td><input class="form-control form-control-sm" name="new_issue_raise[]" value="{{today}}"></td>
    <td><input class="form-control form-control-sm" name="new_issue_finish[]"></td>
    <td><button type="button" class="btn btn-sm btn-outline-danger py-0" onclick="this.closest('tr').remove()">✕</button></td>`;
}
function addShRow() {
    let row = document.querySelector('#shTable tbody').insertRow();
    row.innerHTML = `<td><select class="form-select form-select-sm" name="new_sh_type[]">{% for t in shortage_types %}<option value="{{t}}">{{t}}</option>{% endfor %}</select></td>
    <td><input class="form-control form-control-sm" name="new_sh_content[]"></td>
    <td><input class="form-control form-control-sm" name="new_sh_send[]" value="{{today}}"></td>
    <td><input class="form-control form-control-sm" name="new_sh_report[]"></td>
    <td><input class="form-control form-control-sm" name="new_sh_arrive[]"></td>
    <td><button type="button" class="btn btn-sm btn-outline-danger py-0" onclick="this.closest('tr').remove()">✕</button></td>`;
}
function removeIssueRow(btn, id) {
    var del = document.getElementById('delete_issues');
    var ids = del.value ? del.value.split(',') : [];
    ids.push(String(id));
    del.value = ids.join(',');
    btn.closest('tr').remove();
}
function removeShRow(btn, id) {
    var del = document.getElementById('delete_shortages');
    var ids = del.value ? del.value.split(',') : [];
    ids.push(String(id));
    del.value = ids.join(',');
    btn.closest('tr').remove();
}
</script>
{% endblock %}
''')

DASHBOARD_HTML = BASE_HTML.replace('{% block content %}{% endblock %}','''
{% block content %}
<style>
/* ---------- 看板三列 Flex 布局 ---------- */
.dashboard-row {
  display: flex;
  flex-wrap: nowrap;
  gap: 1rem;
  min-height: calc(100vh - 180px);
}
.dashboard-col {
  display: flex;
  flex-direction: column;
  flex: 1;
  min-width: 0;
}

.col-fixed {
  flex-shrink: 0;
}

.col-fill {
  flex: 1;
  min-height: 0;
  overflow-y: auto;
  display: flex;
  flex-direction: column;
}

.bottom-card {
  flex-shrink: 0;
  overflow-y: auto;
}
</style>

<div class="d-flex justify-content-end mb-3">
  <form method="get" class="d-flex align-items-center" style="gap:8px;">
    <select name="year" class="form-select form-select-sm bg-white shadow-sm border-0" style="width:100px;">
      {% for y in year_range %}<option value="{{y}}" {% if y == selected_year %}selected{% endif %}>{{y}}年</option>{% endfor %}
    </select>
    <select name="month" class="form-select form-select-sm bg-white shadow-sm border-0" style="width:80px;">
      {% for m in range(1,13) %}<option value="{{m}}" {% if m == selected_month %}selected{% endif %}>{{m}}月</option>{% endfor %}
    </select>
    <button type="submit" class="btn btn-primary btn-sm px-3 shadow-sm">刷新面板</button>
  </form>
</div>

<div class="dashboard-row" id="dashboardRow">

  <!-- ========= 第一列 ========= -->
  <div class="dashboard-col" id="col1">
    <div class="col-fill mb-3">
      <div class="section-title">📋 班组当日在制实况</div>
      <div class="dashboard-table-wrap d-flex flex-column">
          <div style="position:relative; height:272px; width:100%; margin-bottom: 10px;">
            <canvas id="inprogressChart"></canvas>
          </div>
          <ul class="nav nav-tabs small px-2 border-bottom-0 flex-nowrap" style="overflow-x:auto;">
            {% for team, items in in_progress_for_template.items() %}
            <li class="nav-item"><a class="nav-link {% if loop.first %}active{% endif %} py-2 px-3 text-dark" data-bs-toggle="tab" href="#team-{{loop.index}}">{{team}}</a></li>
            {% endfor %}
          </ul>
          <div class="tab-content bg-white border border-top-0 rounded-bottom" id="inprogress-tabs-content" style="overflow-y:auto;">
            {% for team, items in in_progress_for_template.items() %}
            <div class="tab-pane fade {% if loop.first %}show active{% endif %}" id="team-{{loop.index}}">
              <table class="dashboard-table small auto-play-table paginate-table border-0">
                <thead><tr><th>序号</th><th>图号</th><th>产品</th><th>具体号型</th><th>批次</th><th>数量</th><th>开始时间</th></tr></thead>
                <tbody>{% for it in items %}<tr><td>{{it.serial_no}}</td><td>{{it.product_draw_no}}</td><td><div style="max-width:75px;overflow:hidden;text-overflow:ellipsis;" title="{{it.product_name}}">{{it.product_name}}</div></td><td>{{it.specific_model}}</td><td>{{it.batch_no}}</td><td>{{it.total_qty}}</td><td>{{fmt_date(it.start_time)}}</td></tr>{% endfor %}</tbody>
              </table>
            </div>
            {% endfor %}
          </div>
      </div>
    </div>

    <div class="bottom-card" id="bottom1">
      <div class="d-flex justify-content-between align-items-center mb-2">
        <div class="section-title mb-0">📦 {{selected_year}}年{{selected_month}}月应交付任务</div>
        <div style="width:45px;"></div>
      </div>
      <div class="dashboard-table-wrap">
          <table class="dashboard-table small paginate-table">
              <thead><tr><th>序号</th><th>图号</th><th>产品</th><th>批次</th><th>数量</th><th>产值</th><th>状态</th></tr></thead>
              <tbody>
                  {% for it in monthly_delivery_list %}
                  <tr><td>{{it.serial_no}}</td><td>{{it.product_draw_no}}</td><td><div style="max-width:70px;overflow:hidden;text-overflow:ellipsis;">{{it.product_name}}</div></td><td>{{it.batch_no}}</td><td>{{it.total_qty}}</td><td>{{it.output_value}}</td>
                  <td><span class="badge {{ 'bg-success' if it.status == '已完成' else 'bg-warning text-dark' }}">{{it.status}}</span></td></tr>
                  {% else %}
                  <tr><td colspan="7" class="text-muted py-4 text-center">当月无主计划要求交付任务</td></tr>
                  {% endfor %}
              </tbody>
          </table>
      </div>
    </div>
  </div>

  <!-- ========= 第二列 ========= -->
  <div class="dashboard-col" id="col2">
    <div class="col-fixed mb-4">
      <div class="section-title">📊 核心效能指标</div>
      <div class="row g-3">
        <div class="col-6"><div class="card-metric"><i class="metric-icon-bg">⏱️</i><div class="card-body p-3"><div><div class="metric-label">🟢 准时完成率</div><div class="metric-value">{{rate1}}%</div><small class="text-muted fw-bold">{{'%.2f'|format(on_time_output)}}万元</small></div></div></div></div>
        <div class="col-6"><div class="card-metric"><i class="metric-icon-bg">📈</i><div class="card-body p-3"><div><div class="metric-label">📊 任务完成率</div><div class="metric-value">{{rate2}}%</div><small class="text-muted fw-bold">{{'%.2f'|format(done_output)}}万元</small></div></div></div></div>
        <div class="col-6"><div class="card-metric"><i class="metric-icon-bg">⏳</i><div class="card-body p-3"><div><div class="metric-label">⏳ 延期滞留率</div><div class="metric-value">{{rate3}}%</div><small class="text-muted fw-bold">{{'%.2f'|format(delayed_done_output)}}万元</small></div></div></div></div>
        <div class="col-6"><div class="card-metric"><i class="metric-icon-bg text-danger">⚠️</i><div class="card-body p-3"><div><div class="metric-label text-danger">⚠️ 未完违约金</div><div class="metric-value text-danger">{{penalty_count}}项</div><small class="text-danger fw-bold">{{'%.2f'|format(penalty_amount)}}万元</small></div></div></div></div>
      </div>
    </div>

    <div class="col-fill" style="margin-bottom: calc(1rem + 2mm);">
      <div class="dashboard-table-wrap d-flex flex-column">
          <ul class="nav nav-tabs small px-2 border-bottom-0">
              <li class="nav-item"><a class="nav-link active py-2 px-3 text-danger fw-bold border-0 shadow-sm" data-bs-toggle="tab" href="#tab-overdue">🚨 超期预警 (7天内未入库)</a></li>
              <li class="nav-item"><a class="nav-link py-2 px-3 text-primary border-0" data-bs-toggle="tab" href="#tab-unstored">✅ 已交总检待入库追踪</a></li>
          </ul>
          <div class="tab-content bg-white border-top-0 rounded-bottom" style="overflow-y:auto;">
              <div class="tab-pane fade show active" id="tab-overdue">
                  <table class="dashboard-table small paginate-table border-0">
                    <thead><tr><th>序号</th><th>图号</th><th>产品</th><th>批次</th><th>交付期限</th><th>产值</th></tr></thead>
                    <tbody>
                      {% for t in warnings %}
                      <tr><td class="text-danger fw-bold">{{t.serial_no}}</td><td>{{t.product_draw_no or ''}}</td><td><div style="max-width:80px;overflow:hidden;text-overflow:ellipsis;">{{t.product_name}}</div></td><td>{{t.batch_no or ''}}</td><td class="text-danger fw-bold">{{fmt_date(t.plan_delivery_time)}}</td><td>{{t.output_value or ''}}</td></tr>
                      {% else %}
                      <tr><td colspan="6" class="text-muted py-4 text-center">暂无即将超期预警</td></tr>
                      {% endfor %}
                    </tbody>
                  </table>
              </div>
              <div class="tab-pane fade" id="tab-unstored">
                  <table class="dashboard-table small paginate-table border-0">
                    <thead><tr><th>序号</th><th>图号</th><th>产品</th><th>批次</th><th>交总检日期</th><th>产值</th></tr></thead>
                    <tbody>
                      {% for t in done_not_stored_tasks %}
                      <tr><td>{{t.serial_no}}</td><td>{{t.product_draw_no}}</td><td><div style="max-width:80px;overflow:hidden;text-overflow:ellipsis;">{{t.product_name}}</div></td><td>{{t.batch_no}}</td><td><span class="text-primary">{{fmt_date(t.final_check_time)}}</span></td><td>{{t.output_value or ''}}</td></tr>
                      {% else %}
                      <tr><td colspan="6" class="text-muted py-4 text-center">无相关任务</td></tr>
                      {% endfor %}
                    </tbody>
                  </table>
              </div>
          </div>
      </div>
    </div>

    <div class="bottom-card" id="bottom2">
      <div class="d-flex justify-content-between align-items-center mb-2">
        <div class="section-title mb-0">⚠️ 当月未完违约金任务列控</div>
        <a href="/export/penalty?year={{selected_year}}&month={{selected_month}}" class="btn btn-sm btn-outline-danger py-0 px-3">导出</a>
      </div>
      <div class="dashboard-table-wrap">
        <table class="dashboard-table small paginate-table">
          <thead><tr><th>序号</th><th>图号</th><th>产品</th><th>批次</th><th>数量</th><th>产值</th><th>操作</th></tr></thead>
          <tbody>
            {% for t in penalty_tasks %}
            <tr>
              <td>{{t.serial_no}}</td><td>{{t.product_draw_no or ''}}</td><td><div style="max-width:70px;overflow:hidden;text-overflow:ellipsis;">{{t.product_name}}</div></td>
              <td>{{t.batch_no or ''}}</td><td>{{t.total_qty}}</td>
              <td>{{t.output_value or ''}}</td>
              <td><button class="btn btn-sm btn-light border py-0 px-2" onclick="popDetail('{{t.id}}', '{{t.product_name}}')">追溯</button></td>
            </tr>
            {% else %}
            <tr><td colspan="7" class="text-muted py-4 text-center">当月无相关违约金任务</td></tr>
            {% endfor %}
          </tbody>
        </table>
      </div>
    </div>
  </div>

  <!-- ========= 第三列 ========= -->
  <div class="dashboard-col" id="col3">
    <div class="col-fixed mb-1">
      <div class="section-title">🥧 当月达成产值比重</div>
      <div class="dashboard-table-wrap"><div style="position:relative; height:176px; width:100%;"><canvas id="completionChart"></canvas></div></div>
    </div>

    <div class="col-fill mb-3">
      <div class="d-flex justify-content-between align-items-center mb-2">
        <div class="section-title mb-0">📋 问题台账</div>
        <a href="/export/quality?year={{selected_year}}&month={{selected_month}}" class="btn btn-sm btn-outline-dark py-0 px-3" target="_blank" download>导出</a>
      </div>
      <div class="dashboard-table-wrap h-100 d-flex flex-column">
        <table class="dashboard-table small paginate-table flex-fill">
          <thead><tr><th>序号</th><th>生产状态</th><th>图号</th><th>批次</th><th>数量</th><th>用时</th><th>定性</th><th>产值</th><th>详情</th></tr></thead>
          <tbody>
          {% for q in quality_list %}
          <tr><td>{{q.serial_no}}</td><td><span class="badge {{ 'bg-danger' if q.prod_status=='未开工' else 'bg-warning text-dark' if q.prod_status=='生产中' else 'bg-success' }}">{{q.prod_status}}</span></td><td>{{q.product_draw_no}}</td><td>{{q.batch_no}}</td><td>{{q.total_qty}}</td>
          <td class="text-danger">{{q.days}}天</td><td>{{q.nature}}</td><td>{{'%.2f'|format(q.output_value)}}</td>
          <td><button class="btn btn-sm btn-light border py-0 px-2" onclick="popDetail('{{q.id}}','{{q.product_name}}')">详情</button></td></tr>
          {% else %}
          <tr><td colspan="9" class="text-muted py-4 text-center">当月无瑕疵记录</td></tr>
          {% endfor %}
          </tbody>
        </table>
      </div>
    </div>

    <div class="bottom-card" id="bottom3">
      <div class="dashboard-table-wrap d-flex flex-column">
          <ul class="nav nav-tabs small px-2 border-bottom-0">
              <li class="nav-item"><a class="nav-link active py-2 px-3 text-dark fw-bold border-0 shadow-sm" data-bs-toggle="tab" href="#month-issue">🚨 技术/管理报错</a></li>
              <li class="nav-item"><a class="nav-link py-2 px-3 text-primary border-0" data-bs-toggle="tab" href="#month-shortage">📦 物料缺件报备</a></li>
          </ul>
          <div class="tab-content bg-white border-top-0 rounded-bottom" style="overflow-y:auto;">
              <div class="tab-pane fade show active h-100" id="month-issue">
                  <table class="dashboard-table small paginate-table border-0">
                    <thead><tr><th>序号</th><th>图号</th><th>产品</th><th>数量</th><th>定性</th><th>提出时间</th></tr></thead>
                    <tbody>
                      {% for issue in m_issues_data %}
                      <tr>
                        <td>{{ issue.serial_no }}</td><td>{{ issue.product_draw_no }}</td>
                        <td><div style="max-width:70px;overflow:hidden;text-overflow:ellipsis;">{{ issue.product_name }}</div></td>
                        <td>{{ issue.total_qty }}</td>
                        <td><span class="badge px-2 py-1" style="background-color:#d4edda; color:#333;">{{ '技术' if issue.issue_type == 'tech' else '管理' }}</span></td>
                        <td>{{ fmt_date(issue.raise_time) }}</td>
                      </tr>
                      {% endfor %}
                    </tbody>
                  </table>
              </div>
              <div class="tab-pane fade" id="month-shortage">
                  <table class="dashboard-table small paginate-table border-0">
                    <thead><tr><th>序号</th><th>图号</th><th>产品</th><th>数量</th><th>物料类型</th><th>报缺日期</th></tr></thead>
                    <tbody>
                      {% for sr in m_shortages_data %}
                      <tr>
                        <td>{{ sr.serial_no }}</td><td>{{ sr.product_draw_no }}</td>
                        <td><div style="max-width:70px;overflow:hidden;text-overflow:ellipsis;">{{ sr.product_name }}</div></td>
                        <td>{{ sr.total_qty }}</td>
                        <td><span class="badge bg-primary text-white px-2 py-1">{{ sr.shortage_type }}</span></td>
                        <td>{{ fmt_date(sr.report_time) }}</td>
                      </tr>
                      {% endfor %}
                    </tbody>
                  </table>
              </div>
          </div>
      </div>
    </div>
  </div>

</div>

<div class="modal fade" id="popModal" tabindex="-1">
  <div class="modal-dialog modal-xl">
    <div class="modal-content">
      <div class="modal-header py-3 bg-light border-0">
        <h5 class="modal-title fw-bold text-dark" id="popTitle">业务详情快照</h5>
        <button type="button" class="btn-close" data-bs-dismiss="modal"></button>
      </div>
      <div class="modal-body p-0 bg-white">
        <iframe id="popFrame" src="" style="width:100%; height:75vh; border:none; display:block;"></iframe>
      </div>
    </div>
  </div>
</div>

<script>
function equalizeBottomCards() {
  var cards = [
    document.getElementById('bottom1'),
    document.getElementById('bottom2'),
    document.getElementById('bottom3')
  ];
  cards.forEach(function(c) { if (c) c.style.height = 'auto'; });
  var maxH = 0;
  cards.forEach(function(c) {
    if (c) {
      var h = c.scrollHeight;
      if (h > maxH) maxH = h;
    }
  });
  cards.forEach(function(c) { if (c) c.style.height = maxH + 'px'; });
}

window.addEventListener('resize', function() { equalizeBottomCards(); });

function initCharts() {
  Chart.register(ChartDataLabels);
  var doughnutVals = [{{'%.2f'|format(on_time_output)}}, {{'%.2f'|format(not_on_time_output)}}];
  var patternGreen = createPattern('#4ADE80', 'dot');
  var patternRed = createPattern('#FB7185', 'stripe');
  new Chart(document.getElementById('completionChart').getContext('2d'), {
    type: 'doughnut',
    data: { 
      labels: ['准时/正常结案','未准时/延期'], 
      datasets: [{ 
          data: [{{on_time}}, {{total_monthly - on_time}}], 
          backgroundColor: [patternGreen, patternRed], 
          hoverBackgroundColor: ['#22C55E', '#F43F5E'],
          borderWidth: 2, borderColor: '#ffffff',
          hoverOffset: 4
      }] 
    },
    options: { 
      responsive: true, maintainAspectRatio: false, cutout: '45%',
      plugins: { 
        legend: { position: 'right', labels:{usePointStyle:true, padding:15, font:{weight:'bold', size: 12}} },
        datalabels: {
          color: '#1e293b', font: {weight: 'bold', size:12}, align:'center',
          backgroundColor: 'rgba(255,255,255,0.7)', borderRadius: 4, padding: 4,
          formatter: function(value, context) {
            var total = context.dataset.data.reduce((a,b) => a+b, 0);
            if (total === 0 || value === 0) return '';
            var pct = Math.round((value / total) * 100);
            return pct + '%\\n' + doughnutVals[context.dataIndex] + '万';
          }
        }
      } 
    }
  });

  let teamLabels = {{ in_progress_for_template.keys() | list | tojson }};
  let teamCounts = []; {% for team, items in in_progress_for_template.items() %} teamCounts.push({{items|length}}); {% endfor %}
  let vividColors = ['#60A5FA', '#A78BFA', '#F472B6', '#FBBF24', '#34D399', '#38BDF8', '#FB923C'];
  let patternsType = ['stripe', 'dot', 'grid', 'zigzag'];
  let mappedPatterns = teamCounts.map((_, i) => createPattern(vividColors[i % vividColors.length], patternsType[i % patternsType.length]));
  new Chart(document.getElementById('inprogressChart').getContext('2d'), {
    type: 'bar',
    data: { labels: teamLabels, datasets: [{ data: teamCounts, backgroundColor: mappedPatterns, borderRadius: 4, hoverBackgroundColor: vividColors }] },
    options: { 
        responsive: true, maintainAspectRatio: false,
        layout: { padding: { top: 25 } }, 
        scales: { 
            y: { 
                border: {display:false}, 
                grid:{color:'#f1f5f9'}, 
                ticks:{precision:0, autoSkip: true, maxTicksLimit: 8},
                beginAtZero: true
            },
            x: { border: {display:false}, grid:{display:false} }
        },
        plugins: { 
            legend: { display: false }, 
            datalinks: { align: 'top', anchor: 'end', color: '#475569', font:{weight:'bold'} } 
        } 
    }
  });
}

function createPattern(color, type) {
    let canvas = document.createElement('canvas');
    canvas.width = 16; canvas.height = 16;
    let ctx = canvas.getContext('2d');
    ctx.fillStyle = color;
    ctx.fillRect(0, 0, 16, 16);
    ctx.fillStyle = 'rgba(255,255,255,0.25)';
    ctx.strokeStyle = 'rgba(255,255,255,0.3)';
    ctx.lineWidth = 2;
    if (type === 'stripe') {
        ctx.beginPath(); ctx.moveTo(0, 16); ctx.lineTo(16, 0); ctx.moveTo(-4, 4); ctx.lineTo(4, -4); ctx.moveTo(12, 20); ctx.lineTo(20, 12); ctx.stroke();
    } else if (type === 'dot') {
        ctx.beginPath(); ctx.arc(8, 8, 3, 0, Math.PI*2); ctx.fill();
    } else if (type === 'grid') {
        ctx.beginPath(); ctx.moveTo(0, 8); ctx.lineTo(16, 8); ctx.moveTo(8, 0); ctx.lineTo(8, 16); ctx.stroke();
    } else if (type === 'zigzag') {
        ctx.beginPath(); ctx.moveTo(0, 4); ctx.lineTo(8, 12); ctx.lineTo(16, 4); ctx.stroke();
    }
    return ctx.createPattern(canvas, 'repeat');
}

document.addEventListener('DOMContentLoaded', function () {
  initCharts();
  setTimeout(function() {
    equalizeBottomCards();
    setupPagination();
  }, 300);
  setInterval(function() {
    var activePane = document.querySelector('#inprogress-tabs-content .tab-pane.active');
    if (!activePane) return;
    var table = activePane.querySelector('.auto-play-table');
    if (!table) return;
    var nav = table.nextElementSibling;
    if (!nav || !nav.classList.contains('pagination-nav')) return;
    var nextBtn = nav.querySelector('.next-btn');
    var prevBtn = nav.querySelector('.prev-btn');
    if (!nextBtn || !prevBtn) return;
    if (nextBtn.disabled) {
      while (prevBtn && !prevBtn.disabled) prevBtn.click();
    } else {
      nextBtn.click();
    }
  }, 8000);
});

document.querySelectorAll('a[data-bs-toggle="tab"]').forEach(tab => {
  tab.addEventListener('shown.bs.tab', function () {
    setTimeout(setupPagination, 100);
    equalizeBottomCards();
  });
});

function setupPagination() {
  document.querySelectorAll('.paginate-table').forEach(table => {
    let tbody = table.querySelector('tbody');
    if (!tbody) return;
    let rows = Array.from(tbody.querySelectorAll('tr')).filter(r => !r.classList.contains('padding-row'));
    if (rows.length === 1 && rows[0].cells.length === 1 && rows[0].cells[0].colSpan > 1) return;
    let rowsPerPage = 10;
    let totalPages = Math.ceil(rows.length / rowsPerPage) || 1;
    let nav = table.parentNode.querySelector('.pagination-nav');
    if (!nav) {
      nav = document.createElement('div');
      nav.className = 'd-flex justify-content-between align-items-center px-3 py-2 pagination-nav bg-light mt-0 border-top';
      nav.innerHTML = '<span class="small text-muted fw-bold">共 ' + rows.length + ' 项</span>' +
        '<div class="d-flex align-items-center">' +
        '<button class="btn-circle-page prev-btn">&lt;</button>' +
        '<span class="page-info small fw-bold text-primary mx-2">1 / ' + totalPages + '</span>' +
        '<button class="btn-circle-page next-btn">&gt;</button></div>';
      table.parentNode.insertBefore(nav, table.nextSibling);
    }
    function showPage(page) {
      page = Math.max(1, Math.min(page, totalPages));
      tbody.querySelectorAll('.padding-row').forEach(r => r.remove());
      let start = (page - 1) * rowsPerPage;
      let end = start + rowsPerPage;
      let visibleCount = 0;
      rows.forEach((row, i) => {
        row.style.display = (i >= start && i < end) ? '' : 'none';
        if (i >= start && i < end) visibleCount++;
      });
      for (let i = visibleCount; i < rowsPerPage; i++) {
        let ptr = document.createElement('tr');
        ptr.className = 'padding-row';
        ptr.innerHTML = '<td colspan="' + rows[0].cells.length + '" style="border:none!important; color:transparent;">-</td>';
        tbody.appendChild(ptr);
      }
      nav.querySelector('.page-info').innerText = page + ' / ' + totalPages;
      nav.querySelector('.prev-btn').disabled = (page === 1);
      nav.querySelector('.next-btn').disabled = (page === totalPages);
    }
    nav.querySelector('.prev-btn').onclick = () => showPage(parseInt(nav.querySelector('.page-info').innerText) - 1);
    nav.querySelector('.next-btn').onclick = () => showPage(parseInt(nav.querySelector('.page-info').innerText) + 1);
    showPage(1);
  });
}

function alignDashboard() {
    var cols = document.querySelectorAll('#col1, #col2, #col3');
    var maxBottom = 0;
    cols.forEach(function(col) {
        var bottom = col.querySelector('.bottom-card');
        if (bottom) bottom.style.minHeight = '';
    });
    cols.forEach(function(col) {
        var bottom = col.querySelector('.bottom-card');
        if (bottom) {
            var h = bottom.getBoundingClientRect().height;
            if (h > maxBottom) maxBottom = h;
        }
    });
    cols.forEach(function(col) {
        var bottom = col.querySelector('.bottom-card');
        if (bottom) bottom.style.minHeight = maxBottom + 'px';
    });
}
window.addEventListener('load', function() { setTimeout(alignDashboard, 200); });
window.addEventListener('resize', alignDashboard);

function popDetail(id, name) {
    document.getElementById('popTitle').innerText = '📑 业务详情快照：' + name;
    document.getElementById('popFrame').src = '/task/' + id + '/details';
    new bootstrap.Modal(document.getElementById('popModal')).show();
}
</script>
{% endblock %}
''')

DATACENTER_HTML = BASE_HTML.replace('{% block content %}{% endblock %}','''
{% block content %}
<style>
{% set mc = ['#3b82f6','#8b5cf6','#f59e0b','#ef4444','#10b981','#06b6d4','#84cc16','#f97316','#6366f1','#ec4899','#14b8a6','#eab308','#a855f7','#dc2626','#ea580c','#c026d3','#0891b2','#f43f5e'] %}
{% set mi = ['📦','📚','⚠️','🔴','📋','📈','✅','🏭','📊','🏆','🔧','⏳','📌','🔥','💰','📋','📅','🔍'] %}
.card-metric-tab {
    border: 1px solid #e2e8f0; border-radius: 8px; background: #fff;
    box-shadow: 0 2px 4px rgba(0,0,0,0.03); position: relative; overflow: hidden;
    padding: 10px 10px; width: 100%; text-align: left; transition: all 0.2s;
    color: #1e293b; cursor: grab;
    border-left: 4px solid #94a3b8;
    min-height: 86px; display: flex; flex-direction: column; justify-content: center;
}
.nav-link.dragging { opacity: 0.5; }
.nav-link.drag-over .card-metric-tab { border: 2px dashed #3b82f6 !important; background-color: #eff6ff !important; }
.datacenter-pills .nav-link {
    padding: 0; background: none; border: none;
    margin-bottom: 4px;
}
.datacenter-pills .nav-link.active .card-metric-tab {
    background-color: #d1fae5 !important;
    color: #1e293b !important;
    box-shadow: 0 4px 12px rgba(16, 185, 129, 0.2);
}
.metric-icon-bg-tab {
    position: absolute; right: -5px; top: 50%; transform: translateY(-50%);
    font-size: 3rem; opacity: 0.08; z-index: 1; user-select: none;
}
.dc-table td { padding: 2px 6px !important; font-size: 15px; line-height: 1.7; }
.dc-table th { font-size: 16px !important; font-weight: bold; }
.dc-tfoot { position: sticky; bottom: 0; z-index: 5; background: #f8f9fa !important; }
.dc-tfoot td { border-top: 2px solid #dee2e6 !important; }
.dc-table tbody tr:hover td { background:#f8fafc; }
#v-pills-tabContent { border-color:#e2e8f0!important; border-radius:8px!important; }
@media (max-width: 1280px){
  .card-metric-tab{min-height:76px; padding:8px;}
  .dc-table td{font-size:14px;}
  .dc-table th{font-size:14px!important;}
}
@media (max-width: 992px){
  .datacenter-pills .nav-link{min-width:150px;}
}
</style>

<div class="d-flex justify-content-between align-items-center mb-3">
  <h4 class="fw-bold text-dark m-0">📊 核心数据中心</h4>
  <form method="get" class="d-flex align-items-center" style="gap:8px;">
    <select name="year" id="yearFilter" class="form-select form-select-sm bg-white shadow-sm border-0" style="width:100px;">
      {% for y in year_range %}<option value="{{y}}" {% if y == selected_year %}selected{% endif %}>{{y}}年</option>{% endfor %}
    </select>
    <select name="month" id="monthFilter" class="form-select form-select-sm bg-white shadow-sm border-0" style="width:80px;">
      {% for m in range(1,13) %}<option value="{{m}}" {% if m == selected_month %}selected{% endif %}>{{m}}月</option>{% endfor %}
    </select>
    <button type="submit" class="btn btn-success btn-sm px-3 shadow-sm">刷新数据</button>
  </form>
</div>

<div class="d-flex" style="gap: 8px;">
  <div class="flex-shrink-0" style="width: 14%; min-width: 160px;">
    <div class="nav flex-column nav-pills datacenter-pills" id="v-pills-tab-left" role="tablist">
      {% for m in metrics[:9] %}
      {% set ci = loop.index0 %}
      <a class="nav-link w-100 border-0 p-0 {% if loop.first %}active{% endif %}" data-bs-toggle="pill" href="#{{m.id}}" role="tab" draggable="true">
         <div class="card-metric-tab" style="border-left-color:{{mc[ci]}};">
             <i class="metric-icon-bg-tab">{{mi[ci]}}</i>
             <div style="position:relative; z-index:2;">
                 <div class="fw-bold text-dark" style="font-size: 14px; white-space: nowrap; overflow: hidden; text-overflow: ellipsis;" title="{{m.title}}">{{mi[ci]}} {{m.title}}</div>
                 <div class="small mt-1 text-dark" style="font-size: 12px; opacity: 0.9;">
                     <div class="d-flex justify-content-between"><span>数: <strong>{{m.count}}</strong></span> <span>产: <strong>{{'%.1f'|format(m.val)}}</strong>W</span></div>
                     {% if m.rate is not none %}<div class="mt-1">达成比率: <strong style="color:{{mc[ci]}}">{{m.rate}}%</strong></div>{% endif %}
                 </div>
             </div>
         </div>
      </a>
      {% endfor %}
    </div>
  </div>
  
  <div class="flex-grow-1" style="min-width: 0;">
    <div class="tab-content bg-white p-3 border rounded shadow-sm" id="v-pills-tabContent" style="min-height: 600px;">
      {% for m in metrics %}
      <div class="tab-pane fade {% if loop.first %}show active{% endif %}" id="{{m.id}}">
          <div class="d-flex justify-content-between align-items-center mb-3">
             <h5 class="fw-bold text-dark m-0">{{m.title}} 明细列表</h5>
             <div class="d-flex gap-2 align-items-center">
                 {% if m.id in ['m16','m17'] %}
                 <select class="form-select form-select-sm dc-filter" style="width:110px;" data-mid="{{m.id}}" data-col="品种" onchange="dcFilterTable(this)">
                   <option value="">全部品种</option>
                   <option value="服装">服装</option><option value="头盔">头盔</option><option value="面罩">面罩</option><option value="船囊">船囊</option><option value="救生衣">救生衣</option>
                 </select>
                 {% endif %}
                 {% if m.id == 'm16' %}
                 <select class="form-select form-select-sm dc-filter" style="width:110px;" data-mid="{{m.id}}" data-col="完成情况" onchange="dcFilterTable(this)">
                   <option value="">全部完成情况</option>
                   <option value="已完成">已完成</option><option value="未完成">未完成</option>
                 </select>
                 {% endif %}
                 <button class="btn btn-sm btn-outline-dark px-3 rounded-pill" data-export-target="table-{{m.id}}">⬇ 导出当前列表</button>
             </div>
         </div>
           <div style="overflow-x: auto;">
              <table class="table table-bordered table-sm table-hover text-center dc-table" id="table-{{m.id}}" data-count="{{m.count}}" data-val="{{'%.2f'|format(m.val)}}">
                   <thead class="bg-light">
                       <tr>
                           {% for col_name, f in metric_cols[m.id] %}
                           <th style="white-space:nowrap;{% if f == '__hover_text' %}display:none;{% endif %}">{{ col_name }}</th>
                           {% endfor %}
                       </tr>
                   </thead>
                   <tbody>
                       {% for row in m.rows %}
                       <tr>
                           {% for col_name, f in metric_cols[m.id] %}
                           <td style="{% if f == '__hover_text' %}display:none;{% endif %}">
                           {% if f == '__hover_text' %}{# hidden #}
                           {% elif f == '__link' %}
                               <a href="/task/{{ row[f] }}/details" target="_blank" class="btn btn-sm btn-outline-secondary py-0" style="font-size:inherit;">详情</a>
                           {% elif f == 'production_status' and m.id in ['m17','m18'] %}
                               <span class="badge {{ 'bg-danger' if row[f]=='未开工' else 'bg-warning text-dark' if row[f]=='生产中' else 'bg-success' }}" title="{{ row['__hover_text'] }}" style="cursor:help;">{{ row[f] }}</span>
                           {% elif f == 'product_name' %}
                               <div style="max-width:120px;overflow:hidden;text-overflow:ellipsis;" title="{{ row[f] }}">{{ row[f] }}</div>
                           {% else %}
                               {{ row[f] }}
                           {% endif %}
                           </td>
                           {% endfor %}
                       </tr>
                       {% else %}
                       <tr><td colspan="{{ metric_cols[m.id]|length }}" class="text-muted py-5">当前条件项下暂无匹配数据</td></tr>
                       {% endfor %}
                   </tbody>
                    <tfoot class="dc-tfoot" style="display:none;">
                        <tr class="fw-bold bg-light">
                            <td colspan="{{ metric_cols[m.id]|length }}" class="text-end pe-3" style="font-size:13px;">
                                产值: <span class="text-primary dc-total-val">{{'%.2f'|format(m.val)}}</span> W
                           </td>
                       </tr>
                   </tfoot>
             </table>
          </div>
          <div class="dc-pagination-nav d-flex justify-content-between align-items-center px-3 py-2 bg-light border-top rounded-bottom" id="pnav-{{m.id}}" style="display:none;">
              <span class="small text-muted">共 <span class="fw-bold dc-pg-count">0</span> 项</span>
              <div class="d-flex align-items-center gap-2">
                  <button class="btn btn-sm btn-outline-secondary py-0 dc-prev" disabled>&lt;</button>
                  <span class="small fw-bold dc-page-info">1 / 1</span>
                  <button class="btn btn-sm btn-outline-secondary py-0 dc-next" disabled>&gt;</button>
              </div>
          </div>
      </div>
      {% endfor %}
    </div>
  </div>

  <div class="flex-shrink-0" style="width: 14%; min-width: 160px;">
    <div class="nav flex-column nav-pills datacenter-pills" id="v-pills-tab-right" role="tablist" style="gap: 2px;">
      {% for m in metrics[9:] %}
      {% set ci = 9 + loop.index0 %}
      <a class="nav-link w-100 border-0 p-0" data-bs-toggle="pill" href="#{{m.id}}" role="tab" draggable="true">
         <div class="card-metric-tab" style="border-left-color:{{mc[ci]}};">
             <i class="metric-icon-bg-tab">{{mi[ci]}}</i>
             <div style="position:relative; z-index:2;">
                 <div class="fw-bold text-dark" style="font-size: 14px; white-space: nowrap; overflow: hidden; text-overflow: ellipsis;" title="{{m.title}}">{{mi[ci]}} {{m.title}}</div>
                 <div class="small mt-1 text-dark" style="font-size: 12px; opacity: 0.9;">
                     <div class="d-flex justify-content-between"><span>数: <strong>{{m.count}}</strong></span> <span>产: <strong>{{'%.1f'|format(m.val)}}</strong>W</span></div>
                     {% if m.rate is not none %}<div class="mt-1">达成比率: <strong style="color:{{mc[ci]}}">{{m.rate}}%</strong></div>{% endif %}
                 </div>
             </div>
         </div>
      </a>
      {% endfor %}
    </div>
  </div>
</div>

<script>
var DC_PAGE_SIZE = 20;
var DC_PAGE_STATE = {};

function getTableDataRows(table) {
    var tbody = table.querySelector('tbody');
    if (!tbody) return [];
    var rows = [];
    for (var i = 0; i < tbody.rows.length; i++) {
        var r = tbody.rows[i];
        if (r.classList.contains('dc-ph')) continue;
        if (r.cells.length === 1 && r.cells[0].colSpan > 1) continue;
        rows.push(r);
    }
    return rows;
}

function dcUpdate(mid) {
    var table = document.getElementById('table-' + mid);
    var nav = document.getElementById('pnav-' + mid);
    if (!table || !nav) return;

    var tbody = table.querySelector('tbody');
    var tfoot = table.querySelector('tfoot');
    var rows = getTableDataRows(table);

    if (rows.length === 0) {
        nav.style.display = 'none';
        if (tfoot) tfoot.style.display = 'none';
        return;
    }

    var totalPages = Math.ceil(rows.length / DC_PAGE_SIZE) || 1;
    var cur = DC_PAGE_STATE[mid] || 1;
    if (cur > totalPages) cur = 1;

    function go(p) {
        if (p < 1) p = 1;
        if (p > totalPages) p = totalPages;
        var s = (p - 1) * DC_PAGE_SIZE, e = s + DC_PAGE_SIZE, shown = 0;
        tbody.querySelectorAll('.dc-ph').forEach(function(ph) { ph.remove(); });
        for (var i = 0; i < rows.length; i++) {
            if (rows[i].classList.contains('dc-filter-hidden')) {
                rows[i].style.display = 'none';
                continue;
            }
            var display = (i >= s && i < e) ? '' : 'none';
            rows[i].style.display = display;
            if (display === '') shown++;
        }
        for (var k = 0; k < DC_PAGE_SIZE - shown; k++) {
            var ph = document.createElement('tr');
            ph.className = 'dc-ph';
            var td = document.createElement('td');
            td.colSpan = rows[0].cells.length;
            td.style.cssText = 'border:none !important; color:transparent;';
            td.innerText = '-';
            ph.appendChild(td);
            tbody.appendChild(ph);
        }
        nav.querySelector('.dc-page-info').textContent = p + ' / ' + totalPages;
        nav.querySelector('.dc-prev').disabled = (p <= 1);
        nav.querySelector('.dc-next').disabled = (p >= totalPages);
        DC_PAGE_STATE[mid] = p;
    }

    nav.style.display = 'flex';
    if (tfoot) tfoot.style.display = '';
    nav.querySelector('.dc-pg-count').textContent = rows.length;
    nav.querySelector('.dc-prev').onclick = function() { go(DC_PAGE_STATE[mid] - 1); };
    nav.querySelector('.dc-next').onclick = function() { go(DC_PAGE_STATE[mid] + 1); };
    go(cur);
    dcAddHeaderFilter(table);
}

function dcFilterTable(sel) {
    var mid = sel.dataset.mid;
    var colName = sel.dataset.col;
    var val = sel.value;
    var table = document.getElementById('table-' + mid);
    if (!table) return;
    var headers = table.querySelectorAll('thead th');
    var colIdx = -1;
    for (var i = 0; i < headers.length; i++) {
        if ((headers[i].getAttribute('data-col') || headers[i].innerText).trim() === colName) { colIdx = i; break; }
    }
    if (colIdx < 0) return;
    var tbody = table.querySelector('tbody');
    tbody.querySelectorAll('.dc-ph').forEach(function(ph) { ph.remove(); });
    var allFilters = table._dcFilters || {};
    allFilters[colName] = val;
    table._dcFilters = allFilters;
    var rows = getTableDataRows(table);
    rows.forEach(function(r) {
        var show = true;
        for (var c in allFilters) {
            if (!allFilters[c]) continue;
            var ci = -1;
            for (var j = 0; j < headers.length; j++) {
                if ((headers[j].getAttribute('data-col') || headers[j].innerText).trim() === c) { ci = j; break; }
            }
            if (ci >= 0 && r.cells[ci]) {
                if (r.cells[ci].innerText.trim() !== allFilters[c]) { show = false; break; }
            }
        }
        if (show) { r.classList.remove('dc-filter-hidden'); r.style.display = ''; }
        else { r.classList.add('dc-filter-hidden'); r.style.display = 'none'; }
    });
    DC_PAGE_STATE[mid] = 1;
    dcUpdate(mid);
    dcAddHeaderFilter(table);
}

function dcAddHeaderFilter(table) {
    if (table._headerFilterDone) return;
    table._headerFilterDone = true;
    var headers = table.querySelectorAll('thead th');
    headers.forEach(function(th, ci) {
        var colName = th.innerText.trim();
        th.setAttribute('data-col', colName);
        var icon = document.createElement('span');
        icon.innerHTML = '🔽';
        icon.style.cssText = 'cursor:pointer;font-size:12px;margin-left:3px;';
        icon.title = '筛选';
        icon.onclick = function(e) {
            e.stopPropagation();
            var existing = document.getElementById('dc-hf-menu');
            if (existing) existing.remove();
            var menu = document.createElement('div');
            menu.id = 'dc-hf-menu';
            menu.style.cssText = 'position:absolute;background:#fff;border:1px solid #ccc;padding:8px;z-index:9999;max-height:300px;overflow-y:auto;box-shadow:0 4px 8px rgba(0,0,0,0.1);min-width:120px;';
            menu.style.left = e.pageX + 'px';
            menu.style.top = e.pageY + 'px';
            var vals = new Set();
            var tbody = table.querySelector('tbody');
            for (var i = 0; i < tbody.rows.length; i++) {
                var r = tbody.rows[i];
                if (r.classList.contains('dc-ph')) continue;
                if (r.cells[ci]) vals.add(r.cells[ci].innerText.trim());
            }
            var sorted = Array.from(vals).sort();
            var html = '<div><label><input type="checkbox" checked class="dc-hf-all"> <b>全选</b></label></div><hr style="margin:4px 0;">';
            sorted.forEach(function(v) {
                var displayVal = v || '(空白)';
                html += '<div><label><input type="checkbox" class="dc-hf-chk" value="' + v.replace(/"/g, '&quot;') + '" checked> ' + displayVal + '</label></div>';
            });
            html += '<hr style="margin:4px 0;"><button class="btn btn-sm btn-primary w-100 dc-hf-ok">确定</button>';
            menu.innerHTML = html;
            document.body.appendChild(menu);
            menu.querySelector('.dc-hf-all').onchange = function() {
                var checked = this.checked;
                menu.querySelectorAll('.dc-hf-chk').forEach(function(cb) { cb.checked = checked; });
            };
            menu.querySelector('.dc-hf-ok').onclick = function() {
                var selected = new Set();
                menu.querySelectorAll('.dc-hf-chk:checked').forEach(function(cb) { selected.add(cb.value); });
                var trows = getTableDataRows(table);
                trows.forEach(function(r) {
                    if (r.cells[ci]) {
                        var show = selected.has(r.cells[ci].innerText.trim());
                        if (show) { r.classList.remove('dc-filter-hidden'); r.style.display = ''; }
                        else { r.classList.add('dc-filter-hidden'); r.style.display = 'none'; }
                    }
                });
                menu.remove();
                var mid = table.id.replace('table-', '');
                DC_PAGE_STATE[mid] = 1;
                dcUpdate(mid);
            };
            setTimeout(function() {
                var closeMenu = function(event) {
                    if (!menu.contains(event.target)) {
                        menu.remove();
                        document.removeEventListener('click', closeMenu);
                    }
                };
                document.addEventListener('click', closeMenu);
            }, 10);
        };
        th.appendChild(icon);
    });
}

function exportTableToCSV(tableId) {
    var mid = tableId.replace('table-', '');
    var year = document.getElementById('yearFilter').value;
    var month = document.getElementById('monthFilter').value;
    window.location.href = '/datacenter/export/' + mid + '?year=' + year + '&month=' + month;
}

var dragMid = null, dragEl = null;
function onDragStartDc(e) {
    var el = e.target.closest('[data-bs-toggle="pill"][draggable="true"]');
    if (!el) return;
    dragEl = el;
    dragMid = el.getAttribute('href').replace('#', '');
    e.dataTransfer.effectAllowed = 'move';
    e.dataTransfer.setData('text/plain', dragMid);
    el.classList.add('dragging');
}
function onDragEndDc(e) {
    if (dragEl) {
        dragEl.classList.remove('dragging');
        var activeTab = document.querySelector('.datacenter-pills .nav-link.active');
        if (activeTab) new bootstrap.Tab(activeTab).show();
    }
    document.querySelectorAll('.datacenter-pills .nav-link.drag-over').forEach(function(el) { el.classList.remove('drag-over'); });
    dragEl = null; dragMid = null;
}
function onDragOverDc(e) {
    e.preventDefault();
    e.dataTransfer.dropEffect = 'move';
    var el = e.target.closest('[data-bs-toggle="pill"][draggable="true"]');
    if (!el || !dragEl || el === dragEl) return;
    document.querySelectorAll('.datacenter-pills .nav-link.drag-over').forEach(function(x) { x.classList.remove('drag-over'); });
    el.classList.add('drag-over');
}
function onDropDc(e) {
    e.preventDefault();
    var targetEl = e.target.closest('[data-bs-toggle="pill"][draggable="true"]');
    if (!targetEl || !dragEl || targetEl === dragEl) return;
    var parent = targetEl.parentNode;
    var rect = targetEl.getBoundingClientRect();
    var midY = rect.top + rect.height / 2;
    if (e.clientY < midY) parent.insertBefore(dragEl, targetEl);
    else parent.insertBefore(dragEl, targetEl.nextSibling);
    var activeHref = dragEl.getAttribute('href');
    if (activeHref) {
        document.querySelector('.datacenter-pills .nav-link[href="' + activeHref + '"]').classList.add('active');
        new bootstrap.Tab(document.querySelector(activeHref)).show();
    }
    onDragEndDc(e);
}
function dcInitDrag() {
    document.querySelectorAll('.datacenter-pills').forEach(function(nav) {
        nav.addEventListener('dragstart', onDragStartDc);
        nav.addEventListener('dragend', onDragEndDc);
        nav.addEventListener('dragover', onDragOverDc);
        nav.addEventListener('drop', onDropDc);
    });
}

document.addEventListener('DOMContentLoaded', function() {
    dcInitDrag();
    setTimeout(function() {
        document.querySelectorAll('.dc-table').forEach(function(t) {
            dcUpdate(t.id.replace('table-', ''));
        });
    }, 350);
    document.addEventListener('click', function(e) {
        var btn = e.target.closest('[data-export-target]');
        if (btn) {
            e.preventDefault();
            var tableId = btn.getAttribute('data-export-target');
            exportTableToCSV(tableId);
        }
    });
});
</script>
{% endblock %}
''')

TEAM_KANBAN_HTML = BASE_HTML.replace('{% block content %}{% endblock %}','''
{% block content %}
<style>
.stats-bar { display: flex; gap: 12px; margin-bottom: 16px; }
.stat-card { display: flex; align-items: center; gap: 8px; background: #fff; border-radius: 10px; padding: 10px 16px; box-shadow: 0 1px 3px rgba(0,0,0,0.05); }
.stat-icon { font-size: 20px; }
.stat-num { font-size: 22px; font-weight: 800; line-height: 1.2; }
.stat-label { font-size: 11px; color: #64748b; }
.stat-card.stat-todo .stat-num { color: #64748b; }
.stat-card.stat-doing .stat-num { color: #d97706; }
.stat-card.stat-done .stat-num { color: #059669; }
.stat-card.stat-overdue .stat-num { color: #ef4444; }

.toolbar { display: flex; justify-content: space-between; align-items: center; margin-bottom: 12px; gap: 10px; flex-wrap:wrap; }
.toolbar .search-box { position: relative; flex: 1; max-width: 360px; }
.toolbar .search-box input { width: 100%; padding: 8px 12px 8px 36px; border: 1px solid #e2e8f0; border-radius: 8px; font-size: 13px; outline: none; }
.toolbar .search-box .search-icon { position: absolute; left: 12px; top: 50%; transform: translateY(-50%); opacity: 0.4; font-size: 15px; }

.team-tabs { display: flex; gap: 6px; flex-wrap: wrap; }

.kanban-board { display: flex; gap: 12px; padding-bottom: 16px; align-items:stretch; }
.kanban-board.single-view { overflow-x: auto; }
.kanban-side-panel { flex: 1; min-width: 0; display: flex; flex-direction: column; gap: 10px; }
.task-chart-box { background: #fff; border: 1px solid #e2e8f0; border-radius: 10px; padding: 10px; }
.task-stat-grid { display: grid; grid-template-columns: repeat(4, 1fr); gap: 6px; margin-top: 8px; }
.task-stat-tile { background: #f8fafc; border: 1px solid #e2e8f0; border-radius: 6px; padding: 6px 4px; text-align: center; }
.task-stat-num { font-size: 16px; font-weight: 800; color: #1e293b; }
.task-stat-label { font-size: 11px; color: #64748b; margin-top: 2px; }
.branch-notice { background: #fff; border: 1px solid #e2e8f0; border-left: 4px solid #2563eb; border-radius: 8px; overflow: hidden; flex: 1; display: flex; flex-direction: column; }
.branch-notice-head { padding: 8px 10px; background: #f8fafc; border-bottom: 1px solid #e2e8f0; display: flex; justify-content: space-between; align-items: center; }
.branch-notice-title { font-size: 14px; font-weight: 800; color: #1e293b; }
.branch-notice-list { padding: 8px; overflow-y: auto; max-height: 200px; }
.branch-notice-item { padding-bottom: 8px; margin-bottom: 8px; border-bottom: 1px dashed #e2e8f0; }
.branch-notice-item:last-child { border-bottom: none; margin-bottom: 0; padding-bottom: 0; }
.branch-notice-name { font-weight: 700; font-size: 13px; color: #0f172a; }
.branch-notice-meta { color: #64748b; font-size: 11px; margin-bottom: 4px; }
.branch-notice-content { white-space: pre-wrap; color: #334155; font-size: 12px; line-height: 1.5; }
.branch-notice-empty { color: #94a3b8; font-size: 12px; }
.kanban-column { flex: 1; min-width: 260px; background: #f8fafc; border-radius: 8px; display: flex; flex-direction: column; max-height: calc(100vh - 170px); border:1px solid #e2e8f0; }
.kanban-column.todo, .kanban-column.doing { flex: 1; }
.kanban-column-header { padding: 12px 16px 10px; font-weight: 700; font-size: 14px; border-radius: 12px 12px 0 0; display: flex; justify-content: space-between; align-items: center; }
.kanban-column-body { padding: 8px 10px; overflow-y: auto; flex: 1; }

.kanban-card { background: #fff; border-radius: 8px; padding: 10px 12px; margin-bottom: 8px; box-shadow: 0 1px 3px rgba(0,0,0,0.05); border-left: 4px solid #94a3b8; cursor: pointer; transition: all 0.15s; }
.kanban-card:hover { box-shadow: 0 2px 8px rgba(0,0,0,0.1); }
.kanban-card.overdue { border-left-color: #ef4444; background: #fef2f2; }
.kanban-card .card-header-row { display: flex; justify-content: space-between; align-items: center; margin-bottom: 4px; }
.kanban-card .card-serial { font-weight: 700; font-size: 15px; color: #1e293b; }
.kanban-card .card-badges { display: flex; gap: 4px; flex-wrap: wrap; }
.kanban-card .card-name { font-weight: 600; font-size: 14px; color: #1e293b; margin-bottom: 4px; }
.kanban-card .card-meta { font-size: 13px; color: #64748b; line-height: 1.2; display: flex; flex-wrap: wrap; gap: 0 10px; }
.overdue-tag { background: #fecaca; color: #991b1b; padding: 1px 6px; border-radius: 4px; font-size: 10px; font-weight: 600; }
.duration-tag { background: #e0e7ff; color: #3730a3; padding: 1px 6px; border-radius: 4px; font-size: 10px; font-weight: 600; }
.done-tag { background: #d1fae5; color: #065f46; padding: 1px 6px; border-radius: 4px; font-size: 10px; font-weight: 600; }
.kanban-card.hidden-by-search { display: none !important; }

.kanban-column-body::-webkit-scrollbar { width: 4px; }
.kanban-column-body::-webkit-scrollbar-thumb { background: #cbd5e1; border-radius: 4px; }

.kanban-column.todo .kanban-column-header { background: #e2e8f0; color: #475569; }
.kanban-column.doing .kanban-column-header { background: #fde68a; color: #92400e; }
.kanban-column.done .kanban-column-header { background: #a7f3d0; color: #065f46; }
.kanban-column.skipped .kanban-column-header { background: #e2e8f0; color: #64748b; }
.kanban-column.todo .kanban-card:not(.overdue) { border-left-color: #94a3b8; }
.kanban-column.doing .kanban-card:not(.overdue) { border-left-color: #f59e0b; }
.kanban-column.done .kanban-card:not(.overdue) { border-left-color: #10b981; }
.kanban-column.skipped .kanban-card:not(.overdue) { border-left-color: #94a3b8; }

.no-match-msg { display: none; text-align: center; padding: 20px; color: #94a3b8; font-size: 13px; }

.member-tag { display: inline-flex; align-items: center; gap: 4px; padding: 3px 10px; background: #eff6ff; border: 1px solid #bfdbfe; border-radius: 12px; font-size: 13px; color: #1e40af; font-weight: 500; }
.member-tag .member-remove { cursor: pointer; color: #94a3b8; font-size: 16px; line-height: 1; margin-left: 4px; font-weight: 700; }
.member-tag .member-remove:hover { color: #ef4444; }
.members-dropdown { position: relative; }
.members-dropdown-toggle { cursor: pointer; padding: 5px 10px; background: #f8fafc; border: 1px solid #e2e8f0; border-radius: 6px; font-size: 13px; font-weight: 700; color: #374151; display: flex; align-items: center; gap: 4px; }
.members-dropdown-toggle:hover { border-color: #3b82f6; }
.members-dropdown-menu { display: none; position: absolute; top: 100%; left: 0; min-width: 100%; z-index: 1000; background: #fff; border: 1px solid #e2e8f0; border-radius: 6px; max-height: 240px; overflow-y: auto; box-shadow: 0 4px 12px rgba(0,0,0,0.1); padding: 4px; }
.member-tag-row { display: flex; justify-content: space-between; align-items: center; padding: 3px 8px; border-radius: 4px; margin: 1px 0; background: #eff6ff; border: 1px solid #bfdbfe; font-size: 13px; color: #1e40af; }
.member-tag-row:hover { background: #dbeafe; }
.member-tag-row .member-remove { cursor: pointer; color: #94a3b8; font-size: 16px; font-weight: 700; }
.member-tag-row .member-remove:hover { color: #ef4444; }
.member-btn-add { padding: 3px 10px; border-radius: 6px; border: 1px solid #3b82f6; background: #3b82f6; color: #fff; font-size: 13px; font-weight: 700; cursor: pointer; }
.member-btn-add:hover { background: #2563eb; }

.kanban-btn-start, .kanban-btn-finish, .kanban-btn-cancel, .kanban-btn-skip { padding: 4px 12px; border-radius: 6px; border: none; font-size: 12px; font-weight: 600; cursor: pointer; white-space: nowrap; }
.kanban-btn-start { background: #3b82f6; color: #fff; }
.kanban-btn-start:hover { background: #2563eb; }
.kanban-btn-finish { background: #10b981; color: #fff; }
.kanban-btn-finish:hover { background: #059669; }
.kanban-btn-cancel { background: #f1f5f9; color: #64748b; border: 1px solid #e2e8f0; }
.kanban-btn-cancel:hover { background: #e2e8f0; }
.kanban-btn-skip { background: #fef3c7; color: #92400e; border: 1px solid #fcd34d; }
.kanban-btn-skip:hover { background: #fde68a; }

.operator-row { display: flex; align-items: center; gap: 6px; margin-bottom: 4px; }
.operator-label { font-size: 11px; color: #64748b; font-weight: 600; white-space: nowrap; }
.operator-select { flex: 1; padding: 3px 6px; border: 1px solid #e2e8f0; border-radius: 5px; font-size: 12px; color: #1e293b; background: #f8fafc; outline: none; cursor: pointer; }
.operator-select:focus { border-color: #3b82f6; background: #fff; }
.operator-dropdown { position: relative; flex: 1; }
.operator-dropdown-btn { display: block; width: 100%; padding: 3px 8px; border: 1px solid #e2e8f0; border-radius: 5px; font-size: 12px; color: #1e293b; background: #f8fafc; cursor: pointer; text-align: left; overflow: hidden; text-overflow: ellipsis; white-space: nowrap; }
.operator-dropdown-btn:hover { border-color: #3b82f6; }
.operator-dropdown-list { position: absolute; top: 100%; left: 0; min-width: 100%; z-index: 1000; background: #fff; border: 1px solid #e2e8f0; border-radius: 6px; max-height: 160px; overflow-y: auto; box-shadow: 0 4px 12px rgba(0,0,0,0.1); padding: 2px 0; }
.operator-checkbox { display: flex; align-items: center; gap: 6px; padding: 3px 10px; font-size: 12px; color: #1e293b; cursor: pointer; white-space: nowrap; margin: 0; font-weight: normal; }
.operator-checkbox:hover { background: #f1f5f9; }
.operator-checkbox input[type=checkbox] { margin: 0; cursor: pointer; }

.member-stat-card { background: #fff; border: 1px solid #e2e8f0; border-radius: 8px; padding: 6px 10px; display: flex; align-items: center; gap: 8px; box-shadow: 0 1px 2px rgba(0,0,0,0.04); transition: all 0.15s; }
.member-stat-card:hover { border-color: #93c5fd; box-shadow: 0 2px 8px rgba(59,130,246,0.12); transform: translateY(-1px); }
.member-stat-card.active-filter { border-color: #3b82f6; background: #eff6ff; box-shadow: 0 0 0 2px rgba(59,130,246,0.2); }
.member-stat-name { font-size: 12px; font-weight: 600; color: #1e293b; min-width: 50px; white-space: nowrap; }
.member-stat-nums { display: flex; gap: 8px; font-size: 13px; font-weight: 700; }
.member-stat-num.todo-num { color: #64748b; }
.member-stat-num.doing-num { color: #d97706; }
.member-stat-num.done-num { color: #059669; }
.member-stat-bar { flex: 1; height: 5px; background: #e2e8f0; border-radius: 3px; overflow: hidden; min-width: 30px; }
.member-stat-bar-fill { height: 100%; background: #10b981; border-radius: 3px; transition: width 0.3s; }

.member-task-list { margin-top: 6px; border-top: 1px solid #f1f5f9; padding-top: 4px; max-height: 140px; overflow-y: auto; background: #fafbfc; border-radius: 4px; padding: 4px; }
.member-task-item { display: flex; align-items: center; gap: 4px; padding: 2px 0; font-size: 13px; color: #475569; cursor: pointer; }
.member-task-item:hover { color: #1e293b; background: #f8fafc; }
.mt-sn { font-weight: 700; color: #1e293b; min-width: 32px; }
.mt-info { flex: 1; overflow: hidden; text-overflow: ellipsis; white-space: nowrap; }
.mt-qty { color: #64748b; font-weight: 600; white-space: nowrap; }
@media (max-width: 1200px){
  .kanban-board{overflow-x:auto;}
  .kanban-side-panel{min-width:320px;}
}
@media (max-width: 768px){
  .toolbar .search-box{max-width:none; min-width:220px;}
  .stats-bar{overflow-x:auto;}
  .stat-card{min-width:96px;}
  .kanban-column{min-width:280px; max-height:calc(100vh - 220px);}
}
</style>

<div class="d-flex justify-content-between align-items-center mb-3">
  <h4 class="fw-bold text-dark m-0">班组模块</h4>
  <div style="display:flex;gap:8px;align-items:center;">
    <select onchange="window.location.href='/teamkanban?month='+this.value+'&team='+encodeURIComponent(document.getElementById('teamSelect').value)" style="width:120px;border-radius:8px;padding:6px 10px;border:1px solid #d1d5db;font-size:13px;font-weight:600;color:#1f2937;background:#fff;cursor:pointer;">
      <option value="" {% if not month_filter %}selected{% endif %}>全部月份</option>
      {% for m in available_months %}
      <option value="{{m}}" {% if month_filter == m %}selected{% endif %}>{{m}}</option>
      {% endfor %}
    </select>
    <select id="teamSelect" onchange="switchTeam(this.value)" style="width:140px;border-radius:8px;padding:6px 10px;border:1px solid #d1d5db;font-size:13px;font-weight:600;color:#1f2937;background:#fff;cursor:pointer;">
      {% for team in teams %}
      <option value="{{team}}" {% if selected_team == team or (not selected_team and loop.first) %}selected{% endif %}>{{team}}</option>
      {% endfor %}
    </select>
    <div class="stat-card stat-todo" style="padding:4px 10px;border-radius:6px;">
      <div class="stat-info"><div class="stat-num" style="font-size:16px;">{{global_stats.todo}}</div><div class="stat-label" style="font-size:10px;">待处理</div></div>
    </div>
    <div class="stat-card stat-doing" style="padding:4px 10px;border-radius:6px;">
      <div class="stat-info"><div class="stat-num" style="font-size:16px;">{{global_stats.doing}}</div><div class="stat-label" style="font-size:10px;">处理中</div></div>
    </div>
    <div class="stat-card stat-done" style="padding:4px 10px;border-radius:6px;">
      <div class="stat-info"><div class="stat-num" style="font-size:16px;">{{global_stats.done}}</div><div class="stat-label" style="font-size:10px;">已完成</div></div>
    </div>
    <span class="small text-muted" style="line-height:32px;" id="lastRefreshTime">刷新于: {{ now.strftime("%H:%M:%S") }}</span>
    <button class="btn btn-sm btn-outline-secondary rounded-pill px-3" onclick="location.reload()">🔄 刷新</button>
  </div>
</div>

    <!-- 工具栏 -->
    <div class="toolbar">
      <div class="search-box">
        <span class="search-icon">🔍</span>
        <input type="text" id="kanbanSearch" placeholder="搜索序号 / 产品名 / 图号 / 批次..." oninput="filterKanbanCards()">
      </div>
      <div class="d-flex gap-2">
      <button id="toggleDoneBtn" class="btn btn-sm btn-outline-success rounded-pill px-3" onclick="openDonePage()" style="font-size:12px;">查看已完成任务</button>
      <button id="toggleSkippedBtn" class="btn btn-sm btn-outline-secondary rounded-pill px-3" onclick="showSkippedModal()" style="font-size:12px;">查看不涉及任务</button>
      </div>
    </div>

<!-- 单班组视图 -->
<div id="singleTeamView">
  {% for team in teams %}
  {% set ns = kanban_data[team] %}
  <div class="kanban-section" id="kanban-{{team}}" style="display:none;">
    <div class="kanban-board single-view">
      <div class="kanban-column todo" data-team="{{team}}">
        <div class="kanban-column-header">待处理<span class="badge rounded-pill bg-secondary">{{ns.todo|length}}</span>{% if ns.todo_overdue > 0 %}<span class="badge rounded-pill bg-danger" style="font-size:10px;">{{ns.todo_overdue}}逾期</span>{% endif %}</div>
        <div class="kanban-column-body">
          {% for item in ns.todo %}
          <div class="kanban-card {% if item.is_overdue %}overdue{% endif %}" data-search-text="{{item.serial_no}} {{item.product_name}} {{item.product_draw_no}} {{item.batch_no}} {{item.specific_model}}" data-task-id="{{item.id}}" data-team="{{team}}" data-operator="{{item.operator or ''}}">
            <div class="card-header-row"><span class="card-serial">{{item.serial_no}}</span><div class="card-badges">{% if item.is_overdue %}<span class="overdue-tag">逾期{{item.overdue_days}}天</span>{% endif %}{% if item.product_category %}<span class="badge bg-light text-dark" style="font-size:9px;">{{item.product_category}}</span>{% endif %}</div></div>
            <div class="card-meta"><span>{{item.product_draw_no or ''}}</span> <span>{{item.product_name or ''}}</span>{% if item.batch_no %} <span>{{item.batch_no}}</span>{% endif %}{% if item.total_qty %} <span>数量: {{item.total_qty}}</span>{% endif %}</div>
            <div class="card-meta">{% if item.plan_delivery_time %}<span>交付: {{fmt_date(item.plan_delivery_time)}}</span>{% endif %}</div>
            <div class="kanban-actions" style="display:flex;justify-content:space-between;align-items:center;">
              <div>
              <button class="kanban-btn-start" onclick="event.stopPropagation();kanbanAction('start',{{item.id}},'{{team}}')">▶ 开始进行</button>
              <button class="kanban-btn-skip" onclick="event.stopPropagation();kanbanAction('skip',{{item.id}},'{{team}}')">⊘ 不涉及</button>
              </div>
              <button class="btn btn-sm btn-outline-secondary py-0" onclick="event.stopPropagation();window.open('/task/{{item.id}}/details','_blank')">详情</button>
            </div>
          </div>
          {% else %}<div class="no-match-msg" style="display:block;">暂无待处理任务</div>{% endfor %}
        </div>
      </div>

      <div class="kanban-column doing" data-team="{{team}}">
        <div class="kanban-column-header">处理中<span class="badge rounded-pill bg-warning text-dark">{{ns.doing|length}}</span>{% if ns.doing_overdue > 0 %}<span class="badge rounded-pill bg-danger" style="font-size:10px;">{{ns.doing_overdue}}逾期</span>{% endif %}</div>
        <div class="kanban-column-body">
          {% for item in ns.doing %}
          <div class="kanban-card {% if item.is_overdue %}overdue{% endif %}" data-search-text="{{item.serial_no}} {{item.product_name}} {{item.product_draw_no}} {{item.batch_no}} {{item.specific_model}}" data-task-id="{{item.id}}" data-team="{{team}}" data-operator="{{item.operator or ''}}">
            <div class="card-header-row"><span class="card-serial">{{item.serial_no}}</span><div class="card-badges">{% if item.is_overdue %}<span class="overdue-tag">逾期{{item.overdue_days}}天</span>{% endif %}<span class="duration-tag">已{{item.duration_days}}天</span>{% if item.product_category %}<span class="badge bg-light text-dark" style="font-size:9px;">{{item.product_category}}</span>{% endif %}</div></div>
            <div class="card-meta"><span>{{item.product_draw_no or ''}}</span> <span>{{item.product_name or ''}}</span>{% if item.batch_no %} <span>{{item.batch_no}}</span>{% endif %}{% if item.total_qty %} <span>数量: {{item.total_qty}}</span>{% endif %}</div>
            <div class="card-meta"><span>开始: {{fmt_date(item.start_time)}}</span>{% if item.plan_delivery_time %} <span>交付: {{fmt_date(item.plan_delivery_time)}}</span>{% endif %}</div>
            <div class="kanban-actions">
              <div class="operator-row"><span class="operator-label">人员:</span>
                {% set member_names = [] %}{% for m in members_data.get(team, []) %}{% if not member_names.append(m.name) %}{% endif %}{% endfor %}
                <div class="operator-dropdown" onclick="event.stopPropagation()">
                  <span class="operator-dropdown-btn" onclick="event.stopPropagation();toggleOperatorDropdown(this.parentElement,{{item.id}})" data-task="{{item.id}}" data-selected="{% if item.operator %}{{item.operator}}{% endif %}">{% if item.operator %}{{item.operator|truncate(12,true,'...')}}{% else %}选择人员{% endif %} ▾</span>
                  <div class="operator-dropdown-list" style="display:none;">
                    {% for mn in member_names %}<label class="operator-checkbox"><input type="checkbox" value="{{mn}}" {% if item.operator and mn in item.operator.split(',') %}checked{% endif %} onchange="handleOperatorCheckbox(this,{{item.id}})"><span>{{mn}}</span></label>{% endfor %}
                    <div style="padding:4px;border-top:1px solid #e2e8f0;position:sticky;bottom:0;background:#fff;"><button class="btn btn-sm btn-secondary w-100 py-0" onclick="event.stopPropagation();confirmOperator(this,{{item.id}})" style="font-size:12px;">确定</button></div>
                  </div>
                </div>
              </div>
              <div style="display:flex;gap:6px;">
                <button class="kanban-btn-finish" onclick="event.stopPropagation();kanbanAction('finish',{{item.id}},'{{team}}')">✓ 完成</button>
                <button class="kanban-btn-cancel" onclick="event.stopPropagation();kanbanAction('cancel',{{item.id}},'{{team}}')">↩ 取消开始</button>
                <button class="btn btn-sm btn-outline-secondary py-0" onclick="event.stopPropagation();window.open('/task/{{item.id}}/details','_blank')" style="margin-left:auto;">详情</button>
              </div>
            </div>
          </div>
          {% else %}<div class="no-match-msg" style="display:block;">暂无处理中任务</div>{% endfor %}
        </div>
      </div>

      <div class="member-panel" style="flex:1;min-width:280px;flex-shrink:0;padding:0 4px;">
        <div class="members-dropdown" style="margin-bottom:6px;">
          <div class="members-dropdown-toggle" onclick="event.stopPropagation();toggleMembersDropdown(this)">
            <span>{{team}} 成员</span><span class="members-count" style="font-weight:500;color:#64748b;"></span><span style="font-size:10px;color:#94a3b8;margin-left:2px;">▾</span>
          </div>
          <div class="members-dropdown-menu">
            <div class="members-list-inner"></div>
            <div style="display:flex;gap:4px;padding:4px;border-top:1px solid #e2e8f0;margin-top:4px;">
              <input type="text" class="new-member-input" placeholder="添加成员" style="flex:1;padding:3px 6px;border:1px solid #d1d5db;border-radius:5px;font-size:12px;outline:none;" onclick="event.stopPropagation()" onkeydown="if(event.key==='Enter'){addMember();}">
              <button class="member-btn-add" onclick="addMember()" style="padding:3px 10px;font-size:12px;">+</button>
            </div>
          </div>
        </div>
        <div style="position:relative;max-height:500px;overflow-y:auto;margin-bottom:6px;background:#fff;border-radius:6px;border:1px solid #e2e8f0;padding:6px;"><div style="font-size:14px;font-weight:700;color:#64748b;margin-bottom:2px;">人员统计</div><div style="position:relative;" class="member-chart-wrap"><canvas class="member-stats-chart"></canvas></div></div>
        <div class="filter-badge" style="display:none;align-items:center;gap:4px;padding:2px 8px;margin-bottom:4px;background:#eff6ff;border:1px solid #bfdbfe;border-radius:4px;font-size:11px;"><span style="color:#1e40af;font-weight:600;">筛选:</span><span class="filter-name" style="color:#1e40af;"></span><span onclick="clearOperatorFilter()" style="cursor:pointer;color:#94a3b8;font-weight:700;margin-left:auto;">x 清除</span></div>
        <div style="font-size:13px;font-weight:700;color:#374151;margin-bottom:2px;">进行中任务详情</div>
        <div class="member-task-panel" style="max-height:282px;overflow-y:auto;border:1px solid #e2e8f0;border-radius:6px;padding:3px;background:#fafbfc;display:none;"></div>
      </div>

      <div class="kanban-side-panel">
        <div class="task-chart-box">
          <div style="font-size:14px;font-weight:700;color:#64748b;margin-bottom:6px;">任务数据统计 <span class="small text-muted task-rate-label"></span></div>
          <div style="position:relative;height:160px;"><canvas class="task-stats-chart"></canvas></div>
          <div class="task-stat-grid">
            <div class="task-stat-tile"><div class="task-stat-num task-total-num">0</div><div class="task-stat-label">总任务</div></div>
            <div class="task-stat-tile"><div class="task-stat-num task-todo-num">0</div><div class="task-stat-label">待处理</div></div>
            <div class="task-stat-tile"><div class="task-stat-num task-doing-num">0</div><div class="task-stat-label">处理中</div></div>
            <div class="task-stat-tile"><div class="task-stat-num task-overdue-num">0</div><div class="task-stat-label">逾期</div></div>
          </div>
        </div>
        <div class="branch-notice">
          <div class="branch-notice-head">
            <div class="branch-notice-title">通知公告</div>
          </div>
          <div class="branch-notice-list">
            {% for notice in branch_notices %}
            <div class="branch-notice-item">
              <div class="branch-notice-meta">{{notice.created_time.strftime("%Y/%m/%d %H:%M")}} · {{(notice.user.username or notice.user.account) if notice.user else '系统'}}</div>
              <div class="branch-notice-name">{{notice.title}}</div>
              <div class="branch-notice-content">{{notice.content}}</div>
            </div>
            {% else %}
            <div class="branch-notice-empty">暂无通知</div>
            {% endfor %}
          </div>
        </div>
      </div>

      <div class="kanban-column done" data-team="{{team}}" style="display:none;">
        <div class="kanban-column-header">已完成 <span class="badge rounded-pill bg-success">{{ns.done|length}}</span></div>
        <div class="kanban-column-body">{% for item in ns.done %}<div class="kanban-card" data-search-text="{{item.serial_no}} {{item.product_name}} {{item.product_draw_no}} {{item.batch_no}} {{item.specific_model}}" data-operator="{{item.operator or ''}}" onclick="window.open('/task/{{item.id}}/details','_blank')"><div class="card-header-row"><span class="card-serial">{{item.serial_no}}</span><div class="card-badges"><span class="done-tag">{{fmt_date(item.end_time)}}</span>{% if item.product_category %}<span class="badge bg-light text-dark" style="font-size:9px;">{{item.product_category}}</span>{% endif %}</div></div><div class="card-name">{{item.product_name or ''}}</div><div class="card-meta"><span>{{item.product_draw_no or ''}}{% if item.batch_no %} | {{item.batch_no}}{% endif %}</span>{% if item.total_qty %}<span>数量: {{item.total_qty}}</span>{% endif %}{% if item.specific_model %}<span>型号: {{item.specific_model}}</span>{% endif %}{% if item.plan_delivery_time %}<span>交付: {{fmt_date(item.plan_delivery_time)}}</span>{% endif %}</div></div>{% else %}<div class="no-match-msg" style="display:block;">暂无已完成任务</div>{% endfor %}</div>
      </div>
      <div class="kanban-column skipped" data-team="{{team}}" style="display:none;">
        <div class="kanban-column-header">不涉及 <span class="badge rounded-pill bg-secondary">{{ns.skipped|length}}</span></div>
        <div class="kanban-column-body">{% for item in ns.skipped %}<div class="kanban-card" data-search-text="{{item.serial_no}} {{item.product_name}} {{item.product_draw_no}} {{item.batch_no}} {{item.specific_model}}" data-task-id="{{item.id}}" data-team="{{team}}" data-operator="{{item.operator or ''}}"><div class="card-header-row"><span class="card-serial">{{item.serial_no}}</span><div class="card-badges"><span class="badge bg-light text-dark" style="font-size:9px;">{{item.product_category or ''}}</span></div></div><div class="card-name">{{item.product_name or ''}}</div><div class="card-meta"><span>{{item.product_draw_no or ''}}{% if item.batch_no %} | {{item.batch_no}}{% endif %}</span>{% if item.total_qty %}<span>数量: {{item.total_qty}}</span>{% endif %}{% if item.specific_model %}<span>型号: {{item.specific_model}}</span>{% endif %}</div><div class="kanban-actions" style="display:flex;justify-content:space-between;align-items:center;"><button class="kanban-btn-cancel" onclick="event.stopPropagation();kanbanAction('cancel',{{item.id}},'{{team}}')">↩ 恢复</button><button class="btn btn-sm btn-outline-secondary py-0" onclick="event.stopPropagation();window.open('/task/{{item.id}}/details','_blank')">详情</button></div></div>{% else %}<div class="no-match-msg" style="display:block;">无不涉及任务</div>{% endfor %}</div>
      </div>
    </div>
  </div>
  {% endfor %}
</div>

<!-- 已完成任务弹窗 -->
<div class="modal fade" id="doneModal" tabindex="-1">
  <div class="modal-dialog modal-lg">
    <div class="modal-content">
      <div class="modal-header"><h5 class="modal-title">已完成任务</h5><button type="button" class="btn-close" data-bs-dismiss="modal"></button></div>
      <div class="modal-body" id="doneModalBody" style="max-height:60vh;overflow-y:auto;"></div>
    </div>
  </div>
</div>

<!-- 不涉及任务弹窗 -->
<div class="modal fade" id="skippedModal" tabindex="-1">
  <div class="modal-dialog modal-lg">
    <div class="modal-content">
      <div class="modal-header"><h5 class="modal-title">不涉及任务</h5><button type="button" class="btn-close" data-bs-dismiss="modal"></button></div>
      <div class="modal-body" id="skippedModalBody" style="max-height:60vh;overflow-y:auto;"></div>
    </div>
  </div>
</div>

<script>
var currentTeam = '{{selected_team if selected_team else teams[0]}}';
var currentOperatorFilter = '';

function $p(sel) { return document.querySelector('#kanban-' + currentTeam + ' ' + sel); }

var membersMap = { {% for team in teams %}'{{team}}': [{% for m in members_data.get(team, []) %}{id:{{m.id}},name:'{{m.name}}'}{% if not loop.last %},{% endif %}{% endfor %}],{% endfor %} };

var memberStatsMap = { {% for team in teams %}'{{team}}': [{% for name, stat in kanban_data[team].member_stats.items() %}{name:'{{name}}',todo:{{stat.todo}},doing:{{stat.doing}},done:{{stat.done}}}{% if not loop.last %},{% endif %}{% endfor %}],{% endfor %} };

var taskStatsMap = { {% for team in teams %}'{{team}}': {todo:{{kanban_data[team].todo|length}}, doing:{{kanban_data[team].doing|length}}, done:{{kanban_data[team].done|length}}, overdue:{{kanban_data[team].overdue_count}}},{% endfor %} };

var memberTasksMap = { {% for team in teams %}'{{team}}': { {% for op, items in kanban_data[team].member_tasks.items() %}'{{op}}': [{% for it in items %}{id:{{it.id}},sn:'{{it.serial_no}}',dn:'{{it.product_draw_no}}',pn:'{{it.product_name}}',bn:'{{it.batch_no}}',qty:{{it.total_qty}}}{% if not loop.last %},{% endif %}{% endfor %}]{% if not loop.last %},{% endif %}{% endfor %}},{% endfor %} };

function switchTeam(team) { currentTeam = team; currentOperatorFilter = ''; updateFilterBadge(); updateActiveStatCard(); document.querySelectorAll('#singleTeamView .kanban-section').forEach(function(s) { s.style.display = s.id === 'kanban-' + team ? 'block' : 'none'; }); document.querySelectorAll('.kanban-column.skipped').forEach(function(c) { c.style.display = 'none'; }); renderMembers(team); renderMemberStats(team); renderTaskStats(team); filterKanbanCards(); }

function kanbanAction(action, taskId, team) { if (!confirm('确定要执行此操作吗？')) return; fetch('/api/kanban/' + action + '/' + taskId, {method:'POST',headers:{'Content-Type':'application/x-www-form-urlencoded'},body:'team='+encodeURIComponent(team)}).then(function(r){return r.json()}).then(function(data){if(data.ok){var m=(new URLSearchParams(window.location.search)).get('month')||'';window.top.location='/teamkanban?team='+encodeURIComponent(team)+(m?'&month='+m:'')+'&_='+Date.now()}else{alert(data.msg)}}).catch(function(e){alert('操作失败: '+e.message)}); }

function saveOperator(taskId, value) { fetch('/api/kanban/operator/' + taskId, {method:'POST',headers:{'Content-Type':'application/x-www-form-urlencoded'},body:'operator='+encodeURIComponent(value)}).then(function(r){return r.json()}).then(function(data){if(!data.ok){alert(data.msg)}}).catch(function(e){alert('保存失败: '+e.message)}); }

function toggleOperatorDropdown(el, taskId) { var list = el.querySelector('.operator-dropdown-list'); var isOpen = list.style.display === 'block'; document.querySelectorAll('.operator-dropdown-list').forEach(function(l) { l.style.display = 'none'; }); if (!isOpen) { list.style.display = 'block'; } }
function handleOperatorCheckbox(cb, taskId) { }

function confirmOperator(btn, taskId) {
  var dropdown = btn.closest('.operator-dropdown');
  var checked = dropdown.querySelectorAll('input[type=checkbox]:checked');
  var names = []; checked.forEach(function(c) { names.push(c.value); });
  var val = names.join(',');
  btn.closest('.operator-dropdown-list').style.display = 'none';
  fetch('/api/kanban/operator/' + taskId, {method:'POST',headers:{'Content-Type':'application/x-www-form-urlencoded'},body:'operator='+encodeURIComponent(val)})
    .then(function(r){return r.json()}).then(function(data){
      if(data.ok){ var m=(new URLSearchParams(window.location.search)).get('month')||''; window.location.href='/teamkanban?team='+encodeURIComponent(currentTeam)+(m?'&month='+m:'')+'&_='+Date.now() }
      else{alert(data.msg)}
    }).catch(function(e){alert('保存失败: '+e.message)});
}
document.addEventListener('click', function(e) { if (!e.target.closest('.operator-dropdown')) { document.querySelectorAll('.operator-dropdown-list').forEach(function(l) { l.style.display = 'none'; }); } if (!e.target.closest('.members-dropdown')) { document.querySelectorAll('.members-dropdown-menu').forEach(function(l) { l.style.display = 'none'; }); } });

function renderMembers(team) { var list = $p('.members-list-inner'); var countEl = $p('.members-count'); if (!list) return; var members = membersMap[team] || []; if (countEl) countEl.textContent = '('+members.length+')'; var html = ''; members.forEach(function(m) { html += '<div class="member-tag-row"><span class="member-name-text" data-member-id="'+m.id+'" onclick="event.stopPropagation();editMemberName(this,'+m.id+')" title="点击编辑">'+m.name+'</span><span class="member-remove" onclick="event.stopPropagation();removeMember('+m.id+')">&times;</span></div>'; }); if (!members.length) { html = '<div style="font-size:12px;color:#94a3b8;text-align:center;padding:8px;">暂无成员</div>'; } list.innerHTML = html; }

function editMemberName(el, id) { var oldName = el.textContent; var input = document.createElement('input'); input.type = 'text'; input.value = oldName; input.className = 'member-edit-input'; input.style.cssText = 'width:60px;padding:1px 4px;border:1px solid #3b82f6;border-radius:4px;font-size:12px;outline:none;'; input.onclick = function(e) { e.stopPropagation(); }; input.onkeydown = function(e) { if (e.key === 'Enter') { input.blur(); } if (e.key === 'Escape') { input.value = oldName; input.blur(); } }; input.onblur = function() { var newName = input.value.trim(); if (newName && newName !== oldName) { fetch('/api/kanban/members/edit/' + id, {method:'POST',headers:{'Content-Type':'application/x-www-form-urlencoded'},body:'name='+encodeURIComponent(newName)}).then(function(r){return r.json()}).then(function(data){if(data.ok){membersMap[currentTeam].forEach(function(m){if(m.id===id)m.name=newName});renderMembers(currentTeam);renderMemberStats(currentTeam)}else{alert(data.msg);renderMembers(currentTeam)}}).catch(function(e){alert('编辑失败');renderMembers(currentTeam)}); } else { renderMembers(currentTeam); } }; el.replaceWith(input); input.focus(); input.select(); }

function removeMember(id) { var team = currentTeam; if (!confirm('确定要移除该成员吗？')) return; fetch('/api/kanban/members/remove/' + id, {method:'POST'}).then(function(r){return r.json()}).then(function(data){if(data.ok){membersMap[team]=membersMap[team].filter(function(m){return m.id!==id});renderMembers(team)}else{alert(data.msg)}}).catch(function(e){alert('移除失败: '+e.message)}); }

function addMember() { var input = $p('.new-member-input'); if (!input) return; var name = input.value.trim(); if (!name) return; fetch('/api/kanban/members/add', {method:'POST',headers:{'Content-Type':'application/x-www-form-urlencoded'},body:'team='+encodeURIComponent(currentTeam)+'&name='+encodeURIComponent(name)}).then(function(r){return r.json()}).then(function(data){if(data.ok){membersMap[currentTeam].push({id:data.id,name:name});renderMembers(currentTeam);input.value=''}else{alert(data.msg)}}).catch(function(e){alert('添加失败: '+e.message)}); }

function toggleMembersDropdown(el) { var menu = el.parentElement.querySelector('.members-dropdown-menu'); var isOpen = menu.style.display === 'block'; document.querySelectorAll('.members-dropdown-menu').forEach(function(m) { m.style.display = 'none'; }); if (!isOpen) { menu.style.display = 'block'; } }

function renderMemberStats(team) { var container = $p('.member-stats-container'); var stats = memberStatsMap[team] || []; var members = membersMap[team] || []; var merged = {}; members.forEach(function(m) { merged[m.name] = {name: m.name, todo: 0, doing: 0, done: 0}; }); stats.forEach(function(s) { if (!merged[s.name]) return; merged[s.name].todo += s.todo; merged[s.name].doing += s.doing; merged[s.name].done += s.done; }); var chartData = Object.values(merged); chartData.sort(function(a, b) { return (b.todo + b.doing + b.done) - (a.todo + a.doing + a.done); }); if (!chartData.length) { if (container) container.innerHTML = '<span style="font-size:12px;color:#94a3b8;">暂无人员任务统计</span>'; if (window.memberChart) { window.memberChart.destroy(); window.memberChart = null; } renderMemberTaskPanel(team); return; } var html = ''; chartData.forEach(function(s) { var total = s.todo + s.doing + s.done; var escapedName = s.name.replace(/'/g, "&#39;"); html += '<div class="member-stat-card" data-op-name="'+escapedName+'" style="cursor:pointer;" title="点击筛选该人员的任务"><div class="member-stat-name">'+s.name+'</div><div class="member-stat-nums"><span class="member-stat-num todo-num">'+s.todo+'</span><span class="member-stat-num doing-num">'+s.doing+'</span><span class="member-stat-num done-num">'+s.done+'</span></div><div class="member-stat-bar"><div class="member-stat-bar-fill" style="width:'+(total?Math.round(s.done/total*100):0)+'%"></div></div></div>'; }); if (container) container.innerHTML = html; renderMemberTaskPanel(team); var canvas = $p('.member-stats-chart'); if (!canvas) return; var ctx = canvas.getContext('2d'); if (window.memberChart) { window.memberChart.destroy(); } var labels = chartData.map(function(s) { return s.name; }); window.memberChart = new Chart(ctx, {type:'bar',data:{labels:labels,datasets:[{label:'待处理',data:chartData.map(function(s){return s.todo}),backgroundColor:'#94a3b8',borderRadius:4},{label:'处理中',data:chartData.map(function(s){return s.doing}),backgroundColor:'#f59e0b',borderRadius:4},{label:'已完成',data:chartData.map(function(s){return s.done}),backgroundColor:'#10b981',borderRadius:4}]},options:{indexAxis:'y',responsive:true,maintainAspectRatio:false,plugins:{legend:{position:'top',labels:{boxWidth:12,padding:8,font:{size:11}}},datalabels:{display:false}},scales:{x:{stacked:true,ticks:{stepSize:1,font:{size:10}},grid:{display:false}},y:{stacked:true,ticks:{font:{size:11}}}}}}); var chartWrap=$p('.member-chart-wrap'); if(chartWrap){var h=Math.max(350,chartData.length*18+50);chartWrap.style.height=h+'px';} }

function renderMemberTaskPanel(team) { var panel = $p('.member-task-panel'); if (!panel) return; var tasks = memberTasksMap[team] || {}; var members = membersMap[team] || []; var memberNames = members.map(function(m) { return m.name; }); var names = Object.keys(tasks).filter(function(n) { return memberNames.indexOf(n) !== -1; }); if (!names.length) { panel.style.display = 'none'; return; } names.sort(function(a, b) { return tasks[b].length - tasks[a].length; }); var html = ''; names.forEach(function(name) { html += '<div style="font-size:14px;font-weight:700;color:#1e293b;padding:2px 6px;background:#e2e8f0;border-radius:3px;margin-bottom:2px;">'+name+' ('+tasks[name].length+')</div>'; tasks[name].forEach(function(t) { html += '<div class="member-task-item" data-task-id="'+t.id+'" style="padding:2px 6px;"><span class="mt-sn">'+t.sn+'</span><span class="mt-info">'+t.pn+' | '+t.dn+(t.bn?' | '+t.bn:'')+'</span><span class="mt-qty">'+t.qty+'</span></div>'; }); }); panel.innerHTML = html; panel.style.display = 'block'; }

function renderTaskStats(team) {
  var stat = taskStatsMap[team] || {todo:0, doing:0, done:0, overdue:0};
  var total = stat.todo + stat.doing + stat.done;
  var rate = total ? Math.round(stat.done / total * 100) : 0;
  var panel = document.querySelector('#kanban-' + team + ' .kanban-side-panel');
  if (!panel) return;
  var label = panel.querySelector('.task-rate-label');
  if (label) label.textContent = '完成率 ' + rate + '%';
  panel.querySelector('.task-total-num').textContent = total;
  panel.querySelector('.task-todo-num').textContent = stat.todo;
  panel.querySelector('.task-doing-num').textContent = stat.doing;
  panel.querySelector('.task-overdue-num').textContent = stat.overdue;
  var canvas = panel.querySelector('.task-stats-chart');
  if (!canvas) return;
  if (window.taskChart) { window.taskChart.destroy(); }
  window.taskChart = new Chart(canvas.getContext('2d'), {
    type: 'doughnut',
    data: { labels: ['已完成','未完成'], datasets: [{ data: [stat.done, stat.todo + stat.doing], backgroundColor: ['#10b981','#e2e8f0'], borderWidth: 2 }] },
    options: { responsive: true, maintainAspectRatio: false, cutout: '60%',
      plugins: { legend: { position: 'bottom', labels: { boxWidth: 10, font: { size: 10 } } },
        datalabels: { color: '#1e293b', font: { weight: '700', size: 11 }, formatter: function(v) { return total ? Math.round(v/total*100)+'%' : ''; } }
      }
    }
  });
}

function filterByOperator(name) { if (name === '未分配') { currentOperatorFilter = '__unassigned__'; } else { currentOperatorFilter = name; } updateFilterBadge(); updateActiveStatCard(); filterKanbanCards(); }
function clearOperatorFilter() { currentOperatorFilter = ''; updateFilterBadge(); updateActiveStatCard(); filterKanbanCards(); }
function updateFilterBadge() { var badge = $p('.filter-badge'); if (!badge) return; if (currentOperatorFilter) { badge.style.display = 'inline-flex'; badge.querySelector('.filter-name').textContent = currentOperatorFilter; } else { badge.style.display = 'none'; } }
function updateActiveStatCard() { document.querySelectorAll('#kanban-'+currentTeam+' .member-stat-card').forEach(function(c) { var op = c.dataset.opName ? c.dataset.opName.replace(/&#39;/g, "'") : ''; c.classList.toggle('active-filter', op === currentOperatorFilter || (currentOperatorFilter === '__unassigned__' && op === '未分配')); }); }

function filterKanbanCards() { var query = (document.getElementById('kanbanSearch').value || '').toLowerCase().trim(); function processCards(container) { if (!container) return; var columns = container.querySelectorAll('.kanban-column'); columns.forEach(function(col) { if (col.style.display === 'none') return; var cards = col.querySelectorAll('.kanban-card'); var body = col.querySelector('.kanban-column-body'); var noMatchMsg = body ? body.querySelector('.no-match-msg') : null; var hasVisible = false; cards.forEach(function(card) { var text = (card.dataset.searchText || '').toLowerCase(); var op = (card.dataset.operator || ''); var shown = query === '' || text.includes(query); if (currentOperatorFilter) { if (currentOperatorFilter === '__unassigned__') { shown = shown && !op; } else { shown = shown && (op === currentOperatorFilter || op.split(',').map(function(s){return s.trim();}).indexOf(currentOperatorFilter) !== -1); } } card.classList.toggle('hidden-by-search', !shown); if (shown) hasVisible = true; }); if (noMatchMsg) { noMatchMsg.style.display = (cards.length === 0 || !hasVisible) ? 'block' : 'none'; } }); } processCards(document.getElementById('singleTeamView')); }

function openDonePage() {
  var col = document.querySelector('#kanban-' + currentTeam + ' .kanban-column.done');
  var body = document.getElementById('doneModalBody');
  if (col) {
    body.innerHTML = col.querySelector('.kanban-column-body').innerHTML;
  } else {
    body.innerHTML = '<div class="text-center text-muted py-3">暂无已完成任务</div>';
  }
  var modal = new bootstrap.Modal(document.getElementById('doneModal'));
  modal.show();
}
function showSkippedModal() {
  var col = document.querySelector('#kanban-' + currentTeam + ' .kanban-column.skipped');
  var body = document.getElementById('skippedModalBody');
  if (col) {
    body.innerHTML = col.querySelector('.kanban-column-body').innerHTML;
  } else {
    body.innerHTML = '<div class="text-center text-muted py-3">无不涉及任务</div>';
  }
  var modal = new bootstrap.Modal(document.getElementById('skippedModal'));
  modal.show();
}
(function init() { var firstTeam = '{{selected_team if selected_team else teams[0]}}'; document.querySelectorAll('#singleTeamView .kanban-section').forEach(function(s) { s.style.display = s.id === 'kanban-' + firstTeam ? 'block' : 'none'; }); renderMembers(firstTeam); renderMemberStats(firstTeam); renderTaskStats(firstTeam); document.getElementById('singleTeamView').addEventListener('click', function(e) { var card = e.target.closest('.member-stat-card'); if (card && card.dataset.opName) { var name = card.dataset.opName.replace(/&#39;/g, "'"); filterByOperator(name); } }); })();

document.addEventListener('keydown', function(e) { if ((e.ctrlKey || e.metaKey) && e.key === 'k') { e.preventDefault(); document.getElementById('kanbanSearch').focus(); } });
</script>
{% endblock %}
''')

ADMIN_USERS_HTML = BASE_HTML.replace('{% block content %}{% endblock %}','''
{% block content %}
<h4>系统用户管理</h4>
<table class="table table-sm bg-white shadow-sm border mt-3">
<thead><tr class="bg-light"><th>ID</th><th>账号</th><th>用户名</th><th>角色</th><th>班组</th><th>产线</th><th style="width:350px;">修改信息</th><th>操作</th></tr></thead>
<tbody>
{% for u in users %}
<tr>
  <td>{{u.id}}</td>
  <td class="fw-bold">{{u.account}}</td>
  <td>{{u.username}}</td>
  <td>{{u.role.name}}</td>
  <td>{{u.team or '' if u.role.name == '班组长' else '--'}}</td>
  <td>{{u.production_line or '' if u.role.name == '计调员' else '--'}}</td>
  <td>
    <form method="post" action="/admin/user/edit/{{u.id}}" class="d-flex gap-1 align-items-center">
      <input type="text" name="account" class="form-control form-control-sm" style="width:80px;" value="{{u.account}}" placeholder="账号">
      <input type="text" name="username" class="form-control form-control-sm" style="width:80px;" value="{{u.username or ''}}" placeholder="用户名">
      <select name="role_id" class="form-select form-select-sm role-sel" style="width:90px;" onchange="var tr=this.closest('tr');var rn=this.selectedOptions[0].text;tr.querySelector('.team-wrap').style.display=rn==='班组长'?'':'none';tr.querySelector('.pline-wrap').style.display=rn==='计调员'?'':'none'">
        {% for r in roles %}<option value="{{r.id}}" {% if u.role_id==r.id %}selected{% endif %}>{{r.name}}</option>{% endfor %}
      </select>
      <span class="team-wrap" style="{% if u.role.name != '班组长' %}display:none{% endif %}">
        <select name="team" class="form-select form-select-sm" style="width:90px;">
          <option value="">--班组--</option>
          {% for t in ['裁剪','缝纫','粘胶','总装1','总装2','氧调','热风热合'] %}
          <option value="{{t}}" {% if u.team==t %}selected{% endif %}>{{t}}</option>
          {% endfor %}
        </select>
      </span>
      <span class="pline-wrap" style="{% if u.role.name != '计调员' %}display:none{% endif %}">
        <select name="production_line" class="form-select form-select-sm" style="width:100px;">
          <option value="">--产线--</option>
          {% for pl in ['头盔面罩','服装','船囊'] %}
          <option value="{{pl}}" {% if u.production_line==pl %}selected{% endif %}>{{pl}}</option>
          {% endfor %}
        </select>
      </span>
      <input type="password" name="password" class="form-control form-control-sm" style="width:100px;" placeholder="新密码" autocomplete="new-password">
      <button class="btn btn-outline-primary btn-sm py-0">保存</button>
    </form>
  </td>
  <td>
    {% if u.account != 'admin' %}
    <form method="post" action="/admin/user/delete/{{u.id}}" onsubmit="return confirm('确认删除用户 {{u.account}} ？')" style="display:inline;">
      <button class="btn btn-outline-danger btn-sm py-0">删除</button>
    </form>
    {% endif %}
  </td>
</tr>
{% endfor %}
</tbody></table>
<h5 class="mt-4">添加用户</h5>
<form method="post" action="/admin/user/add" class="row g-2">
<div class="col-auto"><input class="form-control form-control-sm" name="account" placeholder="账号" required></div>
<div class="col-auto"><input class="form-control form-control-sm" name="username" placeholder="用户名"></div>
<div class="col-auto"><input type="password" class="form-control form-control-sm" name="password" placeholder="密码" required autocomplete="new-password"></div>
<div class="col-auto"><select class="form-select form-select-sm" name="role_id" onchange="var f=this.closest('form');var rn=this.selectedOptions[0].text;f.querySelector('.team-wrap').style.display=rn==='班组长'?'':'none';f.querySelector('.pline-wrap').style.display=rn==='计调员'?'':'none'">{% for r in roles %}<option value="{{r.id}}">{{r.name}}</option>{% endfor %}</select></div>
<span class="team-wrap" style="display:none"><select class="form-select form-select-sm" name="team"><option value="">--班组--</option>{% for t in ['裁剪','缝纫','粘胶','总装1','总装2','氧调','热风热合'] %}<option value="{{t}}">{{t}}</option>{% endfor %}</select></span>
<span class="pline-wrap" style="display:none"><select class="form-select form-select-sm" name="production_line"><option value="">--产线--</option>{% for pl in ['头盔面罩','服装','船囊'] %}<option value="{{pl}}">{{pl}}</option>{% endfor %}</select></span>
<div class="col-auto"><button class="btn btn-primary btn-sm px-4">添加</button></div>
</form>
{% endblock %}''')

DONE_LIST_HTML = BASE_HTML.replace('{% block content %}{% endblock %}','''
{% block content %}
<style>
.done-toolbar { display: flex; justify-content: space-between; align-items: center; margin-bottom: 10px; gap: 10px; }
.done-toolbar .search-box { position: relative; flex: 1; max-width: 360px; }
.done-toolbar .search-box input { width: 100%; padding: 8px 12px 8px 36px; border: 1px solid #e2e8f0; border-radius: 8px; font-size: 13px; outline: none; }
.done-toolbar .search-box .search-icon { position: absolute; left: 12px; top: 50%; transform: translateY(-50%); opacity: 0.4; font-size: 15px; }
</style>
<div class="d-flex justify-content-between align-items-center mb-3">
  <h4 class="fw-bold text-dark m-0">{{team}} - 已完成任务</h4>
  <div style="display:flex;gap:8px;align-items:center;">
    <span class="small text-muted">共 <span id="totalCount">{{items|length}}</span> 项</span>
    <button class="btn btn-sm btn-outline-success rounded-pill px-3" onclick="location.href='/teamkanban/done/{{team}}/export?month={{month_filter or ''}}'">导出Excel</button>
    <button class="btn btn-sm btn-outline-secondary rounded-pill px-3" onclick="window.close()">关闭</button>
  </div>
</div>
<div class="done-toolbar">
  <div class="search-box">
    <span class="search-icon">🔍</span>
    <input type="text" id="doneSearch" placeholder="搜索序号 / 图号 / 产品名 / 型号 / 批次 / 操作人员..." oninput="filterDoneTable()">
  </div>
  <button class="btn btn-sm btn-outline-secondary rounded-pill px-3" onclick="clearDoneFilters()" style="font-size:12px;">清除</button>
</div>
<div style="background:#fff;border-radius:10px;border:1px solid #e2e8f0;overflow:hidden;">
  <table class="table table-sm table-hover m-0" style="font-size:14px;" id="doneTable">
    <thead style="background:#f8fafc;">
      <tr><th>序号</th><th>图号</th><th>产品名称</th><th>型号</th><th>批次</th><th>数量</th><th>操作人员</th><th>开始</th><th>完成</th></tr>
    </thead>
    <tbody>
    {% for it in items %}
    <tr class="done-row" data-search="{{it.serial_no}} {{it.product_draw_no}} {{it.product_name}} {{it.specific_model}} {{it.batch_no}} {{it.operator}}" onclick="window.open('/task/{{it.id}}/details','_blank')" style="cursor:pointer;">
      <td>{{it.serial_no}}</td><td>{{it.product_draw_no}}</td><td>{{it.product_name}}</td>
      <td>{{it.specific_model}}</td><td>{{it.batch_no}}</td><td>{{it.total_qty}}</td>
      <td>{{it.operator}}</td><td>{{it.start_time}}</td><td>{{it.end_time}}</td>
    </tr>
    {% else %}
    <tr id="noDataRow"><td colspan="9" class="text-center text-muted py-4">暂无已完成任务</td></tr>
    {% endfor %}
    </tbody>
  </table>
</div>
<div id="noResult" style="display:none;text-align:center;padding:40px;color:#94a3b8;">无匹配结果</div>
<script>
function filterDoneTable() {
  var query = (document.getElementById('doneSearch').value || '').toLowerCase().trim();
  var rows = document.querySelectorAll('.done-row');
  var visible = 0;
  rows.forEach(function(row) {
    var text = (row.dataset.search || '').toLowerCase();
    var shown = query === '' || text.indexOf(query) !== -1;
    row.style.display = shown ? '' : 'none';
    if (shown) visible++;
  });
  document.getElementById('totalCount').textContent = visible;
  var noResult = document.getElementById('noResult');
  if (rows.length === 0) { noResult.style.display = 'none'; }
  else if (visible === 0) { noResult.style.display = 'block'; }
  else { noResult.style.display = 'none'; }
}
function clearDoneFilters() { document.getElementById('doneSearch').value = ''; filterDoneTable(); }
document.addEventListener('keydown', function(e) { if ((e.ctrlKey || e.metaKey) && e.key === 'k') { e.preventDefault(); document.getElementById('doneSearch').focus(); } });
</script>
{% endblock %}
''')

ADMIN_PERM_HTML = BASE_HTML.replace('{% block content %}{% endblock %}','''
{% block content %}<h4>安全中心：用户级读写权限精准下发</h4>
<form method="post" action="/admin/permissions/set" class="mt-3">
<div class="d-flex align-items-center mb-3">
    <span class="fw-bold me-2">选择配属用户:</span>
    <select name="user_id" class="form-select bg-white shadow-sm border-0" style="width:250px">
      {% for u in users %}<option value="{{u.id}}">{{u.username}} (标签: {{u.role.name}})</option>{% endfor %}
    </select>
</div>
<div class="accordion shadow-sm border-0" id="permAccordion">
{% for gname,gcols in groups.items() %}
<div class="accordion-item border-0 border-bottom">
  <h2 class="accordion-header"><button class="accordion-button collapsed bg-light fw-bold text-dark" type="button" data-bs-toggle="collapse" data-bs-target="#collapse{{loop.index}}">板块配置：{{gname}}</button></h2>
  <div id="collapse{{loop.index}}" class="accordion-collapse collapse">
    <div class="accordion-body p-0"><table class="table table-sm m-0"><thead><tr class="bg-light"><th>关联字段名称</th><th>开启[查阅]权限</th><th>开启[编辑]权限</th></tr></thead><tbody>
    {% for col,name in gcols %}<tr><td>{{name}}</td><td><input type="checkbox" name="view_{{col}}" class="form-check-input"></td><td><input type="checkbox" name="edit_{{col}}" class="form-check-input"></td></tr>{% endfor %}
    </tbody></table></div>
  </div>
</div>{% endfor %}
</div><button class="btn btn-primary mt-4 px-5">保存策略</button></form>{% endblock %}''')

ADMIN_LOGS_HTML = BASE_HTML.replace('{% block content %}{% endblock %}','''
{% block content %}
<h4>操作日志</h4>
<div class="card shadow-sm border-0 mb-3">
  <div class="card-body">
    <div class="row g-2 align-items-end">
      <div class="col-auto">
        <label class="form-label small fw-bold">操作类型</label>
        <select id="filterOpType" class="form-select form-select-sm" style="width:140px">
          <option value="">全部</option>
          <option value="创建任务">创建任务</option>
          <option value="批量新增">批量新增</option>
          <option value="批量更新">批量更新</option>
          <option value="删除任务">删除任务</option>
          <option value="编辑字段">编辑字段</option>
        </select>
      </div>
      <div class="col-auto">
        <label class="form-label small fw-bold">开始日期</label>
        <input type="date" id="filterStartDate" class="form-control form-control-sm" style="width:150px">
      </div>
      <div class="col-auto">
        <label class="form-label small fw-bold">结束日期</label>
        <input type="date" id="filterEndDate" class="form-control form-control-sm" style="width:150px">
      </div>
      <div class="col-auto">
        <button class="btn btn-primary btn-sm" onclick="loadLogs(1)">查询</button>
        <button class="btn btn-outline-secondary btn-sm" onclick="resetLogFilters()">重置</button>
      </div>
    </div>
  </div>
</div>

<div class="card shadow-sm border-0">
  <div class="card-body p-0">
    <div class="table-responsive">
      <table class="table table-hover table-sm m-0" style="table-layout:fixed">
        <thead class="table-dark">
          <tr>
            <th>操作时间</th>
            <th>操作类型</th>
            <th>任务序号</th>
            <th>操作人</th>
            <th>字段</th>
            <th>旧值</th>
            <th>新值</th>
          </tr>
        </thead>
        <tbody id="logTableBody">
          <tr><td colspan="7" class="text-center text-muted py-4">加载中...</td></tr>
        </tbody>
      </table>
    </div>
    <div id="logPagination" class="d-flex justify-content-between align-items-center p-3 border-top">
      <span id="logPageInfo" class="text-muted small"></span>
      <div id="logPageBtns" class="btn-group btn-group-sm"></div>
    </div>
  </div>
</div>

<script>
var FIELD_NAMES = {
  'serial_no':'序号','liquidated_damages':'违约金','help_letter':'求援函项',
  'first_article':'首件鉴定','demand_no':'需求单编号','responsible_person':'负责人',
  'product_category':'品种','stage':'阶段',
  'product_model':'产品型号','product_draw_no':'产品图号','product_name':'产品名称',
  'total_qty':'总数量','plan_deliver_qty':'计划交付数量','fixed_check_qty':'定检数量',
  'batch_no':'批次号','check_party':'验收方','plan_attribute':'计划属性',
  'contract_no':'合同编号','plan_issue_time':'主计划下达时间','plan_delivery_time':'主计划要求交付时间',
  'specific_model':'具体号型','plan_source':'计划来源','contract_no2':'合同编号',
  'demander':'需求人','customer_name':'客户名称','project_no':'机型/项目流程编号',
  'unit_price':'单价(万元)','output_value':'产值','production_time':'出产时间',
  'matching_time':'配套日期','cut_start':'裁剪开始','cut_end':'裁剪结束',
  'sew_start':'缝纫开始','sew_end':'缝纫结束','glue_start':'粘胶开始','glue_end':'粘胶结束',
  'assembly1_start':'总装1开始','assembly1_end':'总装1结束',
  'assembly2_start':'总装2开始','assembly2_end':'总装2结束',
  'oxygen_start':'氧调开始','oxygen_end':'氧调结束',
  'heat_seal_start':'热风热合开始','heat_seal_end':'热风热合结束',
  'estimated_finish_time':'预计提交总检时间','production_status':'生产情况',
  'final_check_time':'交总检日期','fixed_check_deliver_time':'送定检日期',
  'fixed_check_finish_time':'定检完成日期','military_check_plan_time':'报军检计划时间',
  'military_check_time':'军检时间','fixed_submit_item':'固定提交项',
  'fixed_submit_finish_time':'固定提交完成日期','storage_time':'入库时间','remark':'备注'
};

function loadLogs(page) {
  var params = 'page=' + page + '&per_page=30';
  var opType = document.getElementById('filterOpType').value;
  var start = document.getElementById('filterStartDate').value;
  var end = document.getElementById('filterEndDate').value;
  if (opType) params += '&operation_type=' + encodeURIComponent(opType);
  if (start) params += '&start_date=' + encodeURIComponent(start);
  if (end) params += '&end_date=' + encodeURIComponent(end);

  var tbody = document.getElementById('logTableBody');
  tbody.innerHTML = '<tr><td colspan="7" class="text-center text-muted py-4">加载中...</td></tr>';

  fetch('/api/logs?' + params)
    .then(function(r) { return r.json(); })
    .then(function(data) {
      if (data.logs.length === 0) {
        tbody.innerHTML = '<tr><td colspan="7" class="text-center text-muted py-4">暂无操作记录</td></tr>';
      } else {
        var html = '';
        data.logs.forEach(function(log) {
          html += '<tr>';
          html += '<td class="small">' + (log.operated_time || '') + '</td>';
          html += '<td><span class="badge bg-' + typeBadge(log.operation_type) + '">' + (log.operation_type || '') + '</span></td>';
          html += '<td><a href="/task/' + log.task_id + '/details" class="small">' + (log.new_value || '').replace('序号','').replace('删除序号','') + '</a></td>';
          html += '<td class="small">' + (log.operator || '') + '</td>';
          html += '<td class="small">' + (FIELD_NAMES[log.field_name] || log.field_name || '') + '</td>';
          html += '<td class="small text-muted text-truncate">' + (log.old_value || '') + '</td>';
          html += '<td class="small text-truncate">' + (log.new_value || '') + '</td>';
          html += '</tr>';
        });
        tbody.innerHTML = html;
      }

      var info = document.getElementById('logPageInfo');
      info.textContent = '共 ' + data.total + ' 条记录，第 ' + data.page + '/' + data.total_pages + ' 页';

      var btns = document.getElementById('logPageBtns');
      var btnHtml = '';
      btnHtml += '<button class="btn btn-outline-secondary btn-sm" onclick="loadLogs(1)" ' + (data.page <= 1 ? 'disabled' : '') + '>首页</button>';
      btnHtml += '<button class="btn btn-outline-secondary btn-sm" onclick="loadLogs(' + (data.page - 1) + ')" ' + (data.page <= 1 ? 'disabled' : '') + '>上一页</button>';
      btnHtml += '<button class="btn btn-outline-secondary btn-sm" onclick="loadLogs(' + (data.page + 1) + ')" ' + (data.page >= data.total_pages ? 'disabled' : '') + '>下一页</button>';
      btnHtml += '<button class="btn btn-outline-secondary btn-sm" onclick="loadLogs(' + data.total_pages + ')" ' + (data.page >= data.total_pages ? 'disabled' : '') + '>末页</button>';
      btns.innerHTML = btnHtml;
    })
    .catch(function() {
      tbody.innerHTML = '<tr><td colspan="7" class="text-center text-danger py-4">加载失败，请重试</td></tr>';
    });
}

function typeBadge(type) {
  var map = {'创建任务':'primary','批量新增':'success','批量更新':'info','删除任务':'danger','编辑字段':'warning'};
  return map[type] || 'secondary';
}

function resetLogFilters() {
  document.getElementById('filterOpType').value = '';
  document.getElementById('filterStartDate').value = '';
  document.getElementById('filterEndDate').value = '';
  loadLogs(1);
}

loadLogs(1);
</script>
{% endblock %}''')

PLANNING_REPORT_HTML = BASE_HTML.replace('{% block content %}{% endblock %}','''
{% block content %}
<div class="container-fluid" style="max-width:1200px;">
  <h4 class="mb-3 fw-bold text-primary">📋 计划导入生产任务</h4>

  <div class="card shadow-sm mb-3">
    <div class="card-header bg-light d-flex justify-content-between align-items-center">
      <span class="fw-bold">📥 批量粘贴导入（从Excel复制数据，含表头）</span>
      <button type="button" class="btn btn-outline-secondary btn-sm" onclick="toggleHelp()">📖 列名对照</button>
    </div>
    <div class="card-body">
      <div id="helpPanel" class="alert alert-secondary mb-2 py-2 px-3" style="display:none; font-size:12px; max-height:200px; overflow-y:auto;">
        <button type="button" class="btn-close float-end" onclick="document.getElementById('helpPanel').style.display='none'"></button>
        <strong>支持列名（表头需用Tab分隔，从Excel复制即符合）：</strong><br>
        <span class="text-muted">
        序号 负责人 品种 阶段 产品型号 产品图号 产品名称 总数量 计划交付数量 定检数量 批次号
        违约金 求援函项 首件鉴定 需求单编号 验收方 计划属性 合同编号 主计划下达时间
        主计划要求交付时间 具体号型 计划来源 需求人 客户名称 机型/项目流程编号
        单价(万元) 产值 出产时间 配套日期 备注
        </span>
        <br><small class="text-warning">日期格式: 2026/6/22 或 2026-06-22；违约金/求援函项/首件鉴定: 是/否</small>
      </div>
      <textarea id="batchText" class="form-control" rows="14" style="font-family: Consolas,monospace; font-size:13px;" placeholder="从Excel复制数据（含表头）粘贴到此处...&#10;表头和数据列之间用Tab分隔"></textarea>
      <div class="d-flex gap-2 mt-3">
        <button id="batchSubmitBtn" class="btn btn-primary px-4">✅ 导入</button>
        <button type="button" class="btn btn-outline-danger px-4" onclick="document.getElementById('batchText').value=''">🗑 清空</button>
        <span id="submitMsg" class="align-self-center ms-3"></span>
      </div>
    </div>
  </div>

  <!-- 最近填报 -->
  <div class="card shadow-sm">
    <div class="card-header bg-light fw-bold">📝 最近填报记录</div>
    <div class="table-responsive" style="max-height:400px;">
      <table class="table table-sm table-hover mb-0">
        <thead class="table-light"><tr>
          <th>序号</th><th>负责人</th><th>品种</th><th>阶段</th><th>产品图号</th><th>产品名称</th><th>批次号</th><th>总数量</th><th>状态</th><th>填报时间</th>
        </tr></thead>
        <tbody>
        {% for t in recent %}
        <tr>
          <td><a href="/task/{{ t.id }}/details" class="text-decoration-none fw-bold">{{ t.serial_no }}</a></td>
          <td>{{ t.responsible_person or '' }}</td>
          <td>{{ t.product_category or '' }}</td>
          <td>{{ t.stage or '' }}</td>
          <td>{{ t.product_draw_no or '' }}</td>
          <td>{{ t.product_name or '' }}</td>
          <td>{{ t.batch_no or '' }}</td>
          <td>{{ t.total_qty or '' }}</td>
          <td>{% if t.production_time %}已出产{% elif t.cut_start %}生产中{% else %}待开工{% endif %}</td>
          <td>{{ fmt_date(t.created_time) }}</td>
        </tr>
        {% else %}
        <tr><td colspan="10" class="text-center text-muted py-3">暂无填报记录</td></tr>
        {% endfor %}
        </tbody>
      </table>
    </div>
  </div>
</div>

<script>
function toggleHelp() {
  var p = document.getElementById('helpPanel');
  p.style.display = p.style.display === 'none' ? 'block' : 'none';
}

document.getElementById('batchSubmitBtn').addEventListener('click', async () => {
    var text = document.getElementById('batchText').value;
    var msgEl = document.getElementById('submitMsg');
    if (!text.trim()) {
      msgEl.innerHTML = '<span class="text-warning">⚠ 请先粘贴数据</span>';
      return;
    }
    msgEl.innerHTML = '<span class="text-info">提交中...</span>';
    try {
      var res = await fetch('/api/task/batch_add', { method: 'POST', headers: {'Content-Type': 'application/json'}, body: JSON.stringify({text}) });
      var data = await res.json();
      if (data.success) {
        msgEl.innerHTML = '<span class="text-success fw-bold">✅ ' + data.message + '</span>';
        setTimeout(function() { location.reload(); }, 1500);
      } else {
        msgEl.innerHTML = '<span class="text-danger">❌ ' + (data.error || '导入失败') + '</span>';
      }
    } catch(e) {
      msgEl.innerHTML = '<span class="text-danger">❌ 网络错误，请重试</span>';
    }
});
</script>
{% endblock %}''')

PLANNING_DASHBOARD_HTML = BASE_HTML.replace('{% block content %}{% endblock %}','''
{% block content %}
<div class="container-fluid">
  <h4 class="mb-3 fw-bold text-primary">📊 计调模块</h4>

  <!-- 概览卡片 -->
  <div class="row g-3 mb-3">
    <div class="col-md-2"><div class="card card-metric"><div class="card-body"><div><div class="metric-label">总任务</div><div class="metric-value">{{ total }}</div></div><div class="metric-icon-bg">📋</div></div></div></div>
    <div class="col-md-2"><div class="card card-metric" style="border-left: 4px solid #dc3545;"><div class="card-body"><div><div class="metric-label">未开工</div><div class="metric-value text-danger">{{ not_started }}</div></div><div class="metric-icon-bg">⏸</div></div></div></div>
    <div class="col-md-2"><div class="card card-metric" style="border-left: 4px solid #ffc107;"><div class="card-body"><div><div class="metric-label">生产中</div><div class="metric-value text-warning">{{ in_progress }}</div></div><div class="metric-icon-bg">⚙</div></div></div></div>
    <div class="col-md-2"><div class="card card-metric" style="border-left: 4px solid #198754;"><div class="card-body"><div><div class="metric-label">已完成</div><div class="metric-value text-success">{{ completed }}</div></div><div class="metric-icon-bg">✅</div></div></div></div>
    <div class="col-md-2"><div class="card card-metric" style="border-left: 4px solid #dc3545;"><div class="card-body"><div><div class="metric-label">逾期</div><div class="metric-value text-danger">{{ overdue }}</div></div><div class="metric-icon-bg">⚠</div></div></div></div>
  </div>

   <div class="row g-3">
    <!-- 班组任务概览 -->
    <div class="col-md-3">
      <div class="card shadow-sm h-100">
        <div class="card-header bg-light fw-bold">班组任务概览</div>
        <div class="card-body p-0" style="max-height:220px;overflow-y:auto;">
          <table class="table table-sm table-hover mb-0">
            <thead class="table-light"><tr><th>工序</th><th class="text-center">待开</th><th class="text-center">进行</th><th class="text-center">完成</th><th class="text-center">逾期</th></tr></thead>
            <tbody>
            {% for team in teams %}
            {% set s = team_stats[team] %}
            <tr>
              <td class="fw-bold">{{ team }}</td>
              <td class="text-center">{{ s.todo }}</td>
              <td class="text-center text-warning fw-bold">{{ s.doing }}</td>
              <td class="text-center text-success">{{ s.done }}</td>
              <td class="text-center">{% if s.overdue %}<span class="badge bg-danger">{{ s.overdue }}</span>{% else %}-{% endif %}</td>
            </tr>
            {% endfor %}
            </tbody>
          </table>
        </div>
      </div>
    </div>

    <!-- 班组动态 -->
    <div class="col-md-3">
      <div class="card shadow-sm h-100">
        <div class="card-header bg-light fw-bold">🔧 班组动态</div>
         <div style="max-height:220px;overflow-y:auto;">
          <table class="table table-sm table-hover mb-0">
            <thead class="table-light"><tr><th>时间</th><th>动态</th></tr></thead>
            <tbody>
            {% for d in team_dynamics %}
            <tr>
              <td class="small text-muted text-nowrap">{{ fmt_date(d.operated_time)[5:] }}</td>
              <td class="small">{{ d.description }}</td>
            </tr>
            {% else %}
            <tr><td colspan="2" class="text-center text-muted py-3">暂无动态</td></tr>
            {% endfor %}
            </tbody>
          </table>
        </div>
      </div>
    </div>

    <!-- 逾期 -->
    <div class="col-md-3">
      <div class="card shadow-sm h-100">
        <div class="card-header bg-light fw-bold">⚠ 逾期 ({{ overdue_detail|length }})</div>
        <div style="max-height:220px;overflow-y:auto;">
          <table class="table table-sm table-hover mb-0">
            <thead class="table-light"><tr><th>序号</th><th>图号/批次</th><th>逾期</th><th>状态</th></tr></thead>
            <tbody>
            {% for t in overdue_detail %}
            <tr>
              <td><a href="/task/{{ t.id }}/details" class="text-decoration-none small fw-bold">{{ t.serial_no }}</a></td>
              <td class="small">{{ t.product_draw_no }}{% if t.batch_no %}/{{ t.batch_no }}{% endif %}</td>
              <td><span class="badge bg-danger">{{ t.overdue_days }}天</span></td>
              <td><span class="badge {% if t.production_status == '未开工' %}bg-secondary{% else %}bg-warning text-dark{% endif %}">{{ t.production_status[:2] }}</span></td>
            </tr>
            {% else %}
            <tr><td colspan="4" class="text-center text-muted py-3">✅ 无逾期</td></tr>
            {% endfor %}
            </tbody>
          </table>
        </div>
      </div>
    </div>

    <!-- 通知公告 -->
    <div class="col-md-3">
      <div class="card shadow-sm h-100">
        <div class="card-header bg-light fw-bold">📢 通知公告</div>
        <div class="card-body p-2" style="max-height:220px;overflow-y:auto;">
          {% for n in branch_notices %}
          <div class="border-bottom py-1 {% if not loop.last %}mb-1{% endif %}">
            <span class="fw-bold small">{{ n.title }}</span>
            <small class="text-muted d-block">{{ fmt_date(n.created_time) }}</small>
            <div class="small text-muted" style="white-space:pre-wrap;">{{ n.content }}</div>
          </div>
          {% else %}
          <div class="text-center text-muted py-2 small">暂无通知</div>
          {% endfor %}
        </div>
      </div>
    </div>
  </div>

  <!-- 问题台账 -->
  <div class="row g-3 mt-3">
    <div class="col-md-6">
      <div class="card shadow-sm">
        <div class="card-header bg-light fw-bold d-flex justify-content-between align-items-center">
          <span>📝 问题台账 ({{ issue_ledger|length }})</span>
          <button class="btn btn-sm btn-outline-primary py-0 px-2" data-bs-toggle="modal" data-bs-target="#issueModal">+ 添加</button>
        </div>
        <div style="max-height:400px;overflow-y:auto;">
          <table class="table table-sm table-hover mb-0">
            <thead class="table-light"><tr><th>序号</th><th>图号/批次</th><th>类型</th><th>部门</th><th>内容</th><th>提出</th><th>天数</th></tr></thead>
            <tbody>
            {% for i in issue_ledger %}
            <tr>
              <td><a href="/task/{{ i.task_id }}/details" class="text-decoration-none small fw-bold">{{ i.serial_no }}</a></td>
              <td class="small">{{ i.product_draw_no }}{% if i.batch_no %}/{{ i.batch_no }}{% endif %}</td>
              <td><span class="badge {% if i.issue_type == '技术' %}bg-primary{% else %}bg-warning text-dark{% endif %}">{{ i.issue_type }}</span></td>
              <td class="small">{{ i.dept }}</td>
              <td class="small text-muted">{{ i.content }}</td>
              <td class="small">{{ fmt_date(i.raise_time) }}</td>
              <td><span class="badge bg-{% if i.days > 7 %}danger{% elif i.days > 3 %}warning text-dark{% else %}success{% endif %}">{{ i.days }}天</span></td>
            </tr>
            {% else %}
            <tr><td colspan="7" class="text-center text-muted py-3">✅ 无问题记录</td></tr>
            {% endfor %}
            </tbody>
          </table>
        </div>
      </div>
    </div>
    <div class="col-md-6">
      <div class="card shadow-sm">
        <div class="card-header bg-light fw-bold d-flex justify-content-between align-items-center">
          <span>📦 缺件台账 ({{ shortage_ledger|length }})</span>
          <button class="btn btn-sm btn-outline-primary py-0 px-2" data-bs-toggle="modal" data-bs-target="#shortageModal">+ 添加</button>
        </div>
        <div style="max-height:400px;overflow-y:auto;">
          <table class="table table-sm table-hover mb-0">
            <thead class="table-light"><tr><th>序号</th><th>图号/批次</th><th>类型</th><th>内容</th><th>报缺</th><th>天数</th></tr></thead>
            <tbody>
            {% for s in shortage_ledger %}
            <tr>
              <td><a href="/task/{{ s.task_id }}/details" class="text-decoration-none small fw-bold">{{ s.serial_no }}</a></td>
              <td class="small">{{ s.product_draw_no }}{% if s.batch_no %}/{{ s.batch_no }}{% endif %}</td>
              <td><span class="badge bg-secondary">{{ s.shortage_type }}</span></td>
              <td class="small text-muted">{{ s.content }}</td>
              <td class="small">{{ fmt_date(s.report_time) }}</td>
              <td><span class="badge bg-{% if s.days > 7 %}danger{% elif s.days > 3 %}warning text-dark{% else %}success{% endif %}">{{ s.days }}天</span></td>
            </tr>
            {% else %}
            <tr><td colspan="6" class="text-center text-muted py-3">✅ 无缺件记录</td></tr>
            {% endfor %}
            </tbody>
          </table>
        </div>
      </div>
    </div>
  </div>
</div>

<!-- 添加问题弹窗 -->
<div class="modal fade" id="issueModal" tabindex="-1">
  <div class="modal-dialog">
    <div class="modal-content">
      <div class="modal-header"><h5 class="modal-title">添加问题</h5><button type="button" class="btn-close" data-bs-dismiss="modal"></button></div>
      <div class="modal-body">
        <div class="mb-2">
          <label class="form-label small mb-0">任务</label>
          <select class="form-select form-select-sm" id="issueTask">
            <option value="">请选择任务</option>
            {% for t in all_tasks %}<option value="{{ t.id }}">{{ t.serial_no }} {{ t.product_draw_no or '' }}/{{ t.batch_no or '' }}</option>{% endfor %}
          </select>
        </div>
        <div class="mb-2"><label class="form-label small mb-0">类型</label>
          <select class="form-select form-select-sm" id="issueType"><option value="tech">技术</option><option value="mgmt">管理</option></select></div>
        <div class="mb-2"><label class="form-label small mb-0">部门</label>
          <select class="form-select form-select-sm" id="issueDept"><option value="">--选填--</option>{% for d in dept_options %}<option value="{{ d }}">{{ d }}</option>{% endfor %}</select></div>
        <div class="mb-2"><label class="form-label small mb-0">内容</label><textarea class="form-control form-control-sm" id="issueContent" rows="2"></textarea></div>
        <div class="mb-2"><label class="form-label small mb-0">提出日期</label><input type="date" class="form-control form-control-sm" id="issueDate"></div>
      </div>
      <div class="modal-footer"><button class="btn btn-secondary btn-sm" data-bs-dismiss="modal">取消</button><button class="btn btn-primary btn-sm" id="issueSubmitBtn">提交</button></div>
    </div>
  </div>
</div>

<!-- 添加缺件弹窗 -->
<div class="modal fade" id="shortageModal" tabindex="-1">
  <div class="modal-dialog">
    <div class="modal-content">
      <div class="modal-header"><h5 class="modal-title">添加缺件</h5><button type="button" class="btn-close" data-bs-dismiss="modal"></button></div>
      <div class="modal-body">
        <div class="mb-2">
          <label class="form-label small mb-0">任务</label>
          <select class="form-select form-select-sm" id="shortageTask">
            <option value="">请选择任务</option>
            {% for t in all_tasks %}<option value="{{ t.id }}">{{ t.serial_no }} {{ t.product_draw_no or '' }}/{{ t.batch_no or '' }}</option>{% endfor %}
          </select>
        </div>
        <div class="mb-2"><label class="form-label small mb-0">缺件类型</label>
          <select class="form-select form-select-sm" id="shortageType">{% for st in shortage_types %}<option value="{{ st }}">{{ st }}</option>{% endfor %}</select></div>
        <div class="mb-2"><label class="form-label small mb-0">内容</label><textarea class="form-control form-control-sm" id="shortageContent" rows="2"></textarea></div>
        <div class="mb-2"><label class="form-label small mb-0">报缺日期</label><input type="date" class="form-control form-control-sm" id="shortageDate"></div>
      </div>
      <div class="modal-footer"><button class="btn btn-secondary btn-sm" data-bs-dismiss="modal">取消</button><button class="btn btn-primary btn-sm" id="shortageSubmitBtn">提交</button></div>
    </div>
  </div>
</div>


</script>
<script>
// 添加问题
document.getElementById('issueSubmitBtn').addEventListener('click', function() {
  var taskId = document.getElementById('issueTask').value;
  if (!taskId) { alert('请选择任务'); return; }
  var content = document.getElementById('issueContent').value.trim();
  if (!content) { alert('请填写内容'); return; }
  var data = {
    task_id: parseInt(taskId),
    issue_type: document.getElementById('issueType').value,
    dept: document.getElementById('issueDept').value,
    content: content,
    raise_time: document.getElementById('issueDate').value
  };
  fetch('/api/issue/add', { method:'POST', headers:{'Content-Type':'application/json'}, body:JSON.stringify(data) })
    .then(function(r){return r.json()})
    .then(function(d){ if(d.success) location.reload(); else alert(d.error||'添加失败'); });
});

// 添加缺件
document.getElementById('shortageSubmitBtn').addEventListener('click', function() {
  var taskId = document.getElementById('shortageTask').value;
  if (!taskId) { alert('请选择任务'); return; }
  var content = document.getElementById('shortageContent').value.trim();
  if (!content) { alert('请填写内容'); return; }
  var data = {
    task_id: parseInt(taskId),
    shortage_type: document.getElementById('shortageType').value,
    content: content,
    report_time: document.getElementById('shortageDate').value
  };
  fetch('/api/shortage/add', { method:'POST', headers:{'Content-Type':'application/json'}, body:JSON.stringify(data) })
    .then(function(r){return r.json()})
    .then(function(d){ if(d.success) location.reload(); else alert(d.error||'添加失败'); });
});
</script>
{% endblock %}''')

FEEDBACK_HTML = BASE_HTML.replace('{% block content %}{% endblock %}','''
{% block content %}
<div class="container-fluid" style="max-width:1000px;">
  <h4 class="mb-3 fw-bold text-primary">💬 修改意见</h4>

  <!-- 提交意见 -->
  <div class="card shadow-sm mb-4">
    <div class="card-header bg-light fw-bold">📝 提交修改意见</div>
    <div class="card-body">
      <div class="row g-2 mb-2">
        <div class="col-md-8">
          <input type="text" id="fbTitle" class="form-control form-control-sm" placeholder="意见标题" maxlength="200">
        </div>
        <div class="col-md-4">
          <select id="fbCategory" class="form-select form-select-sm">
            <option value="功能建议">功能建议</option>
            <option value="问题反馈">问题反馈</option>
            <option value="其他">其他</option>
          </select>
        </div>
      </div>
      <textarea id="fbContent" class="form-control form-control-sm" rows="3" placeholder="请详细描述您的修改意见..."></textarea>
      <div class="d-flex gap-2 mt-2">
        <button id="fbSubmitBtn" class="btn btn-primary btn-sm px-4">提交</button>
        <span id="fbMsg" class="align-self-center"></span>
      </div>
    </div>
  </div>

  <!-- 意见列表 -->
  <div class="card shadow-sm">
    <div class="card-header bg-light fw-bold">📋 意见列表 ({{ feedbacks|length }})</div>
    <div class="card-body p-0">
    {% for fb in feedbacks %}
      <div class="border-bottom p-3" id="fb-{{ fb.id }}">
        <div class="d-flex justify-content-between align-items-start">
          <div>
            <span class="fw-bold">{{ fb.title }}</span>
            <span class="badge {% if fb.category == '功能建议' %}bg-primary{% elif fb.category == '问题反馈' %}bg-warning text-dark{% else %}bg-secondary{% endif %} ms-2">{{ fb.category }}</span>
            <span class="badge {% if fb.status == '待处理' %}bg-secondary{% elif fb.status == '已采纳' %}bg-success{% else %}bg-danger{% endif %} ms-1">{{ fb.status }}</span>
          </div>
          <small class="text-muted">{{ fb.user.username or fb.user.account }} · {{ fmt_date(fb.created_time) }}</small>
        </div>
        <div class="mt-2 text-muted" style="white-space:pre-wrap;">{{ fb.content }}</div>
        {% if fb.reply %}
        <div class="mt-2 p-2 bg-light rounded" style="border-left:3px solid #198754;">
          <div class="d-flex justify-content-between"><small class="fw-bold text-success">🔔 管理员回复</small><small class="text-muted">{{ fmt_date(fb.replied_time) }}</small></div>
          <div class="small" style="white-space:pre-wrap;">{{ fb.reply }}</div>
        </div>
        {% endif %}
        {% if current_user.role.name == '管理员' %}
        <div class="mt-2 admin-reply-area" id="replyArea-{{ fb.id }}" style="display:none;">
          <select class="form-select form-select-sm mb-1 reply-status" style="width:120px;">
            <option value="已采纳" {% if fb.status == '已采纳' %}selected{% endif %}>已采纳</option>
            <option value="已回复" {% if fb.status == '已回复' %}selected{% endif %}>已回复</option>
            <option value="已拒绝" {% if fb.status == '已拒绝' %}selected{% endif %}>已拒绝</option>
          </select>
          <textarea class="form-control form-control-sm reply-text" rows="2" placeholder="回复内容...">{{ fb.reply or '' }}</textarea>
          <button class="btn btn-sm btn-success mt-1 reply-save-btn" data-id="{{ fb.id }}">保存回复</button>
        </div>
        <button class="btn btn-sm btn-outline-secondary mt-1 reply-toggle-btn" data-id="{{ fb.id }}">{% if fb.reply %}编辑回复{% else %}回复{% endif %}</button>
        {% endif %}
      </div>
    {% else %}
      <div class="text-center text-muted py-4">暂无修改意见，欢迎提交</div>
    {% endfor %}
    </div>
  </div>
</div>

<script>
document.getElementById('fbSubmitBtn').addEventListener('click', function() {
  var title = document.getElementById('fbTitle').value.trim();
  var content = document.getElementById('fbContent').value.trim();
  var category = document.getElementById('fbCategory').value;
  var msg = document.getElementById('fbMsg');
  if (!title || !content) { msg.innerHTML = '<span class="text-warning">请填写标题和内容</span>'; return; }
  msg.innerHTML = '<span class="text-info">提交中...</span>';
  fetch('/api/feedback/submit', {
    method: 'POST', headers: {'Content-Type': 'application/json'},
    body: JSON.stringify({title:title, content:content, category:category})
  })
  .then(function(r) { return r.json(); })
  .then(function(data) {
    if (data.success) {
      msg.innerHTML = '<span class="text-success">✅ 提交成功</span>';
      setTimeout(function() { location.reload(); }, 1000);
    } else {
      msg.innerHTML = '<span class="text-danger">❌ ' + (data.error || '提交失败') + '</span>';
    }
  })
  .catch(function() { msg.innerHTML = '<span class="text-danger">网络错误</span>'; });
});

// Admin reply toggle
document.querySelectorAll('.reply-toggle-btn').forEach(function(btn) {
  btn.addEventListener('click', function() {
    var id = this.getAttribute('data-id');
    var area = document.getElementById('replyArea-' + id);
    area.style.display = area.style.display === 'none' ? 'block' : 'none';
  });
});

// Admin reply save
document.querySelectorAll('.reply-save-btn').forEach(function(btn) {
  btn.addEventListener('click', function() {
    var id = this.getAttribute('data-id');
    var area = document.getElementById('replyArea-' + id);
    var reply = area.querySelector('.reply-text').value.trim();
    var status = area.querySelector('.reply-status').value;
    if (!reply) { alert('请填写回复内容'); return; }
    fetch('/api/feedback/reply/' + id, {
      method: 'POST', headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({reply:reply, status:status})
    })
    .then(function(r) { return r.json(); })
    .then(function(data) {
      if (data.success) location.reload();
      else alert(data.error || '保存失败');
    });
  });
});
</script>
{% endblock %}''')

TASK_REPORT_HTML = BASE_HTML.replace('{% block content %}{% endblock %}','''
{% block content %}
<div class="container-fluid">
  <h4 class="mb-3 fw-bold text-primary">📋 计调任务填报</h4>
  <div class="d-flex mb-2">
    <div class="btn-group-actions">
      <button id="reportClearFilterBtn" class="btn">✕ 取消筛选</button>
      <button id="reportColSettingsBtn" class="btn">⚙ 列设置</button>
    </div>
  </div>

  <!-- 列设置弹窗 -->
  <div class="modal fade" id="reportColModal" tabindex="-1">
    <div class="modal-dialog modal-sm">
      <div class="modal-content">
        <div class="modal-header"><h5 class="modal-title">列设置</h5><button type="button" class="btn-close" data-bs-dismiss="modal"></button></div>
        <div class="modal-body">
          <div class="d-flex gap-1 mb-2">
            <button type="button" class="btn btn-outline-secondary btn-sm" id="rptSelectAll">全选</button>
            <button type="button" class="btn btn-outline-secondary btn-sm" id="rptSelectNone">全不选</button>
            <button type="button" class="btn btn-outline-secondary btn-sm" id="rptResetDefault">恢复默认</button>
          </div>
          <div id="rptColSortList" class="list-group" style="max-height:400px;overflow-y:auto;"></div>
        </div>
        <div class="modal-footer"><button type="button" class="btn btn-secondary" data-bs-dismiss="modal">取消</button><button type="button" class="btn btn-primary" id="rptColApplyBtn">应用</button></div>
      </div>
    </div>
  </div>

  <div class="table-outer shadow-sm" style="height:calc(100vh - 200px);">
    <div class="table-inner">
      <table class="table table-bordered table-sm" id="reportTable">
        <thead>
          <tr>
            <th data-field="serial_no">序号</th>
            <th data-field="product_category">品种</th>
            <th data-field="product_draw_no">产品图号</th>
            <th data-field="product_name">产品名称</th>
            <th data-field="batch_no">批次号</th>
            <th data-field="matching_time">配套日期</th>
            <th data-field="fixed_check_deliver_time">送定检日期</th>
            <th data-field="fixed_check_finish_time">定检完成日期</th>
            <th data-field="military_check_plan_time">报军检计划日期</th>
            <th data-field="military_check_time">军检时间</th>
            <th data-field="fixed_submit_finish_time">固定提交完成日期</th>
            <th data-field="storage_time">入库时间</th>
            <th>操作</th>
          </tr>
        </thead>
        <tbody>
        {% for t in tasks %}
        <tr data-task-id="{{ t.id }}">
          <td class="seq" data-field="serial_no">{{ t.serial_no }}</td>
          <td data-field="product_category">{{ t.product_category }}</td>
          <td data-field="product_draw_no">{{ t.product_draw_no }}</td>
          <td data-field="product_name">{{ t.product_name }}</td>
          <td data-field="batch_no">{{ t.batch_no }}</td>
          <td data-field="matching_time" data-id="{{ t.id }}" ondblclick="makeDateEditable(this)" class="can-edit">{{ fmt_date(t.matching_time) }}</td>
          <td data-field="fixed_check_deliver_time" data-id="{{ t.id }}" ondblclick="makeDateEditable(this)" class="can-edit">{{ fmt_date(t.fixed_check_deliver_time) }}</td>
          <td data-field="fixed_check_finish_time" data-id="{{ t.id }}" ondblclick="makeDateEditable(this)" class="can-edit">{{ fmt_date(t.fixed_check_finish_time) }}</td>
          <td data-field="military_check_plan_time" data-id="{{ t.id }}" ondblclick="makeDateEditable(this)" class="can-edit">{{ fmt_date(t.military_check_plan_time) }}</td>
          <td data-field="military_check_time" data-id="{{ t.id }}" ondblclick="makeDateEditable(this)" class="can-edit">{{ fmt_date(t.military_check_time) }}</td>
          <td data-field="fixed_submit_finish_time" data-id="{{ t.id }}" ondblclick="makeDateEditable(this)" class="can-edit">{{ fmt_date(t.fixed_submit_finish_time) }}</td>
          <td data-field="storage_time" data-id="{{ t.id }}" ondblclick="makeDateEditable(this)" class="can-edit">{{ fmt_date(t.storage_time) }}</td>
          <td class="text-center"><a href="/task/{{ t.id }}/details" class="btn btn-outline-secondary btn-sm py-0">详情</a></td>
        </tr>
        {% endfor %}
        </tbody>
      </table>
    </div>
  </div>
</div>

<script>
// ===== 日期编辑 =====
function makeDateEditable(td) {
  if (td.querySelector("input")) return;
  var oldText = td.innerText;
  var input = document.createElement("input");
  input.type = "date";
  input.className = "form-control form-control-sm";
  input.style.width = "130px"; input.style.padding = "2px 4px"; input.style.fontSize = "14px";
  if (oldText) {
    var parts = oldText.split("/");
    if (parts.length === 3) {
      input.value = parts[0] + "-" + parts[1].padStart(2,"0") + "-" + parts[2].padStart(2,"0");
    }
  }
  td.innerText = ""; td.appendChild(input); input.focus();
  function save() {
    var newVal = input.value; input.remove(); td.innerText = oldText;
    if (newVal === (oldText ? oldText.replace(/\//g,"-") : "")) return;
    td.innerText = "保存中...";
    var taskId = td.getAttribute("data-id");
    var field = td.getAttribute("data-field");
    fetch("/api/task/" + taskId + "/field", {
      method: "PUT", headers: {"Content-Type": "application/json"},
      body: JSON.stringify({field: field, value: newVal})
    })
    .then(function(r) { return r.json(); })
    .then(function(data) {
      if (data.success) {
        var resultDate = data.new_value;
        if (resultDate && resultDate.indexOf("-") > -1) {
          var parts2 = resultDate.split("-");
          resultDate = parts2[0] + "/" + parts2[1] + "/" + parts2[2];
        }
        td.innerText = resultDate || "";
      } else { td.innerText = oldText; alert(data.error || "保存失败"); }
    })
    .catch(function() { td.innerText = oldText; });
  }
  input.addEventListener("blur", save);
  input.addEventListener("keypress", function(e) { if (e.key === "Enter") save(); });
}

// ===== 表头筛选 =====
(function() {
  var table = document.getElementById("reportTable");
  var rows = Array.from(table.querySelectorAll("tbody tr"));
  window._reportFilters = {};

  table.querySelectorAll("thead th").forEach(function(th, colIndex) {
    var icon = document.createElement("span");
    icon.innerHTML = " 🔽";
    icon.style.cursor = "pointer"; icon.style.fontSize = "10px";
    icon.onclick = function(e) {
      e.stopPropagation();
      showRptFilter(th, colIndex, e.pageX, e.pageY);
    };
    th.appendChild(icon);
  });

  function showRptFilter(th, colIndex, pageX, pageY) {
    var existing = document.getElementById("rptFilterMenu");
    if (existing) existing.remove();

    var uniqueValues = new Set();
    rows.forEach(function(r) {
      var td = r.cells[colIndex];
      if (td) uniqueValues.add(td.innerText.trim());
    });
    var sortedValues = Array.from(uniqueValues).sort();

    var menu = document.createElement("div");
    menu.id = "rptFilterMenu";
    Object.assign(menu.style, {
      position: "absolute", left: pageX+"px", top: pageY+"px", background: "#fff",
      border: "1px solid #ccc", padding: "10px", zIndex: "9999",
      maxHeight: "400px", overflowY: "auto", boxShadow: "0 4px 8px rgba(0,0,0,0.1)", minWidth: "180px"
    });
    document.body.appendChild(menu);

    var cur = window._reportFilters[colIndex] || { text: "", values: uniqueValues };
    var html = '<div><input type="text" id="rptFilterText" class="form-control form-control-sm" placeholder="输入关键词筛选..." value="' + (cur.text||"").replace(/"/g,"&quot;") + '" style="margin-bottom:6px;"></div>';
    html += '<div><label><input type="checkbox" id="rptSelectAll"> <strong>(全选)</strong></label></div><hr style="margin:5px 0;">';
    sortedValues.forEach(function(v) {
      html += '<div><label><input type="checkbox" class="rpt-filter-val" value="' + v.replace(/"/g,"&quot;") + '" ' + (cur.values.has(v)?"checked":"") + '> ' + (v||"(空白)") + '</label></div>';
    });
    html += '<hr style="margin:5px 0;"><button class="btn btn-sm btn-primary w-100" id="rptApplyFilter">确定</button>';
    html += '<button class="btn btn-sm btn-outline-secondary w-100 mt-1" id="rptClearOne">清除筛选</button>';
    menu.innerHTML = html;

    var selAll = menu.querySelector("#rptSelectAll");
    var cbs = menu.querySelectorAll(".rpt-filter-val");
    var txtIn = menu.querySelector("#rptFilterText");
    selAll.checked = Array.from(cbs).every(function(c) { return c.checked; });

    selAll.onchange = function() { cbs.forEach(function(cb) { cb.checked = selAll.checked; }); };
    cbs.forEach(function(cb) {
      cb.onchange = function() {
        if (!cb.checked) selAll.checked = false;
        else { var ac = Array.from(cbs).every(function(c) { return c.checked; }); if (ac) selAll.checked = true; }
      };
    });
    txtIn.addEventListener("input", function() {
      var q = txtIn.value.toLowerCase();
      cbs.forEach(function(cb) {
        var row = cb.parentNode.parentNode;
        row.style.display = (cb.value||"(空白)").toLowerCase().includes(q) ? "" : "none";
      });
    });

    menu.querySelector("#rptApplyFilter").onclick = function() {
      var sel = new Set(); cbs.forEach(function(cb) { if (cb.checked) sel.add(cb.value); });
      var t = txtIn.value.trim();
      if (sel.size === uniqueValues.size && !t) delete window._reportFilters[colIndex];
      else window._reportFilters[colIndex] = { text: t, values: sel };
      applyRptFilters(); menu.remove();
    };
    menu.querySelector("#rptClearOne").onclick = function() {
      delete window._reportFilters[colIndex]; applyRptFilters(); menu.remove();
    };
    setTimeout(function() {
      txtIn.focus();
      document.addEventListener("click", function closeM(e) {
        if (!menu.contains(e.target)) { menu.remove(); document.removeEventListener("click", closeM); }
      });
    }, 10);
  }

  function applyRptFilters() {
    rows.forEach(function(r) {
      var show = true;
      for (var col in window._reportFilters) {
        var td = r.cells[col];
        if (td) {
          var text = td.innerText.trim();
          var f = window._reportFilters[col];
          if (f.text && !text.toLowerCase().includes(f.text.toLowerCase())) { show = false; break; }
          if (f.values && f.values.size > 0 && !f.values.has(text)) { show = false; break; }
        }
      }
      r.style.display = show ? "" : "none";
    });
  }

  document.getElementById("reportClearFilterBtn").addEventListener("click", function() {
    window._reportFilters = {}; applyRptFilters();
  });
})();

// ===== 列设置 =====
(function() {
  var STORAGE_KEY = "reportColSettings";
  var table = document.getElementById("reportTable");

  function getDefaultOrder() {
    var order = [];
    table.querySelectorAll("thead th").forEach(function(th) {
      var f = th.getAttribute("data-field");
      if (f) order.push({ field: f, name: th.innerText.replace(" 🔽","").trim(), visible: true });
    });
    return order;
  }

  function loadSettings() { try { return JSON.parse(localStorage.getItem(STORAGE_KEY)); } catch(e) { return null; } }
  function saveSettings(s) { localStorage.setItem(STORAGE_KEY, JSON.stringify(s)); }

  function buildList(settings) {
    var list = document.getElementById("rptColSortList");
    list.innerHTML = "";
    settings.forEach(function(item) {
      if (!item.field) return;
      var div = document.createElement("div");
      div.className = "list-group-item d-flex align-items-center py-1 px-2";
      div.setAttribute("draggable", "true");
      div.dataset.field = item.field;
      div.innerHTML = '<span class="me-1 text-muted" style="cursor:grab;font-size:12px;">☰</span>' +
        '<input type="checkbox" class="form-check-input me-2 rpt-col-chk" ' + (item.visible !== false ? "checked" : "") + '>' +
        '<span style="font-size:13px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;">' + item.name + '</span>';

      div.addEventListener("dragstart", function(e) { e.dataTransfer.effectAllowed = "move"; div.style.opacity = "0.4"; e.dataTransfer.setData("text/plain", item.field); });
      div.addEventListener("dragend", function() { div.style.opacity = "1"; });
      div.addEventListener("dragover", function(e) { e.preventDefault(); div.classList.add("drag-over"); });
      div.addEventListener("dragleave", function() { div.classList.remove("drag-over"); });
      div.addEventListener("drop", function(e) {
        e.preventDefault(); div.classList.remove("drag-over");
        var fromField = e.dataTransfer.getData("text/plain");
        var fromEl = list.querySelector('[data-field="' + fromField + '"]');
        if (!fromEl || fromEl === div) return;
        var nodes = Array.from(list.children);
        if (nodes.indexOf(fromEl) < nodes.indexOf(div)) list.insertBefore(fromEl, div.nextSibling);
        else list.insertBefore(fromEl, div);
      });
      list.appendChild(div);
    });
  }

  function readSettings() {
    var items = document.getElementById("rptColSortList").children;
    var settings = [];
    for (var i = 0; i < items.length; i++) {
      var el = items[i], chk = el.querySelector(".rpt-col-chk"), span = el.querySelector("span:last-child");
      settings.push({ field: el.dataset.field, name: span ? span.innerText : "", visible: chk.checked });
    }
    return settings;
  }

  function applyColOrder(settings) {
    var theadTr = table.querySelector("thead tr");
    var headerArr = Array.from(table.querySelectorAll("thead th"));
    var rows = table.querySelectorAll("tbody tr");

    var fieldToTh = {}; headerArr.forEach(function(th) { var f = th.getAttribute("data-field"); if (f) fieldToTh[f] = th; });
    var orderedFields = [], visibleSet = {};
    settings.forEach(function(s) { orderedFields.push(s.field); if (s.visible) visibleSet[s.field] = true; });

    headerArr.forEach(function(th) { theadTr.removeChild(th); });
    orderedFields.forEach(function(f) { if (fieldToTh[f]) theadTr.appendChild(fieldToTh[f]); });

    rows.forEach(function(r) {
      var cells = Array.from(r.cells);
      var fieldToCell = {}; cells.forEach(function(td) { var f = td.getAttribute("data-field"); if (f) fieldToCell[f] = td; });
      cells.forEach(function(td) { r.removeChild(td); });
      orderedFields.forEach(function(f) { if (fieldToCell[f]) { r.appendChild(fieldToCell[f]); fieldToCell[f].style.display = visibleSet[f] ? "" : "none"; } });
    });

    var newHeaders = table.querySelectorAll("thead th");
    newHeaders.forEach(function(th) { var f = th.getAttribute("data-field"); if (f) th.style.display = visibleSet[f] ? "" : "none"; });

    saveSettings(settings);
  }

  document.getElementById("reportColSettingsBtn").addEventListener("click", function() {
    var settings = loadSettings();
    if (!settings || settings.length === 0) settings = getDefaultOrder();
    var currentFields = {};
    table.querySelectorAll("thead th").forEach(function(th) { var f = th.getAttribute("data-field"); if (f) currentFields[f] = th.innerText.replace(" 🔽","").trim(); });
    var existing = new Set(settings.map(function(s) { return s.field; }));
    for (var f in currentFields) { if (!existing.has(f)) settings.push({ field: f, name: currentFields[f], visible: true }); }
    settings = settings.filter(function(s) { return currentFields[s.field]; });
    buildList(settings);
    var modal = new bootstrap.Modal(document.getElementById("reportColModal"));
    modal.show();
  });

  document.getElementById("rptSelectAll").addEventListener("click", function() { document.querySelectorAll("#rptColSortList .rpt-col-chk").forEach(function(c) { c.checked = true; }); });
  document.getElementById("rptSelectNone").addEventListener("click", function() { document.querySelectorAll("#rptColSortList .rpt-col-chk").forEach(function(c) { c.checked = false; }); });
  document.getElementById("rptResetDefault").addEventListener("click", function() { buildList(getDefaultOrder()); });
  document.getElementById("rptColApplyBtn").addEventListener("click", function() {
    var settings = readSettings();
    var modal = bootstrap.Modal.getInstance(document.getElementById("reportColModal"));
    applyColOrder(settings);
    if (modal) modal.hide();
  });
})();
</script>
{% endblock %}''')

NOTICES_HTML = BASE_HTML.replace('{% block content %}{% endblock %}','''
{% block content %}
<div class="container-fluid" style="max-width:700px;">
  <h4 class="mb-3 fw-bold text-primary">📢 通知公告</h4>

  <div class="card shadow-sm mb-4">
    <div class="card-header bg-light fw-bold">发布通知</div>
    <div class="card-body">
      <form method="post" action="/admin/notices/add">
        <div class="mb-2"><input class="form-control form-control-sm" name="title" maxlength="100" required placeholder="通知标题"></div>
        <div class="mb-2"><textarea class="form-control form-control-sm" name="content" rows="3" required placeholder="通知内容..."></textarea></div>
        <div class="d-flex gap-2">
          <select class="form-select form-select-sm" name="target" style="width:auto;">
            <option value="全部">发往全部</option>
            <option value="班组">发往班组</option>
            <option value="计调">发往计调</option>
          </select>
          <button class="btn btn-primary btn-sm px-4">发布</button>
        </div>
      </form>
    </div>
  </div>

  <div class="card shadow-sm">
    <div class="card-header bg-light fw-bold">已发布通知 ({{ notices|length }})</div>
    <div class="card-body p-0">
    {% for n in notices %}
      <div class="border-bottom p-3">
        <div class="d-flex justify-content-between align-items-start">
          <div>
            <span class="fw-bold">{{ n.title }}</span>
            <span class="badge {% if n.target == '全部' %}bg-success{% elif n.target == '班组' %}bg-warning text-dark{% else %}bg-info{% endif %} ms-1">{{ n.target }}</span>
            <small class="text-muted ms-2">{{ fmt_date(n.created_time) }} · {{ (n.user.username or n.user.account) if n.user else '' }}</small>
          </div>
          <form method="post" action="/admin/notices/delete/{{ n.id }}" onsubmit="return confirm('确认删除？')" style="display:inline;">
            <button class="btn btn-sm btn-outline-danger py-0">删除</button>
          </form>
        </div>
        <div class="mt-2 text-muted" style="white-space:pre-wrap;">{{ n.content }}</div>
      </div>
    {% else %}
      <div class="text-center text-muted py-4">暂无通知</div>
    {% endfor %}
    </div>
  </div>
</div>
{% endblock %}''')

# ---------- 初始化 ----------
with app.app_context():
    db.session.execute(text("PRAGMA journal_mode=DELETE"))
    db.session.execute(text("PRAGMA synchronous=FULL"))
    import shutil, glob as _glob
    backup_dir = os.path.join(EXE_DIR, 'backups')
    os.makedirs(backup_dir, exist_ok=True)
    db_path = os.path.join(EXE_DIR, 'data.db')
    def do_backup():
        if os.path.exists(db_path):
            ts = bj_now().strftime('%Y%m%d_%H%M%S')
            fname = os.path.join(backup_dir, f'data_backup_{ts}.db')
            shutil.copy2(db_path, fname)
            cutoff = (bj_now() - timedelta(days=7)).strftime('%Y%m%d_%H%M%S')
            for old in _glob.glob(os.path.join(backup_dir, 'data_backup_*.db')):
                try:
                    bn = os.path.basename(old).replace('data_backup_','').replace('.db','')
                    if bn < cutoff:
                        os.remove(old)
                except:
                    pass
    do_backup()
    def backup_timer():
        while True:
            import time
            time.sleep(1800)
            do_backup()
    threading.Thread(target=backup_timer, daemon=True).start()
    try:
        db.session.execute(text("SELECT user_id FROM column_permissions LIMIT 1"))
    except Exception as e:
        db.session.rollback()
        db.session.execute(text("DROP TABLE IF EXISTS column_permissions"))
        db.create_all()
        admin_user = User.query.filter_by(username='admin').first()
        if admin_user:
            for col,_ in COLUMN_ORDER:
                db.session.add(ColumnPermission(user_id=admin_user.id, column_name=col, can_view=True, can_edit=True))
            db.session.commit()

    try:
        db.session.execute(text("SELECT product_category FROM production_tasks LIMIT 1"))
    except Exception as e:
        db.session.rollback()
        db.session.execute(text("ALTER TABLE production_tasks ADD COLUMN product_category VARCHAR(50)"))
        db.session.commit()

    try:
        db.session.execute(text("SELECT stage FROM production_tasks LIMIT 1"))
    except Exception as e:
        db.session.rollback()
        db.session.execute(text("ALTER TABLE production_tasks ADD COLUMN stage VARCHAR(50)"))
        db.session.commit()

    db.create_all()

    try:
        db.session.execute(text("SELECT account FROM users LIMIT 1"))
    except Exception as e:
        db.session.rollback()
        db.session.execute(text("ALTER TABLE users ADD COLUMN account VARCHAR(50)"))
        db.session.commit()
        db.session.execute(text("UPDATE users SET account = username WHERE account IS NULL"))
        db.session.commit()

    try:
        db.session.execute(text("SELECT team FROM users LIMIT 1"))
    except Exception as e:
        db.session.rollback()
        db.session.execute(text("ALTER TABLE users ADD COLUMN team VARCHAR(50)"))
        db.session.commit()

    try:
        db.session.execute(text("SELECT production_line FROM users LIMIT 1"))
    except Exception as e:
        db.session.rollback()
        db.session.execute(text("ALTER TABLE users ADD COLUMN production_line VARCHAR(50)"))
        db.session.commit()

    admin_user = User.query.filter_by(username='admin').first()
    if admin_user:
        existing = {p.column_name for p in ColumnPermission.query.filter_by(user_id=admin_user.id).all()}
        for col, _ in COLUMN_ORDER:
            if col not in existing:
                db.session.add(ColumnPermission(user_id=admin_user.id, column_name=col, can_view=True, can_edit=True))
        if len(existing) != len(COLUMN_ORDER):
            db.session.commit()
    try:
        db.session.execute(text("SELECT production_time FROM production_tasks LIMIT 1"))
    except Exception as e:
        db.session.rollback()
        db.session.execute(text("ALTER TABLE production_tasks ADD COLUMN production_time DATETIME"))
        db.session.commit()

    try:
        db.session.execute(text("SELECT operator FROM production_tasks LIMIT 1"))
    except Exception as e:
        db.session.rollback()
        db.session.execute(text("ALTER TABLE production_tasks ADD COLUMN operator VARCHAR(200)"))
        db.session.commit()

    try:
        db.session.execute(text("SELECT operation_type FROM operation_logs LIMIT 1"))
    except Exception as e:
        db.session.rollback()
        db.session.execute(text("ALTER TABLE operation_logs ADD COLUMN operation_type VARCHAR(50)"))
        db.session.commit()

    try:
        db.session.execute(text("UPDATE operation_logs SET operation_type = '编辑字段' WHERE operation_type IS NULL"))
        db.session.commit()
    except:
        db.session.rollback()

    try:
        db.session.execute(text("SELECT target FROM branch_notices LIMIT 1"))
    except Exception as e:
        db.session.rollback()
        db.session.execute(text("ALTER TABLE branch_notices ADD COLUMN target VARCHAR(20) DEFAULT '全部'"))
        db.session.commit()

    # Ensure base roles exist, delete any others
    base_roles = ['管理员','部门领导','计划员','班组长','工艺员','计调员','检验员']
    for rname in base_roles:
        if not Role.query.filter_by(name=rname).first():
            db.session.add(Role(name=rname))
    db.session.commit()
    # Clean up truly obsolete roles (keep all base_roles)
    for r in Role.query.all():
        if r.name not in base_roles:
            fallback = Role.query.filter_by(name='计划员').first()
            if fallback:
                for u in User.query.filter_by(role_id=r.id).all():
                    u.role_id = fallback.id
            db.session.delete(r)
    db.session.commit()
    admin_role = Role.query.filter_by(name='管理员').first()
    if admin_role and not User.query.filter_by(account='admin').first():
        new_admin = User(account='admin', username='admin', password_hash=generate_password_hash('123456'), role_id=admin_role.id)
        db.session.add(new_admin)
        db.session.commit()
        for col,_ in COLUMN_ORDER:
            db.session.add(ColumnPermission(user_id=new_admin.id, column_name=col, can_view=True, can_edit=True))
        db.session.commit()

if __name__ == '__main__':
    app.run(debug=False, host='0.0.0.0', port=5000)
